from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
import build_book2_full_audit as base


ROOT = Path("/Users/rosslauda/Documents/HAL")
SRC = ROOT / "Documents/Book II/markdown/HAL_Book_II_Final_Constitutional_Certification_Audit_v1.0.md"
OUT = ROOT / "Documents/Book II/DOC/HAL_Book_II_Final_Constitutional_Certification_Audit_v1.0.docx"
ORIGINAL_SETUP_SECTION = base.setup_section


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    r = p.add_run("HAL")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(base.TEAL)

    p = doc.add_paragraph(style="Title")
    p.add_run("Book II Final Certification")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Constitutional and Cross-Chapter Architecture Audit")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    base.add_rule(p, base.TEAL, "30")

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    metadata = [
        ("Version", "1.0"),
        ("Date", "July 27, 2026"),
        ("Scope", "Book II Chapters 1–35"),
        ("Authority", "Book I Constitution v1.0 and Principles v1.1"),
        ("Disposition", "Certified constitutionally conformant"),
        ("Owner Review", "Closed; no unresolved decisions"),
    ]
    for i, (key, value) in enumerate(metadata):
        table.cell(i, 0).text = key.upper()
        table.cell(i, 1).text = value
        base.set_cell_shading(table.cell(i, 0), base.NAVY)
        table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        base.set_cell_shading(table.cell(i, 1), base.LIGHT_GRAY if i % 2 == 0 else "FFFFFF")
        for cell in table.rows[i].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            base.set_cell_margins(cell)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BOOK I REMAINS THE SUPREME AUTHORITY")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(base.RED)
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("Contents", level=1)
    items = [
        ("Executive Disposition", "Certification conclusion"),
        ("1. Certification Basis", "Authority and review method"),
        ("2. Authoritative Chapter Set", "Controlling editions"),
        ("3. Constitutional Coverage", "Decisions 1–58"),
        ("4. Closed Owner Decisions", "Final adopted rules"),
        ("5. Resolved Cross-Chapter Findings", "Canonical ownership corrections"),
        ("6. Certification Tests", "Final pass results"),
        ("7. Certification Boundary", "What this certification covers"),
        ("8. Final Certification Statement", "Formal disposition"),
        ("Document Control", "Revision history"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Shading Accent 1"
    table.cell(0, 0).text = "SECTION"
    table.cell(0, 1).text = "PURPOSE"
    for left, right in items:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    for row in table.rows:
        for cell in row.cells:
            base.set_cell_margins(cell)
    doc.add_paragraph(
        "This record certifies the architecture document set. It does not certify a future implementation or deployment."
    ).italic = True
    doc.add_page_break()


def setup_section(section):
    ORIGINAL_SETUP_SECTION(section)
    footer = section.footer.paragraphs[0]
    footer.text = "Final constitutional architecture certification  •  v1.0"
    footer.style = "Caption"
    base.add_page_number(footer)


def add_revision_history(doc):
    # The Markdown already contains the final Document Control table.
    return


def build():
    base.SRC = SRC
    base.OUT = OUT
    base.add_title_page = add_title_page
    base.add_contents = add_contents
    base.setup_section = setup_section
    base.add_revision_history = add_revision_history
    base.build()
    doc = Document(OUT)
    doc.core_properties.title = "HAL Book II — Final Constitutional Certification Audit"
    doc.core_properties.subject = "Final certification of Book II Chapters 1–35 against Book I"
    doc.core_properties.author = "HAL Architecture Review"
    doc.core_properties.keywords = "HAL, Book II, architecture, constitution, final certification"
    doc.save(OUT)


if __name__ == "__main__":
    build()
