from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/rosslauda/Documents/HAL")
SRC = ROOT / "Documents/Book II/markdown/HAL_Book_II_Full_Constitutional_and_Owner_Decision_Audit_v1.0.md"
OUT = ROOT / "Documents/Book II/DOC/HAL_Book_II_Full_Constitutional_and_Owner_Decision_Audit_v1.0.docx"

NAVY = "17324D"
BLUE = "2F648C"
TEAL = "3D7B78"
LIGHT_BLUE = "EAF1F6"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D5DDE3"
DARK = RGBColor(31, 42, 51)
RED = "9E3B3B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("  •  ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*.+?\*\*|`.+?`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Aptos Mono"
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(37, 78, 104)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


def add_rule(paragraph, color=NAVY, width="24"):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), width)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 30, NAVY),
        ("Subtitle", 13, BLUE),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 11.5, TEAL),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name != "Subtitle" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 18)
        style.paragraph_format.space_after = Pt(6)

    if "Audit Metadata" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("Audit Metadata", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Aptos"
        st.font.size = Pt(9.5)
        st.font.color.rgb = RGBColor(65, 80, 90)
        st.paragraph_format.space_after = Pt(3)


def setup_section(section):
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    header = section.header.paragraphs[0]
    header.text = ""
    footer = section.footer.paragraphs[0]
    footer.text = "Independent constitutional and owner-decision review  •  v1.0"
    footer.style = "Caption"
    add_page_number(footer)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.add_run("HAL").bold = True
    p.runs[0].font.size = Pt(15)
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph(style="Title")
    p.add_run("Book II Architecture Audit")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Full Constitutional, Cross-Chapter, and Owner-Decision Review")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    add_rule(p, TEAL, "30")

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.45)
    table.columns[1].width = Inches(5.45)
    metadata = [
        ("Version", "1.0"),
        ("Date", "July 27, 2026"),
        ("Scope", "Book II Chapters 1–35"),
        ("Authority", "Book I Constitution v1.0 and Principles v1.1"),
        ("Disposition", "Conditionally conformant; Owner review required"),
        ("Prepared for", "Constitutional Owner"),
    ]
    for i, (k, v) in enumerate(metadata):
        table.cell(i, 0).text = k.upper()
        table.cell(i, 1).text = v
        set_cell_shading(table.cell(i, 0), NAVY)
        table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        set_cell_shading(table.cell(i, 1), LIGHT_GRAY if i % 2 == 0 else "FFFFFF")
        for c in table.rows[i].cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BOOK I REMAINS THE SUPREME AUTHORITY")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(RED)
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("Contents", level=1)
    items = [
        ("Executive Disposition", "Audit conclusion and counts"),
        ("1. Audit Method", "Authority, editions, and tests"),
        ("2. Constitutional Coverage", "Coverage across all 58 decisions"),
        ("3. Owner Decisions Required", "Three matters reserved to the Owner"),
        ("4. Constitutional Conflict and Tension Register", "Blocking and material findings"),
        ("5. Cross-Chapter Inconsistency Register", "Engineering corrections"),
        ("6. Terminology and Editorial Findings", "Consolidation cleanup"),
        ("7. Certification Readiness", "Required path to certification"),
        ("8. Owner Review Summary", "Compact decision register"),
        ("9. Final Audit Statement", "Independent disposition"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Shading Accent 1"
    table.cell(0, 0).text = "SECTION"
    table.cell(0, 1).text = "PURPOSE"
    for left, right in items:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    doc.add_paragraph(
        "This audit is a decision and conformance document. It is not a replacement for Book I or for the Book II chapters."
    ).italic = True
    doc.add_page_break()


def parse_markdown(doc, source):
    lines = source.splitlines()
    # Skip title and document-control metadata; those are represented on the title page.
    start = next(i for i, line in enumerate(lines) if line.startswith("## Executive Disposition"))
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            headers = [c.strip() for c in line.strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for j, h in enumerate(headers):
                table.cell(0, j).text = h
                set_cell_shading(table.cell(0, j), NAVY)
                for r in table.cell(0, j).paragraphs[0].runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.font.bold = True
            for ridx, row in enumerate(rows):
                cells = table.add_row().cells
                for j in range(len(headers)):
                    value = row[j] if j < len(row) else ""
                    cells[j].text = value
                    set_cell_shading(cells[j], "FFFFFF" if ridx % 2 == 0 else LIGHT_GRAY)
                    set_cell_margins(cells[j], 70, 90, 70, 90)
                    cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            doc.add_paragraph()
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            if line[4:].startswith("IC-06"):
                doc.add_page_break()
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif re.match(r"^\d+\.\s", line):
            match = re.match(r"^(\d+\.)\s+(.*)$", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.22)
            marker = p.add_run(match.group(1) + "  ")
            marker.bold = True
            add_inline(p, match.group(2))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1


def add_revision_history(doc):
    doc.add_heading("Document Control", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["VERSION", "DATE", "CHANGE", "STATUS"]
    vals = ["1.0", "July 27, 2026", "Initial independent whole-book audit", "Owner review required"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        set_cell_shading(table.cell(0, j), NAVY)
        table.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.cell(0, j).paragraphs[0].runs[0].font.bold = True
        table.cell(1, j).text = vals[j]
        set_cell_margins(table.cell(0, j))
        set_cell_margins(table.cell(1, j))


def build():
    doc = Document()
    configure_styles(doc)
    setup_section(doc.sections[0])
    add_title_page(doc)
    add_contents(doc)
    parse_markdown(doc, SRC.read_text())
    add_revision_history(doc)
    doc.core_properties.title = "HAL Book II — Full Constitutional and Owner-Decision Audit"
    doc.core_properties.subject = "Independent audit of Book II Chapters 1–35 against Book I"
    doc.core_properties.author = "HAL Architecture Review"
    doc.core_properties.keywords = "HAL, Book II, architecture, constitution, audit, Owner review"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
