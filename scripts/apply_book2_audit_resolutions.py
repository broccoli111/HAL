from pathlib import Path
from copy import deepcopy
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


ROOT = Path("/Users/rosslauda/Documents/HAL")
BOOK = ROOT / "Documents/Book II"
DOC = BOOK / "DOC"
MD = BOOK / "markdown"


CHAPTERS = {
    1: {
        "old": "HAL_Book_II_Chapter_01_Overall_System_Architecture_v0.1",
        "new": "HAL_Book_II_Chapter_01_Overall_System_Architecture_v0.2",
        "replacements": [
            ("Identity & Authority Service", "Identity Service / Authority Service"),
            (
                "root identity, authorization, policy issuance and evaluation authority, protected audit rules, transaction commit coordination, cluster membership authority, and recovery control.",
                "root identity, protected authorization admission, protected policy validation, protected audit rules, transaction commit coordination, cluster membership authority, and recovery control.",
            ),
            (
                "The initial deployment uses a Mac mini as Primary Constitutional Host and HAL Core control plane;",
                "The initial deployment uses a Mac mini as the normal Primary Constitutional Host and active HAL Core control plane;",
            ),
        ],
        "resolutions": [
            "Identity and Authority are separate authoritative services within one logical constitutional domain; each owns only its declared records.",
            "The Authority Service evaluates ordinary effects and prepares protected decisions. The Constitutional Kernel independently validates and commits protected decisions.",
            "The persistence layer is a physical custodian, never the semantic owner of domain state.",
            "HAL has one active constitutional control plane. The normal Mac mini host may be replaced only through the governed lease and recovery process.",
        ],
    },
    3: {
        "old": "HAL_Book_II_Chapter_03_Constitutional_Kernel_v0.1",
        "new": "HAL_Book_II_Chapter_03_Constitutional_Kernel_v0.2",
        "replacements": [],
        "resolutions": [
            "The Authority Service may prepare protected decisions, but only the Constitutional Kernel may validate and commit them.",
            "A recovery constitutional lease is limited to 24 hours. Any extension requires fresh Owner authorization bound to the exact successor, constitutional state, and duration; expiry stops protected canonical mutation and returns the successor to Restricted or Safe Recovery mode.",
            "An invariant change requires two exact-change Owner Authorization Ceremonies separated by a 72-hour cooling-off period, with a Constitutional Mirror, independently verified recovery point, and explicit continuity classification.",
        ],
    },
    5: {
        "old": "HAL_Book_II_Chapter_05_Authority_and_Delegation_v0.1",
        "new": "HAL_Book_II_Chapter_05_Authority_and_Delegation_v0.2",
        "replacements": [
            (
                "capability-class activation, Treaty activation, protected deletion, and changes to trust/authentication requirements.",
                "capability-class activation, Treaty activation, cryptographic protected deletion, and changes to trust/authentication requirements.",
            ),
        ],
        "resolutions": [
            "Protected deletion never removes an Experience Object from the immutable ledger sequence.",
            "Authorized deletion may cryptographically erase a protected payload and remove non-authoritative copies, caches, indexes, and external replicas.",
            "A minimal non-sensitive tombstone remains with the authority, time, scope, and proof of deletion.",
            "Protected deletion requires the Owner Authorization Ceremony. Removal of the event identity or tombstone would require a constitutional amendment.",
        ],
    },
    8: {
        "old": "HAL_Book_II_Chapter_08_Attention_and_Resource_Allocation_v0.2",
        "new": "HAL_Book_II_Chapter_08_Attention_and_Resource_Allocation_v0.3",
        "replacements": [
            (
                "**Scheduler:** Policy-constrained assignment across feasible nodes and providers",
                "**Scheduler:** Sole work-admission and placement authority across feasible nodes and providers; consumes temporal constraints from Chapter 13",
            )
        ],
        "resolutions": [
            "The Chapter 8 Scheduler is the sole authority for work admission and placement.",
            "Temporal commitments, recurrence, windows, and deadlines are authoritative inputs from Chapter 13; they do not create a second placement authority.",
        ],
    },
    10: {
        "old": "HAL_Book_II_Chapter_10_Knowledge_Architecture_v0.2",
        "new": "HAL_Book_II_Chapter_10_Knowledge_Architecture_v0.3",
        "replacements": [],
        "resolutions": [
            "The Knowledge Service is the semantic owner of Knowledge Objects and relationships.",
            "Databases, replicas, indexes, embeddings, caches, and observability projections are physical or derived custodians and cannot become knowledge authority.",
        ],
    },
    12: {
        "old": "HAL_Book_II_Chapter_12_Memory_and_Experience_Ledger_v0.2",
        "new": "HAL_Book_II_Chapter_12_Memory_and_Experience_Ledger_v0.3",
        "replacements": [
            (
                "Apply visibility, retention and protected-deletion policies without silently rewriting the ledger.",
                "Apply visibility and retention policies; protected deletion may cryptographically erase payload access and remove derived copies but never remove the Experience Object or its tombstone from the ledger sequence.",
            ),
        ],
        "resolutions": [
            "The Conversation Service is the sole semantic owner of Conversation Objects. The Chapter 23 Thread Service maintains transport and delivery projections only.",
            "The Experience Ledger sequence is immutable. Protected deletion uses cryptographic erasure, deletes derived copies, and preserves a minimal non-sensitive tombstone.",
            "Removal of an Experience Object identity, sequence position, or tombstone is prohibited without a constitutional amendment.",
        ],
    },
    13: {
        "old": "HAL_Book_II_Chapter_13_Temporal_Architecture_v0.2",
        "new": "HAL_Book_II_Chapter_13_Temporal_Architecture_v0.3",
        "replacements": [
            (
                "**Scheduler:** Commitments, dependencies, windows, recurrence and deadlines",
                "**Temporal Commitment Service:** Authoritative commitments, dependencies, windows, recurrence and deadlines supplied to the Chapter 8 Scheduler",
            )
        ],
        "resolutions": [
            "The Temporal Service and Temporal Commitment Service own time facts and scheduling constraints.",
            "Only the Chapter 8 Scheduler admits and places work. Chapter 13 does not independently dispatch execution.",
        ],
    },
    14: {
        "old": "HAL_Book_II_Chapter_14_Presence_and_Human_Interaction_v0.2",
        "new": "HAL_Book_II_Chapter_14_Presence_and_Human_Interaction_v0.3",
        "replacements": [
            (
                "**Interaction Coordinator:** Modality, audience, handoff and response rendering",
                "**Presence Interaction Adapter:** Presence modality, audience, handoff and rendering context supplied to the canonical Interaction Session",
            ),
            (
                "**Interaction Session:** Participants, identity confidence, conversation, modality, disclosure scope and handoff token.",
                "**Interaction Session Reference:** Identifier and Presence/audience context linked to the canonical Interaction Session owned by Chapter 31.",
            ),
        ],
        "resolutions": [
            "The Chapter 31 Interaction Manager is the sole semantic owner of Interaction Sessions.",
            "Chapter 14 owns Presence Objects, Audience Context, and Presence Handoffs and contributes those facts to the canonical session.",
        ],
    },
    18: {
        "old": "HAL_Book_II_Chapter_18_Trust_Architecture_v0.2",
        "new": "HAL_Book_II_Chapter_18_Trust_Architecture_v0.3",
        "replacements": [],
        "resolutions": [
            "The Evidence Service is the semantic owner of Evidence Objects, custody, signatures, and verification state.",
            "Observability systems may produce evidence candidates and projections but cannot mutate evidentiary meaning or trust conclusions.",
            "The Audit Ledger owns protected action and access audit, not general Evidence Objects.",
        ],
    },
    19: {
        "old": "HAL_Book_II_Chapter_19_Privacy_and_Data_Governance_v0.2",
        "new": "HAL_Book_II_Chapter_19_Privacy_and_Data_Governance_v0.3",
        "replacements": [
            (
                "**Retention/Deletion Coordinator:** Policy lifecycle, legal/constitutional holds and protected deletion transactions",
                "**Retention/Deletion Coordinator:** Policy lifecycle, legal/constitutional holds, cryptographic payload erasure, derived-copy deletion and tombstone transactions",
            ),
            (
                "**Retention Disposition:** Keep, cool, archive, compress, expire access, or protected delete with evidence.",
                "**Retention Disposition:** Keep, cool, archive, compress, expire access, or cryptographically erase protected payload access while preserving an immutable tombstone.",
            ),
            (
                "Execute protected deletion only through authorized, auditable transactions and record lawful residual obligations.",
                "Execute protected deletion only through the Owner Authorization Ceremony and an auditable transaction: cryptographically erase the protected payload, delete non-authoritative copies, and preserve the immutable event sequence and minimal tombstone.",
            ),
        ],
        "resolutions": [
            "Protected deletion balances privacy with historical integrity by making authorized content unrecoverable while preserving proof that an event and lawful deletion occurred.",
            "No retention policy may silently remove an Experience Object identity, sequence position, or deletion tombstone.",
        ],
    },
    20: {
        "old": "HAL_Book_II_Chapter_20_Constitutional_Firewall_v0.2",
        "new": "HAL_Book_II_Chapter_20_Constitutional_Firewall_v0.3",
        "replacements": [
            (
                "**Treaty Manager:** Treaty identity, scope, lifecycle, renewal, revocation and evidence",
                "**Treaty Decision Consumer:** Verified Treaty state and decisions supplied by the sole Treaty Manager defined in Chapter 21",
            )
        ],
        "resolutions": [
            "Chapter 21's Treaty Manager is the sole owner of Treaty proposal, approval, lifecycle, and history.",
            "The Constitutional Firewall owns exchange enforcement, redaction, ingress/egress validation, and cross-domain incident records and consumes signed Treaty state.",
        ],
    },
    21: {
        "old": "HAL_Book_II_Chapter_21_External_Trust_Domains_and_Treaties_v0.1",
        "new": "HAL_Book_II_Chapter_21_External_Trust_Domains_and_Treaties_v0.2",
        "replacements": [
            (
                "**Treaty Manager:** Treaty proposal, protected approval, activation, renewal, suspension, revocation, and history",
                "**Treaty Manager:** Sole semantic owner of Treaty proposal, protected approval, activation, renewal, suspension, revocation, and history",
            )
        ],
        "resolutions": [
            "The Treaty Manager in this chapter is the sole authoritative owner of Treaty lifecycle state.",
            "The Chapter 20 Constitutional Firewall consumes signed Treaty decisions and cannot create, renew, or revoke a Treaty independently.",
        ],
    },
    22: {
        "old": "HAL_Book_II_Chapter_22_Distributed_Coordination_v0.1",
        "new": "HAL_Book_II_Chapter_22_Distributed_Coordination_v0.2",
        "replacements": [],
        "resolutions": [
            "A recovery constitutional lease may be issued only to a pre-registered, independently attested successor after the prior lease has provably expired and the required quorum agrees on the latest valid constitutional state.",
            "Every recovery lease has a hard maximum duration of 24 hours. Any extension requires fresh Owner authorization bound to the exact successor, constitutional state, lease identifier, and duration; expiry ends protected canonical mutation and returns the successor to Restricted or Safe Recovery mode. Permanent Primary Constitutional Host reassignment requires the Owner Authorization Ceremony.",
        ],
    },
    23: {
        "old": "HAL_Book_II_Chapter_23_Event_and_Messaging_Architecture_v0.1",
        "new": "HAL_Book_II_Chapter_23_Event_and_Messaging_Architecture_v0.2",
        "replacements": [
            (
                "**Thread Service:** Correlation, participants, conversation state, decisions, questions, and continuation",
                "**Thread Service:** Transport correlation, participant routing, delivery state, continuation handles, and projections of the canonical Conversation Object",
            ),
            (
                "**Conversation Object:** Purpose, participants, state, decisions, open questions, dependencies, confidence, and next actions.",
                "**Conversation Object Reference:** Identifier and delivery projection of the canonical Conversation Object owned by the Chapter 12 Conversation Service.",
            ),
        ],
        "resolutions": [
            "The Chapter 12 Conversation Service is the sole semantic owner of Conversation Objects.",
            "The Thread Service owns message/thread delivery mechanics only and references the canonical conversation rather than duplicating it.",
        ],
    },
    24: {
        "old": "HAL_Book_II_Chapter_24_State_and_Persistence_Architecture_v0.1",
        "new": "HAL_Book_II_Chapter_24_State_and_Persistence_Architecture_v0.2",
        "replacements": [
            (
                "**Authoritative State Store:** Constitution, identity, authority, policy, configuration, transactions, and other canonical records",
                "**Persistence Custodian:** Physical durability for Constitution, identity, authority, policy, configuration, transactions, and other canonical records; semantic authority remains with each domain service",
            ),
            (
                "**Event and Experience Ledgers:** Append-only history, causal linkage, integrity, and replay",
                "**Ledger Storage Custody:** Physical storage, integrity, causal linkage, and replay for append-only Event and Experience Ledgers owned by their domain services",
            ),
        ],
        "resolutions": [
            "Persistence components are physical custodians. They never become semantic authorities merely because they store, replicate, or restore a record.",
            "Experience Ledger, Evidence Service, Audit Ledger, Knowledge Service, and each domain service retain separate canonical responsibilities.",
            "Protected deletion may erase payload keys and derived copies but may not remove the Experience Object sequence or minimal tombstone.",
        ],
    },
    25: {
        "old": "HAL_Book_II_Chapter_25_Observability_and_Evidence_v0.1",
        "new": "HAL_Book_II_Chapter_25_Observability_and_Evidence_v0.2",
        "replacements": [],
        "resolutions": [
            "The Observability Fabric produces telemetry and evidence candidates; it is not the semantic owner of general Evidence Objects.",
            "The Chapter 18 Evidence Service owns Evidence Objects and custody. The Audit Ledger owns protected action/access audit records.",
            "Explanations and health views are derived projections and cannot become authoritative state.",
        ],
    },
    28: {
        "old": "HAL_Book_II_Chapter_28_Recovery_and_Continuity_v0.1",
        "new": "HAL_Book_II_Chapter_28_Recovery_and_Continuity_v0.2",
        "replacements": [],
        "resolutions": [
            "A successor may receive a recovery constitutional lease only after independent state, identity, integrity, and quorum verification.",
            "The lease expires after no more than 24 hours. Every extension requires fresh, exact, Owner-specific authorization.",
            "Without extension, recovery continues only in Restricted or Safe Recovery mode and protected canonical mutation stops.",
            "Permanent reassignment of the Primary Constitutional Host requires the Owner Authorization Ceremony.",
        ],
    },
    31: {
        "old": "HAL_Book_II_Chapter_31_Human_Interaction_Architecture_v0.1",
        "new": "HAL_Book_II_Chapter_31_Human_Interaction_Architecture_v0.2",
        "replacements": [
            (
                "**Interaction Manager:** Conversation state, modality, response mode, handoff, and interruption behavior",
                "**Interaction Manager:** Sole semantic owner of Interaction Sessions, conversation linkage, modality, response mode, handoff, and interruption behavior",
            )
        ],
        "resolutions": [
            "The Interaction Manager in this chapter is the sole semantic owner of Interaction Sessions.",
            "Presence, audience, modality, and handoff facts supplied by Chapter 14 are referenced as governed context rather than copied into a competing session record.",
        ],
    },
    33: {
        "old": "HAL_Book_II_Chapter_33_Constitutional_Evolution_Support_v0.1",
        "new": "HAL_Book_II_Chapter_33_Constitutional_Evolution_Support_v0.2",
        "replacements": [],
        "resolutions": [
            "An invariant change requires two separate Owner Authorization Ceremonies bound to the exact unchanged proposal.",
            "The ceremonies are separated by a mandatory 72-hour cooling-off period.",
            "Before the second ceremony, HAL creates a signed Constitutional Mirror and independently verified recovery point for the predecessor system.",
            "The final compatibility report explicitly classifies whether the result preserves HAL identity or creates a successor constitutional system.",
            "Any material proposal change restarts the complete process. The predecessor Constitution, state, authorization, and migration evidence are preserved permanently.",
            "Until the extraordinary process completes, invariant changes may be analyzed and simulated but never executed.",
        ],
    },
    34: {
        "old": "HAL_Book_II_Chapter_34_Deployment_Topologies_v0.1",
        "new": "HAL_Book_II_Chapter_34_Deployment_Topologies_v0.2",
        "replacements": [
            (
                "preserve one central constitutional control plane",
                "preserve one active constitutional control plane",
            ),
            (
                "Bootstrap the dedicated HAL Core control plane and establish the constitutional lease.",
                "Bootstrap the dedicated HAL Core control plane and establish the constitutional lease; a recovery successor lease is capped at 24 hours and every extension requires fresh Owner authorization.",
            ),
        ],
        "resolutions": [
            "The Mac mini is the normal Primary Constitutional Host, not an irreplaceable identity anchor.",
            "Only one active constitutional control plane may hold a valid lease.",
            "A pre-registered successor may hold a verified recovery lease for no more than 24 hours; extensions require fresh Owner authorization and permanent reassignment requires the Owner Authorization Ceremony.",
        ],
    },
    35: {
        "old": "HAL_Book_II_Chapter_35_Architecture_Conformance_and_Certification_v0.1",
        "new": "HAL_Book_II_Chapter_35_Architecture_Conformance_and_Certification_v0.2",
        "replacements": [
            (
                "**Certification Authority:** Scope, version, environment, result, limitations, expiry, and revocation",
                "**Architecture Certification Service:** Scope, version, environment, result, limitations, expiry, and revocation; certification is not constitutional authority",
            )
        ],
        "resolutions": [
            "Architecture certification is an evidence-backed conformance status and never a source of constitutional authority.",
            "Certification requires one semantic owner per authoritative object and explicit separation of domain authority from persistence and observability custody.",
            "Invariant-change conformance requires two exact Owner ceremonies, a 72-hour cooling-off period, a Constitutional Mirror, independent recovery proof, and continuity classification.",
            "Recovery-lease conformance requires a 24-hour maximum and fresh Owner authorization for every extension.",
            "Experience deletion conformance requires cryptographic payload erasure, derived-copy removal, and an immutable minimal tombstone.",
        ],
    },
}

# Chapters without a substantive conflict still receive a new edition so that
# embedded author self-approval is removed from the authoritative set.
CHAPTERS.update({
    2: {
        "old": "HAL_Book_II_Chapter_02_Runtime_Model_v0.1",
        "new": "HAL_Book_II_Chapter_02_Runtime_Model_v0.2",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    4: {
        "old": "HAL_Book_II_Chapter_04_Identity_and_Continuity_v0.1",
        "new": "HAL_Book_II_Chapter_04_Identity_and_Continuity_v0.2",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    6: {
        "old": "HAL_Book_II_Chapter_06_Intent_and_Planning_Architecture_v0.2",
        "new": "HAL_Book_II_Chapter_06_Intent_and_Planning_Architecture_v0.3",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    7: {
        "old": "HAL_Book_II_Chapter_07_Cognitive_Orchestration_v0.2",
        "new": "HAL_Book_II_Chapter_07_Cognitive_Orchestration_v0.3",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    9: {
        "old": "HAL_Book_II_Chapter_09_Judgment_and_Decision_Objects_v0.2",
        "new": "HAL_Book_II_Chapter_09_Judgment_and_Decision_Objects_v0.3",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    11: {
        "old": "HAL_Book_II_Chapter_11_Learning_and_Wisdom_v0.2",
        "new": "HAL_Book_II_Chapter_11_Learning_and_Wisdom_v0.3",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    15: {
        "old": "HAL_Book_II_Chapter_15_Capability_Architecture_v0.2",
        "new": "HAL_Book_II_Chapter_15_Capability_Architecture_v0.3",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    16: {
        "old": "HAL_Book_II_Chapter_16_Action_and_Transaction_Architecture_v0.2",
        "new": "HAL_Book_II_Chapter_16_Action_and_Transaction_Architecture_v0.3",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    17: {
        "old": "HAL_Book_II_Chapter_17_Verification_and_Simulation_v0.2",
        "new": "HAL_Book_II_Chapter_17_Verification_and_Simulation_v0.3",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    26: {
        "old": "HAL_Book_II_Chapter_26_Security_Architecture_v0.1",
        "new": "HAL_Book_II_Chapter_26_Security_Architecture_v0.2",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    27: {
        "old": "HAL_Book_II_Chapter_27_Failure_Containment_v0.1",
        "new": "HAL_Book_II_Chapter_27_Failure_Containment_v0.2",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    29: {
        "old": "HAL_Book_II_Chapter_29_Software_Lifecycle_and_Change_Governance_v0.1",
        "new": "HAL_Book_II_Chapter_29_Software_Lifecycle_and_Change_Governance_v0.2",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    30: {
        "old": "HAL_Book_II_Chapter_30_Self-Description_and_Constitutional_Mirror_v0.1",
        "new": "HAL_Book_II_Chapter_30_Self-Description_and_Constitutional_Mirror_v0.2",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
    32: {
        "old": "HAL_Book_II_Chapter_32_Outcome_and_Success_Architecture_v0.1",
        "new": "HAL_Book_II_Chapter_32_Outcome_and_Success_Architecture_v0.2",
        "replacements": [],
        "resolutions": [
            "Embedded author approval is not conformance evidence. This chapter's constitutional status is determined only by the independent whole-book audit.",
            "No substantive constitutional or cross-chapter correction was required in this edition.",
        ],
    },
})


def iter_block_items(doc):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def replace_paragraph_text(paragraph, replacements):
    original = paragraph.text
    updated = original
    for old, new in replacements:
        updated = updated.replace(old.replace("**", ""), new.replace("**", ""))
    if updated != original:
        paragraph.text = updated


def apply_docx(spec):
    src = DOC / f"{spec['old']}.docx"
    dst = DOC / f"{spec['new']}.docx"
    doc = Document(src)
    replacements = list(spec["replacements"])
    replacements.extend([
        ("Architect draft complete; constitutional and engineering review complete", "Revised after whole-book constitutional and cross-chapter audit"),
        ("Source-aligned rewrite; architecture audit passed", "Revised after whole-book constitutional and cross-chapter audit"),
        ("Source-aligned; architecture audit passed", "Revised after whole-book constitutional and cross-chapter audit"),
        ("None required at this stage", "No unresolved Owner decision after whole-book audit"),
        ("None required", "No unresolved Owner decision after whole-book audit"),
    ])
    old_ver = re.search(r"_v([0-9.]+)$", spec["old"]).group(1)
    new_ver = re.search(r"_v([0-9.]+)$", spec["new"]).group(1)
    replacements.extend([
        (f"Version\n{old_ver}", f"Version\n{new_ver}"),
        (f"Version: {old_ver}", f"Version: {new_ver}"),
        (f"Version {old_ver}", f"Version {new_ver}"),
    ])
    for p in doc.paragraphs:
        replace_paragraph_text(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_paragraph_text(p, replacements)
                if cell.text.strip() == old_ver:
                    cell.text = new_ver

    heading = doc.add_paragraph("Whole-Book Audit Resolution")
    heading.style = "Heading 2"
    heading.paragraph_format.keep_with_next = True
    intro = doc.add_paragraph(
        "The following requirements supersede any ambiguous earlier wording in this chapter and are incorporated into this edition:"
    )
    intro.paragraph_format.keep_with_next = True
    for item in spec["resolutions"]:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.keep_together = True
    doc.core_properties.version = new_ver
    doc.core_properties.comments = "Revised after full Book II constitutional and cross-chapter audit."
    doc.save(dst)
    return dst


def apply_markdown(spec):
    src = MD / f"{spec['old']}.md"
    if not src.exists():
        return None
    dst = MD / f"{spec['new']}.md"
    text = src.read_text()
    old_ver = re.search(r"_v([0-9.]+)$", spec["old"]).group(1)
    new_ver = re.search(r"_v([0-9.]+)$", spec["new"]).group(1)
    text = text.replace(f"**Version:** {old_ver}", f"**Version:** {new_ver}", 1)
    text = text.replace(
        "**Status:** Source-aligned rewrite; architecture audit passed",
        "**Status:** Revised after whole-book constitutional and cross-chapter audit",
        1,
    )
    text = text.replace(
        "**Status:** Source-aligned; architecture audit passed",
        "**Status:** Revised after whole-book constitutional and cross-chapter audit",
        1,
    )
    for old, new in spec["replacements"]:
        text = text.replace(old, new)
    section = "\n## Whole-Book Audit Resolution\n\n"
    section += (
        "The following requirements supersede any ambiguous earlier wording in this chapter "
        "and are incorporated into this edition:\n\n"
    )
    section += "\n".join(f"- {item}" for item in spec["resolutions"]) + "\n\n"
    marker = "\n## Source Alignment and Review"
    if marker in text:
        text = text.replace(marker, section + marker, 1)
    else:
        text += section
    dst.write_text(text)
    return dst


def docx_to_markdown(docx_path, md_path):
    doc = Document(docx_path)
    out = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name if block.style else ""
            if style == "Title":
                out.append(f"# {text}")
            elif style.startswith("Heading 1"):
                out.append(f"## {text}")
            elif style.startswith("Heading 2"):
                out.append(f"## {text}")
            elif style.startswith("Heading 3"):
                out.append(f"### {text}")
            elif style.startswith("List Bullet"):
                out.append(f"- {text}")
            elif style.startswith("List Number"):
                out.append(f"1. {text}")
            else:
                out.append(text)
            out.append("")
        else:
            rows = []
            for row in block.rows:
                rows.append([cell.text.replace("\n", " ").strip() for cell in row.cells])
            if not rows or not any(rows[0]):
                continue
            out.append("| " + " | ".join(rows[0]) + " |")
            out.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            for row in rows[1:]:
                out.append("| " + " | ".join(row) + " |")
            out.append("")
    md_path.write_text("\n".join(out).strip() + "\n")


def main():
    DOC.mkdir(parents=True, exist_ok=True)
    MD.mkdir(parents=True, exist_ok=True)
    for chapter, spec in CHAPTERS.items():
        out_doc = apply_docx(spec)
        out_md = apply_markdown(spec)
        if out_md is None:
            out_md = MD / f"{spec['new']}.md"
            docx_to_markdown(out_doc, out_md)
        print(f"{chapter:02}: {out_doc.name} | {out_md.name}")


if __name__ == "__main__":
    main()
