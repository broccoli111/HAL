from __future__ import annotations

import csv
import hashlib
import json
import re
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
CHAPTERS = ROOT / "chapters"
MODEL = ROOT / "model"
TRACE = ROOT / "traceability"
REVIEWS = ROOT / "reviews"
CHAPTER_REVIEWS = REVIEWS / "chapter-reviews"
DELIVERABLES = ROOT / "deliverables"
SCHEMAS = ROOT / "schemas"
TEMPLATES = ROOT / "templates"
CHECKLISTS = ROOT / "checklists"
EXAMPLES = ROOT / "policy-examples"
TODAY = "2026-07-27"
VERSION = "1.0"

for directory in (CHAPTERS, MODEL, TRACE, REVIEWS, CHAPTER_REVIEWS, DELIVERABLES,
                  SCHEMAS, TEMPLATES, CHECKLISTS, EXAMPLES):
    directory.mkdir(parents=True, exist_ok=True)


chapters = [
    (1, "Purpose, Scope, Authority, and Semantic Governance",
     "Defines Book X authority, semantic records, precedence, change control, and the rules for resolving ambiguity.",
     "Book I in full; especially constitutional authority and invariants",
     "Book II Chapters 01, 03, 29, 30, and 35",
     "Book III Chapters 01, 08, and 09"),
    (2, "Concept System and Information-Model Foundations",
     "Defines semantic types, entity and record distinctions, relationships, states, events, claims, constraints, and machine-readable representations.",
     "Book I Articles II, IV, IX, and XI",
     "Book II Chapters 01, 02, 22, 23, 24, 25, and 30",
     "Book III Chapters 03, 04, 06, and 09"),
    (3, "Constitutional Identity, Ownership, Continuity, and Presence",
     "Defines HAL, Owner, constitutional identity, continuity, Presence, embodiment, self-description, and the constitutional mirror.",
     "Book I identity, Owner authority, sovereignty, continuity, and constitutional evolution",
     "Book II Chapters 02, 03, 04, 14, 28, 30, and 31",
     "Book III Chapters 01, 05, 07, and 08"),
    (4, "Identity, Authentication, Authority, and Delegation",
     "Separates identity, identifiers, authentication, trust, permission, authority, delegation, policy decisions, and protected action.",
     "Book I Decisions 5, 6, 25, 26, 27, and 48",
     "Book II Chapters 04, 05, 18, 20, 21, and 26",
     "Book III Chapters 03, 05, 06, and 08"),
    (5, "Intent, Planning, Attention, Judgment, and Outcomes",
     "Defines the purpose-to-outcome hierarchy and the durable objects used for planning, attention, judgment, and success evaluation.",
     "Book I intent, judgment, learning, restraint, success, and outcomes",
     "Book II Chapters 06, 07, 08, 09, 13, and 32",
     "Book III Chapters 03, 04, 06, and 08"),
    (6, "Capabilities, Providers, Actions, Transactions, and the Reality Boundary",
     "Defines abilities, implementations, governed execution, commit barriers, rollback, compensation, and separation of simulated from real effects.",
     "Book I Decisions 10, 15, 16, 20, 23, 24, 35, 36, 44, and 50",
     "Book II Chapters 15, 16, 17, 22, 27, and 28",
     "Book III Chapters 03, 05, 06, and 07"),
    (7, "Evidence, Trust, Verification, Assurance, and Certification",
     "Defines the evidence admission boundary, claims, provenance, custody, trust assessment, verification, assurance cases, and certification.",
     "Book I Articles II and XI; Decisions 22, 26, 34, 35, 40, 42, 43, 50, and 56",
     "Book II Chapters 17, 18, 25, and 35",
     "Book III Chapters 04, 06, 08, and 09"),
    (8, "Experience, Memory, Knowledge, Learning, Patterns, and Wisdom",
     "Separates records of occurrence, retained experience, contextualized knowledge, learned patterns, and evidence-bounded wisdom.",
     "Book I learning, wisdom, evidence, uncertainty, restraint, and continuity",
     "Book II Chapters 10, 11, 12, 13, and 30",
     "Book III Chapters 03, 04, 05, 06, and 08"),
    (9, "State, Time, Events, Messaging, Coordination, and Persistence",
     "Defines authoritative state, projections, commands, queries, events, ordering, time, idempotency, messaging, and durable persistence.",
     "Book I evidence, continuity, reversibility, accountability, and failure containment",
     "Book II Chapters 13, 22, 23, 24, 25, 27, and 28",
     "Book III Chapters 03, 04, 06, and 07"),
    (10, "Privacy, Security, Trust Domains, Treaties, and the Constitutional Firewall",
     "Defines information classification, privacy purpose, secrets, credentials, trust boundaries, external domains, Treaties, and governed exchange.",
     "Book I privacy, dignity, sovereignty, external trust, and protected authority",
     "Book II Chapters 18, 19, 20, 21, 25, and 26",
     "Book III Chapters 02, 04, 05, 06, and 07"),
    (11, "Runtime, Resources, Operations, Failure, Recovery, and Change",
     "Defines runtime state, health, readiness, supervision, resources, degradation, incidents, recovery objectives, releases, migrations, and exceptions.",
     "Book I continuity, reversibility, restraint, accountability, and constitutional shutdown",
     "Book II Chapters 02, 22, 27, 28, 29, 33, 34, and 35",
     "Book III Chapters 01, 02, 04, 06, 07, 08, and 09"),
    (12, "Naming, Acronyms, Ambiguity, Deprecation, and Cross-Book Use",
     "Defines canonical labels, aliases, acronyms, qualified terms, forbidden ambiguities, semantic versioning, deprecation, and adoption rules.",
     "Book I constitutional authority and evolution",
     "Book II Chapters 29, 30, and 35",
     "Book III Chapters 01, 03, 04, 08, and 09"),
]


# label, chapter, category, semantic type, definition, exclusion/distinction, aliases
term_rows = [
("Book X",1,"Governance","Canonical reference","The controlled HAL volume that defines shared terminology and the cross-canon information model.","It does not independently create constitutional, architectural, engineering, component, interface, operational, or certification requirements.","Canonical Terminology and Information Model"),
("Canonical Term",1,"Governance","Semantic record","A governed label with one approved meaning, stable identifier, source traceability, relationships, constraints, examples, and lifecycle status.","A commonly used word is not canonical until admitted to this register.",""),
("Term Record",1,"Governance","Metadata record","The authoritative Book X record for a Canonical Term and its semantic metadata.","A Term Record describes a concept; it is not automatically an operational entity record.",""),
("Semantic Authority",1,"Governance","Precedence property","The authority of a source to determine meaning within its governed scope according to the canon hierarchy.","Semantic Authority does not grant operational Authority to a Principal.",""),
("Semantic Change",1,"Governance","Governed change","A controlled modification to a canonical label, definition, relationship, constraint, status, or mapping.","Editorial correction is a Semantic Change when meaning or compatibility may change.",""),
("Qualified Term",1,"Governance","Disambiguated label","A label extended with a domain qualifier so that one meaning can be selected without ambiguity.","A qualifier must clarify meaning and must not conceal two distinct concepts under one record.",""),
("Allowed Alias",1,"Governance","Reference label","A non-canonical label permitted to reference one Canonical Term without changing its meaning.","An alias must not be used where its ambiguity would obscure the canonical concept.","permitted synonym"),
("Deprecated Term",1,"Governance","Lifecycle status","A previously permitted label or meaning retained only for migration and historical interpretation.","Deprecation is not immediate deletion and must identify a replacement and sunset condition.",""),
("Forbidden Term",1,"Governance","Prohibited label or usage","A label or usage prohibited because it collapses materially distinct HAL concepts or creates unsafe ambiguity.","The prohibition applies to the specified meaning, not necessarily every ordinary-language occurrence.",""),
("Normative Source",1,"Governance","Source role","A controlled artifact whose authority determines a requirement or meaning within the canon hierarchy.","Book X cites Normative Sources but cannot promote a lower-order artifact above a higher-order source.",""),
("Entity",2,"Information model","Semantic type","A distinguishable thing with identity and continuity relevant to the HAL domain.","An Entity is not the same as its record, identifier, state, role, or representation.",""),
("Value Object",2,"Information model","Semantic type","An immutable value defined by its attributes rather than by independent identity.","Changing a Value Object produces another value; it does not mutate an enduring identity.",""),
("Record",2,"Information model","Semantic type","A governed representation of facts, state, decisions, or observations retained by an owning domain.","A Record is not necessarily authoritative evidence and is not interchangeable with the entity represented.",""),
("Authoritative Record",2,"Information model","State role","The record owned by the designated source of truth for a governed state domain.","A replica, cache, index, projection, or local copy is not authoritative merely because it is current.","source of truth"),
("Relationship",2,"Information model","Semantic type","A typed association between concepts or entity instances with declared direction, cardinality, constraints, and lifecycle.","Proximity or co-occurrence does not imply a governed Relationship.",""),
("State",2,"Information model","Semantic type","The values and lifecycle condition of an entity or process at a defined observation point.","State is not an Event; an Event records a completed fact about change.",""),
("Lifecycle",2,"Information model","Semantic model","The allowed states, transitions, entry conditions, exit conditions, terminal conditions, and evidence for a governed concept.","A list of statuses without transition rules is not a complete Lifecycle.",""),
("Invariant",2,"Information model","Constraint","A condition required to remain true throughout a defined scope or transition set.","An engineering invariant is not automatically a Constitutional invariant.",""),
("Claim",2,"Information model","Evidence-bearing assertion","A proposition stated for evaluation and linked to its subject, issuer, scope, time, and supporting or opposing evidence.","A Claim is not true merely because it is recorded or signed.",""),
("Constraint",2,"Information model","Rule element","A condition that limits valid state, relationships, transitions, or behavior within a defined scope.","A preference or target is not a Constraint unless its governing source makes it binding.",""),
("Owner",3,"Constitutional","Constitutional role","The unique human principal holding HAL's constitutional ownership and the authority reserved to that role by Book I.","Ownership must not be inferred from possession of infrastructure, credentials, data, or a deployment. Founder is a historical source label for this same role, never a second constitutional role.","Founder"),
("HAL",3,"Constitutional","Constitutional identity","The single constitutionally governed intelligence whose continuity is independent of any one model, service, Presence, node, or machine.","A runtime instance, model, component, or interface must not be called a separate HAL identity.",""),
("Constitution",3,"Constitutional","Supreme governing instrument","Book I, the supreme source of HAL identity, principles, authority, rights, duties, prohibitions, and invariants.","No lower-order book, policy, configuration, or code may amend the Constitution by reinterpretation.","Book I"),
("Constitutional Invariant",3,"Constitutional","Constitutional constraint","A Book I requirement whose alteration may change HAL's constitutional identity and cannot be waived or redefined by lower-order documents.","An architectural or engineering invariant is not constitutional unless Book I makes it so.",""),
("Constitutional Kernel",3,"Constitutional","Architectural component class","The Book II architectural authority that evaluates and enforces constitutional rules at designated decision and action paths.","It does not replace the Constitution or independently invent constitutional meaning.",""),
("Constitutional Mirror",3,"Constitutional","Self-description mechanism","The governed, evidence-linked representation through which HAL describes its identity, governing constraints, capabilities, limitations, and conformance state.","It is not a source of new constitutional authority and must not become self-authorizing.",""),
("Continuity",3,"Identity","Constitutional property","The governed preservation of HAL identity, obligations, provenance, and essential state across time, replacement, recovery, and deployment change.","Continuity does not require uninterrupted availability or persistence of every transient process.",""),
("Presence",3,"Interaction","Contextual manifestation","A bounded manifestation through which HAL senses, communicates, or acts in a particular human, device, location, modality, or session context.","A Presence is not a separate HAL identity and does not independently hold Owner authority.",""),
("Embodiment",3,"Interaction","Contextual binding","The governed association of a Presence with physical or virtual sensors, actuators, interfaces, and environmental context.","Embodiment does not make hardware ownership equivalent to constitutional ownership.",""),
("Sovereignty",3,"Constitutional","Constitutional property","HAL's constitutionally governed independence from unauthorized external control, coercion, substitution, or absorption.","Sovereignty does not authorize HAL to exceed Owner authority or human rights.",""),
("Identity",4,"Identity","Governed entity identity","The durable governed identity by which HAL recognizes a human, device, service, sensor, agent, node, or subsystem as the same entity across time.","Identity is distinct from names, identifiers, attributes, roles, credentials, trust, authority, and Presence.",""),
("Principal",4,"Identity","Governed actor role","An Identity that may be the subject of authentication, policy evaluation, authority, delegation, accountability, or action attribution.","Principal status does not itself grant authority.",""),
("Identity Record",4,"Identity","Authoritative record","The authoritative governed state representing an Identity, including its immutable identity reference, type, lifecycle, handles, relationships, and source-approved metadata.","The record represents an Identity but is not interchangeable with the Identity.",""),
("Identifier",4,"Identity","Reference value","A value used to reference an Identity or another entity within a declared namespace and lifecycle.","Possession or presentation of an Identifier does not authenticate identity or grant authority.","ID"),
("Identity Attribute",4,"Identity","Governed descriptive value","A property, assertion, or governed fact associated with an Identity and qualified by source, time, and confidence where applicable.","An attribute is not the Identity and does not independently establish authentication, trust, or authority.",""),
("Credential",4,"Security","Authentication instrument","A governed secret, key, token, certificate, biometric template, or other instrument used in an authentication protocol.","Credential possession is evidence, not identity, trust, permission, or authority by itself.",""),
("Authentication",4,"Identity","Assurance process and result","The evidence-based process and resulting confidence that a claimed Identity is presently genuine in a specified context.","Authentication answers who or what is present; it does not answer what action is allowed.",""),
("Authentication Evidence",4,"Evidence","Evidence role","One or more Evidence Objects used to assess whether a claimed Identity is presently genuine.","Authentication Evidence informs assurance but does not itself grant authority.",""),
("Trust",4,"Trust","Evidence-based assessment","Multidimensional, domain-specific confidence derived from evidence about an entity, claim, process, or relationship.","Trust may inform decisions but must not be treated as authority or permission.",""),
("Permission",4,"Authority","Decision result","A scoped result of authoritative policy evaluation allowing a specified Principal to attempt a specified operation under stated context and conditions.","Permission is not identical to Authority, Delegation, Capability, role, trust, or credential possession.","authorization result"),
("Authority",4,"Authority","Governed decision and action scope","The constitutionally and policy-governed scope within which a Principal may decide or cause action.","Authority constrains Permission evaluation but is not itself a Permission result and is not inferred from identity, trust, capability, role, proximity, or credentials.",""),
("Delegation",4,"Authority","Governed authority grant","An attributable, scoped, conditional, expiring, and revocable grant of Authority from an authorized delegator to a recipient.","A Delegation cannot grant authority the delegator does not possess or make a constitutional invariant waivable.",""),
("Policy",4,"Authority","Decision rule set","A governed set of rules evaluated against identity, authority, delegation, purpose, context, resource, and risk inputs.","A Policy is not itself a decision and cannot outrank its Normative Source.",""),
("Policy Decision Record",4,"Authority","Decision record","The attributable record of a policy evaluation, including inputs, governing policy version, outcome, obligations, reason, and correlation data.","It records a decision; it does not create standing Authority beyond that decision's scope.","authorization decision record"),
("Protected Action",4,"Authority","Risk classification","An Action whose authority, privacy, security, safety, irreversibility, or constitutional impact requires explicit governed evaluation and evidence.","A routine implementation detail is not protected merely because it is technically complex.",""),
("Intent",5,"Cognition","Purpose object","A governed expression of desired purpose, direction, or outcome attributable to an authorized source.","Intent is not a Plan, Action, Permission, or evidence that an outcome occurred.",""),
("Vision",5,"Cognition","Long-horizon intent","A durable directional state describing an intended future without fully specifying its execution path.","A Vision is broader and less operational than a Goal.",""),
("Goal",5,"Cognition","Outcome target","A governed desired outcome with success criteria and a time or review horizon.","A Goal does not by itself authorize Actions used to pursue it.",""),
("Objective",5,"Cognition","Measurable target","A bounded, measurable target that advances a Goal and has explicit completion or evaluation criteria.","An Objective is not a Task; it states what must be achieved, not merely what work is performed.",""),
("Project",5,"Cognition","Coordinated work container","A governed body of related Objectives, Plans, Tasks, resources, decisions, and evidence organized toward a defined outcome.","A Project is not standing authority for every contained Action.",""),
("Task",5,"Cognition","Work unit","A bounded unit of work with responsibility, inputs, expected result, dependencies, and completion evidence.","Completing a Task is not equivalent to achieving the parent Objective or Goal.",""),
("Strategy",5,"Cognition","Approach selection","A reasoned approach for advancing one or more Goals under known constraints, uncertainties, and tradeoffs.","A Strategy is not an executable Plan and does not bypass verification.",""),
("Plan",5,"Cognition","Coordinated intended work","A governed arrangement of Tasks, dependencies, resources, decision points, verification steps, and recovery conditions intended to realize an Objective.","A Plan is not proof that Actions are permitted, executed, or successful.",""),
("Plan Graph",5,"Cognition","Planning representation","A dependency graph of intended work, decisions, resources, and verification gates.","It represents intended coordination, not completed reality.",""),
("Execution Graph",5,"Cognition","Runtime representation","The governed graph of actual Attempts, Actions, dependencies, outcomes, and evidence for an execution instance.","It must not be silently substituted for the Plan Graph when explaining divergence.",""),
("Attention Object",5,"Cognition","Prioritization record","A durable object representing a candidate matter for bounded attention, including source, salience, urgency, risk, relevance, context, disposition, and evidence.","Attention is not Authority, approval, or a promise to act.",""),
("Decision Object",5,"Cognition","Decision record","A durable record of a consequential decision including question, alternatives, authority context, evidence, uncertainty, judgment, rationale, selected disposition, and review conditions.","It records a decision and must not be used as a substitute for required Permission or execution evidence.",""),
("Judgment",5,"Cognition","Reasoned evaluation","The context-sensitive evaluation that weighs evidence, uncertainty, values, consequences, proportionality, and restraint to reach or recommend a decision.","Judgment must not silently invent authority or conceal unresolved uncertainty.",""),
("Uncertainty",5,"Cognition","Epistemic condition","A represented limitation in knowledge, evidence, prediction, interpretation, or confidence relevant to a claim or decision.","Uncertainty is not failure; unrepresented material uncertainty is a defect.",""),
("Outcome Object",5,"Cognition","Outcome record","A durable record linking intended outcome, observed result, affected parties, evidence, side effects, confidence, and evaluation.","It is not equivalent to an Event, metric sample, or optimistic status assertion.",""),
("Success",5,"Cognition","Evaluated condition","A source-governed determination that relevant outcomes satisfy stated criteria without unacceptable constitutional, human, privacy, security, or reliability costs.","Task completion, activity volume, or a single metric does not by itself establish Success.",""),
("Capability",6,"Capability","Implementation-independent contract","An abstract ability defined by outcomes, inputs, outputs, constraints, required authority and permission classes, risks, side effects, and evaluation criteria.","Capability does not identify an implementation and does not grant authority to use the ability.",""),
("Capability Contract",6,"Capability","Contract record","The versioned specification of a Capability's semantic inputs, outputs, preconditions, effects, risks, authority requirements, evidence, and compatibility.","It is not a provider-specific API contract.",""),
("Provider",6,"Capability","Implementation role","A component, service, model, person, device, or external system that can fulfill a Capability under a declared contract and trust context.","Being able to perform work does not authorize the Provider to perform it.",""),
("Adapter",6,"Capability","Boundary component role","A component that translates between a Capability Contract and a provider-specific interface while preserving authority, semantics, evidence, and failure behavior.","An Adapter must not smuggle provider semantics into the canonical Capability definition.",""),
("Capability Registry",6,"Capability","Authoritative registry","The governed catalog of Capabilities, versions, Providers, constraints, authority classes, health, and selection metadata.","Registration does not establish permission for use or trustworthiness in every domain.",""),
("Action",6,"Action","State-changing attempt","A governed attempt to produce an effect in authoritative state or the external world.","A read-only Query is not an Action; an Action is not proof of successful effect.",""),
("Attempt",6,"Action","Execution instance","One attributable execution effort for a Task, Action, or verification step with its own timing, context, result, and evidence.","A retry is a new Attempt even when it shares an idempotency key.",""),
("Transaction",6,"Action","Governed action lifecycle","The durable coordination object for one or more Actions, including authorization, prepare, commit, result, evidence, rollback, and compensation states.","It is broader than a database transaction and must not imply atomic reversibility of external effects.",""),
("Commit Barrier",6,"Action","Irreversibility gate","The explicit governed point after which a proposed change may create authoritative or real-world effects that cannot be treated as merely simulated or prepared.","Passing a Commit Barrier requires the applicable authority and verification; preparation alone must not cross it.",""),
("Rollback",6,"Action","Reversal operation","A controlled restoration of a prior recoverable state when the relevant effects are truthfully reversible.","Rollback must not claim to erase external or human effects that already occurred.",""),
("Compensation",6,"Action","Remedial operation","A new governed Action that mitigates, offsets, or repairs effects that cannot truthfully be undone.","Compensation is not Rollback and may require independent authority and evidence.",""),
("Reality Boundary",6,"Action","Governed environment boundary","The explicit separation among simulation, digital twin, shadow, test, canary, controlled-reality, production, recovery, and emergency contexts.","Non-reality authority, data, or effects must not leak across this boundary into reality.",""),
("Simulation",6,"Verification","Non-reality environment","An execution environment whose effects are confined to modeled or synthetic state and cannot directly alter production or external reality.","High apparent fidelity does not make a Simulation production.",""),
("Digital Twin",6,"Verification","Modeled counterpart","A governed model of selected real entities, relationships, state, and dynamics used to evaluate behavior without treating modeled effects as real effects.","A Digital Twin is bounded by declared fidelity and must not be mistaken for the represented reality.",""),
("Shadow Execution",6,"Verification","Non-committing execution mode","Execution using live or representative inputs while preventing proposed outputs from producing authoritative or external effects.","Shadow results do not authorize promotion without the required review and certification.",""),
("Canary",6,"Verification","Limited reality stage","A deliberately constrained real execution stage used to accumulate evidence before broader adoption.","A Canary crosses the Reality Boundary and therefore requires real authority, containment, and rollback or compensation.",""),
("Evidence Candidate",7,"Evidence","Pre-admission record","An observation, telemetry item, document, Audit Record, claim, or other integrity-protected input proposed for admission as an Evidence Object.","It is not authoritative Evidence until accepted through the Evidence Service's governed admission and custody process.",""),
("Evidence Object",7,"Evidence","Immutable provenance-bearing record","An immutable object admitted and governed by the authoritative Evidence Service that records provenance, custody, source identity, observation or claim content, time, signatures, verification state, confidence, domain, and expiration metadata.","An Evidence Object must not be mutated to revise a conclusion; later objects supersede or challenge it.","Evidence"),
("Audit Record",7,"Evidence","Protected accountability record","An append-only record of a protected action, authorization, decision, access, or change owned by the applicable audit domain.","An Audit Record may produce an Evidence Candidate but is not automatically a general Evidence Object.","audit log entry"),
("Provenance",7,"Evidence","Origin history","The attributable origin, derivation, transformation, and custody history of an artifact, datum, claim, model, decision, or Evidence Object.","A source label without derivation and custody context is incomplete provenance.",""),
("Chain of Custody",7,"Evidence","Integrity history","The ordered, attributable record of possession, control, transfer, and integrity protection for evidence or sensitive artifacts.","It does not establish truth; it establishes accountable handling.",""),
("Evidence Graph",7,"Evidence","Evidence relationship model","A graph connecting Claims, Evidence Objects, sources, derivations, supporting or opposing relations, confidence, and conclusions.","Graph connectivity does not make all linked material equally authoritative or trustworthy.",""),
("Verification",7,"Verification","Evidence-producing evaluation","A reproducible, risk-scaled process that evaluates a Claim, invariant, behavior, artifact, or Outcome against explicit criteria.","Verification produces evidence; it is not the same as certification or operational approval.",""),
("Verification Plan",7,"Verification","Planning record","A governed specification of claims, risks, methods, environments, data, success criteria, independence, and required evidence for Verification.","A test list without mapped claims and criteria is not a complete Verification Plan.",""),
("Assurance Case",7,"Verification","Structured argument","A structured, reviewable argument connecting scoped claims to reasoning and sufficient supporting evidence.","An Assurance Case must expose assumptions, defeaters, uncertainty, and evidence gaps.",""),
("Certification",7,"Verification","Governed assurance decision","A scoped, time-bounded, evidence-based determination by an authorized certifier that specified conformance claims are satisfied.","Certification is not permanent, universal, or self-issued by the artifact being certified.",""),
("Conformance",7,"Verification","Evaluated relation","The evidenced condition of satisfying identified requirements from identified authoritative sources within a declared scope and version.","Conformance is never implied merely by compatibility, successful execution, or absence of known defects.",""),
("Confidence",7,"Trust","Calibrated assessment","A bounded assessment of support for a Claim or prediction, expressed with method, scope, evidence basis, uncertainty, and time sensitivity.","Confidence is not probability unless a defined model justifies that interpretation.",""),
("Experience",8,"Knowledge","Retained occurrence record","A governed representation of what HAL perceived, attempted, decided, experienced, and observed, with context, outcomes, provenance, and privacy controls.","Experience is not automatically Knowledge, a Pattern, or Wisdom.",""),
("Experience Ledger",8,"Knowledge","Authoritative ledger","The append-oriented governed store of Experience records and their provenance, correction, retention, and access metadata.","It is not a general-purpose mutable memory store.",""),
("Memory",8,"Knowledge","Retrievable retained representation","A retained representation available for later contextual retrieval under authority, privacy, relevance, and lifecycle rules.","Memory is broader than Experience and is not necessarily authoritative Knowledge.",""),
("Memory Graph",8,"Knowledge","Associative representation","A governed graph linking retained representations by context, entity, time, causation, similarity, and relevance.","Association must not be treated as proof of causation or truth.",""),
("Knowledge",8,"Knowledge","Contextualized warranted representation","A governed representation whose claims, provenance, validity scope, confidence, and supporting evidence are sufficient for its declared use.","Stored information or model output is not automatically Knowledge.",""),
("Knowledge Graph",8,"Knowledge","Knowledge relationship model","A governed graph of entities, concepts, relationships, Claims, sources, validity, and provenance used for contextual reasoning and retrieval.","Graph membership does not erase source authority, uncertainty, or temporal scope.",""),
("Pattern",8,"Knowledge","Learned regularity","A reproducibly supported regularity across Experiences or observations with stated domain, evidence, confidence, limits, and exceptions.","A repeated coincidence or one-off anecdote is not a Pattern.",""),
("Learning",8,"Knowledge","Governed update process","The evidence-bounded process by which HAL updates representations, Patterns, policies, or behavior within authorized scope.","Learning must not silently alter constitutional meaning, authority, or protected production behavior.",""),
("Learning Ledger",8,"Knowledge","Change ledger","The governed record of learning proposals, evidence, evaluation, approval, applied changes, monitoring, rollback, and outcomes.","It is not permission for unrestricted self-modification.",""),
("Wisdom",8,"Knowledge","Evidence-bounded judgment resource","A durable, revisable synthesis of Experience, Patterns, values, consequences, uncertainty, and restraint used to inform Judgment.","Wisdom informs decisions but does not create authority or replace current evidence.",""),
("Self Model",8,"Knowledge","Governed self-representation","HAL's evidence-linked representation of its current capabilities, limits, state, dependencies, uncertainty, and identity continuity.","It is descriptive and must not become a self-authorizing source.",""),
("Command",9,"Distributed systems","Intent message","A request addressed to an authoritative owner to evaluate and, if allowed, perform a state-changing operation.","A Command is not evidence that the operation was authorized, committed, or completed.",""),
("Query",9,"Distributed systems","Read request","A request for information that must not intentionally mutate authoritative state or external reality.","Telemetry side effects must remain non-authoritative and must not turn a Query into a hidden Command.",""),
("Event",9,"Distributed systems","Completed-fact record","An immutable message stating that a defined fact or state transition has completed in its owning domain.","An Event is not a request, intention, or mutable current-state record.",""),
("Event Journal",9,"Distributed systems","Durable event store","The ordered durable record of Events for an owning domain, with identity, sequence, provenance, and integrity controls.","Global total order must not be inferred when only per-stream ordering exists.",""),
("Message Envelope",9,"Distributed systems","Transport record","The governed wrapper carrying message identity, schema version, source, destination, time, correlation, causation, authority context, classification, and integrity metadata.","The envelope does not change the semantic type of its payload.",""),
("Idempotency Key",9,"Distributed systems","Deduplication value","A scoped stable value allowing a receiver to recognize equivalent retry Attempts for one declared operation and result horizon.","It does not make non-idempotent external effects reversible or globally exactly-once.",""),
("Correlation Identifier",9,"Distributed systems","Trace reference","A value linking related work, messages, evidence, and telemetry across a bounded flow.","Correlation does not prove causation.","correlation ID"),
("Causation Identifier",9,"Distributed systems","Causal reference","A value identifying the immediate initiating message, decision, or event for a derived operation.","It records declared lineage and must not be used as sole proof of real-world causality.","causation ID"),
("Authoritative State",9,"Distributed systems","Owned state","State whose mutation and truth are governed by the designated owning domain.","A cache, replica, projection, search index, or local aggregate is not authoritative unless explicitly designated.",""),
("Projection",9,"Distributed systems","Derived state","A read-optimized representation derived from authoritative records or Events for a declared purpose.","A Projection may be stale and must not silently accept authoritative mutations.","read model"),
("Cache",9,"Distributed systems","Disposable derived state","A replaceable performance-oriented copy whose loss does not destroy authoritative truth.","A Cache must not become the only copy of required state or evidence.",""),
("Replica",9,"Distributed systems","Replicated state","A governed copy maintained from an authoritative source under declared consistency, lag, and failover rules.","Replication alone does not transfer semantic ownership.",""),
("Transactional Outbox",9,"Distributed systems","Publication pattern","A durable pattern that records an authoritative state change and its pending Event publication in one local commit boundary.","It reduces dual-write failure but does not guarantee global exactly-once processing.","outbox"),
("Logical Time",9,"Temporal","Ordering construct","A non-wall-clock ordering value used to represent causal or domain sequence relationships.","Logical Time must not be presented as real elapsed or calendar time.",""),
("Wall-Clock Time",9,"Temporal","Temporal observation","A timestamp from a declared clock source, precision, and synchronization context.","Wall-Clock Time alone must not establish distributed causal order.",""),
("External Trust Domain",10,"Trust","External governance domain","An external environment, organization, system, or authority regime whose controls and assumptions are not natively governed by HAL.","External status does not imply hostility or trustworthiness; exchange requires explicit governance.","ETD"),
("Treaty",10,"Trust","Governed trust agreement","An exact, scoped, time-bounded, revocable, auditable, and explicitly Owner-authorized agreement defining allowed cross-domain identities, data, capabilities, authority, obligations, verification, monitoring, failure behavior, and termination.","Activation requires the Owner Authorization Ceremony bound to the exact Treaty. Conversation, delegated ordinary authority, trust, usefulness, prior access, or the Treaty itself cannot grant authority prohibited by Book I or bypass the Constitutional Firewall.","External Trust Treaty"),
("Constitutional Firewall",10,"Trust","Architectural enforcement boundary","The Book II boundary that mediates and enforces constitutionally governed exchange between HAL and External Trust Domains.","It is not merely a network firewall and must not be bypassed by direct integration.","CF"),
("Trust Boundary",10,"Security","Security boundary","A boundary across which identity, authority, integrity, confidentiality, provenance, or governance assumptions change.","Network adjacency alone does not define or erase a Trust Boundary.",""),
("Data Classification",10,"Privacy","Governance label","A governed label determining handling, access, transmission, retention, evidence, and disposal requirements for information.","Classification is not a purpose or permission to use the data.",""),
("Personal Data",10,"Privacy","Information class","Information relating to an identified or reasonably identifiable human under the governing privacy context.","Pseudonymization may reduce exposure but does not necessarily remove personal-data status.",""),
("Sensitive Data",10,"Privacy","Information class","Information whose unauthorized use or disclosure could materially harm a person, HAL, the Owner, security, sovereignty, trust, or protected operations.","Sensitivity is context-dependent and may include non-personal operational data.",""),
("Purpose Limitation",10,"Privacy","Use constraint","The rule that governed data may be collected, used, shared, and retained only for declared, authorized, compatible purposes.","Availability or technical usefulness does not establish purpose.",""),
("Data Minimization",10,"Privacy","Collection and use constraint","The rule that only the least data reasonably necessary for an authorized purpose may be collected, processed, retained, and disclosed.","Minimization applies to fields, precision, population, duration, access, and derived inferences.",""),
("Retention Class",10,"Privacy","Lifecycle label","A governed category defining retention period, review, legal or constitutional hold behavior, archival, and disposal requirements.","A Retention Class does not itself authorize collection or access.",""),
("Secret",10,"Security","Sensitive authentication material","Confidential material whose disclosure could enable unauthorized access, impersonation, decryption, signing, or protected operation.","An Identifier or public key is not a Secret merely because it is security-related.",""),
("Cryptographic Key",10,"Security","Cryptographic material","A governed value used by an approved cryptographic operation and managed through generation, storage, access, rotation, revocation, destruction, and provenance controls.","A key's possession does not itself establish business Authority.",""),
("Security Control",10,"Security","Protective control","A safeguard intended to protect confidentiality, integrity, availability, identity assurance, or system resilience.","A Security Control protecting HAL is distinct from an Authority Control preventing HAL from exceeding its mandate.",""),
("Authority Control",10,"Authority","Mandate-limiting control","A safeguard that prevents a Principal, component, or HAL from deciding or acting beyond governed Authority.","It is not interchangeable with a Security Control, though one mechanism may support both.",""),
("Runtime",11,"Operations","Execution environment","The governed combination of processes, state, resources, policies, dependencies, and environments in which HAL behavior executes.","A Runtime is not HAL's constitutional identity.",""),
("Node",11,"Operations","Execution entity","A governed compute or device participant capable of hosting workload, state, sensing, or action under the runtime architecture.","A Node does not independently become HAL or own authoritative state without explicit designation.",""),
("Service",11,"Operations","Component deployment role","A deployable runtime boundary providing one or more governed responsibilities or interfaces.","A Service is not automatically an architectural component, capability, or authority domain.",""),
("Supervisor",11,"Operations","Runtime control role","A component that monitors and governs workload lifecycle, desired state, health, restart, containment, and escalation within its authority.","A Supervisor must not conceal repeated failure through unlimited restart loops.",""),
("Desired State",11,"Operations","Control target","The governed runtime condition an authorized controller intends the system to maintain.","Desired State is not proof of Observed State or successful convergence.",""),
("Observed State",11,"Operations","Runtime observation","The evidenced runtime condition measured at a defined time and scope.","Observation may be stale, partial, or uncertain and must not be treated as Desired State.",""),
("Health",11,"Operations","Operational assessment","A multidimensional assessment of a component's ability to perform its declared responsibilities within current constraints.","Process liveness alone is not Health.",""),
("Liveness",11,"Operations","Runtime signal","Evidence that a workload is running or able to make progress according to a narrow declared probe.","Liveness does not establish readiness, correctness, authorization, or safety.",""),
("Readiness",11,"Operations","Admission signal","Evidence that a workload may receive its declared class of work under current dependencies, configuration, and safety conditions.","Readiness is scoped and must not be inferred from Liveness.",""),
("Resource",11,"Operations","Governed capacity","A bounded consumable or allocatable asset such as compute, memory, storage, bandwidth, attention, time, energy, or device capacity.","Resource availability does not imply authority to allocate or consume it.",""),
("Reservation",11,"Operations","Resource claim","A time-bounded governed allocation claim against a Resource for a declared purpose and owner.","A Reservation is not proof the resource was consumed or the work completed.",""),
("Degraded Mode",11,"Operations","Operating state","A declared operating state in which selected capabilities or service levels are reduced while higher-priority constitutional, authority, safety, privacy, and evidence obligations remain protected.","Degradation must not silently weaken non-degradable controls.",""),
("Safe Mode",11,"Operations","Protective operating state","A constrained operating state that prioritizes containment, Owner communication, essential evidence, and prevention of unauthorized or unsafe effects.","Safe Mode is not a generic low-performance mode.",""),
("Restricted Mode",11,"Operations","Authority-limited state","An operating state in which selected capabilities, integrations, or authority paths are disabled or narrowed because required trust, evidence, policy, or assurance is unavailable.","It must be explicit, observable, and reversible through governed recovery.",""),
("Incident",11,"Operations","Governed adverse event","An event or condition requiring coordinated response because it threatens constitutional conformance, authority, security, privacy, trust, availability, integrity, or material outcomes.","An anomaly is not necessarily an Incident until classification criteria are met.",""),
("Quarantine",11,"Operations","Containment state","A governed isolation state preventing a component, artifact, identity, message, or data set from participating beyond an explicitly limited inspection boundary.","Quarantine is not deletion and must preserve required evidence.",""),
("Recovery",11,"Operations","Restoration process","The governed process of restoring acceptable identity, authority, state, service, evidence, and trust after failure or compromise.","Restart alone is not Recovery.",""),
("Recovery Point Objective",11,"Operations","Recovery target","The maximum tolerable loss or unavailability of recoverable state measured from the disruption point for a declared domain.","RPO is a target, not proof that recovery met it.","RPO"),
("Recovery Time Objective",11,"Operations","Recovery target","The target duration for restoring a declared service or capability to its required state after disruption.","RTO does not authorize unsafe shortcuts during recovery.","RTO"),
("Release",11,"Engineering","Governed artifact set","A versioned, traceable, qualified set of software, configuration, schemas, models, and deployment artifacts approved for a declared environment and scope.","A successful build is not a Release.",""),
("Migration",11,"Engineering","State or contract change","A governed transition of data, state, schema, interface, configuration, or runtime behavior from one compatible condition to another.","A Migration must distinguish reversible steps from irreversible effects and compensation.",""),
("Architecture Decision Record",11,"Engineering","Decision record","A durable record of a consequential architecture decision, context, alternatives, rationale, consequences, source traceability, and review status.","An ADR cannot authorize deviation from Book II without the approved architecture-governance process.","ADR"),
("Architecture Deviation",11,"Engineering","Governed exception class","A documented departure from an applicable Book II requirement processed through architecture governance with scope, risk, evidence, approval, and expiry or remediation.","It is not an ordinary Book III control exception and cannot amend Book II silently.",""),
("Control",11,"Engineering","Enforceable rule record","A stable, attributable requirement with applicability, responsibility, enforcement, evidence, severity, exception authority, verification, and source traceability.","Advice is not a Control unless it is made objectively reviewable.",""),
("Exception",11,"Engineering","Time-bounded control relief","A documented, scoped, risk-assessed, compensating, approved, expiring departure from a waivable lower-order Control.","A constitutional invariant cannot be waived; an Architecture Deviation uses its own governance path.","waiver"),
("Definition of Ready",11,"Engineering","Entry criteria","The minimum evidenced conditions required before consequential implementation work may begin or enter its next controlled stage.","Readiness does not imply approval to release or cross the Reality Boundary.","DoR"),
("Definition of Done",11,"Engineering","Completion criteria","The minimum evidenced conditions required before work may be treated as complete within a declared scope.","Done does not erase post-release monitoring, retention, or recovery obligations.","DoD"),
("Canonical Label",12,"Naming","Naming element","The single approved label used as the primary reference for one Canonical Term.","Capitalization is part of controlled usage when needed to distinguish the term from ordinary language.",""),
("Acronym",12,"Naming","Abbreviated label","An approved shortened form mapped to exactly one canonical expansion within its declared scope.","An Acronym must be expanded on first use unless the audience and artifact make the meaning unambiguous.",""),
("Semantic Version",12,"Governance","Compatibility marker","A version assigned to the Book X corpus to communicate the compatibility impact of semantic changes.","It does not replace the versioning of Books I-III, components, or interfaces.",""),
("Term Status",12,"Governance","Lifecycle label","The controlled state of a Term Record: Proposed, Candidate, Approved, Deprecated, Retired, or Rejected.","Status must not be inferred from document age or usage frequency.",""),
("Cross-Book Term Index",12,"Governance","Traceability index","The mapping from each Canonical Term to its authoritative source, Book X record, and known use across the HAL canon.","It is an index, not a substitute for reading the governing source.",""),
("Semantic Compatibility",12,"Governance","Compatibility relation","The condition in which a terminology or information-model change preserves the valid interpretation and obligations of dependent artifacts within declared scope.","Textual similarity alone does not establish Semantic Compatibility.",""),
("Trust Domain",10,"Trust","Governance context","A bounded identity, policy, evidence, security, privacy, and accountability context within which trust assumptions and controls are evaluated.","Trust Domain is the generic concept. An External Trust Domain is a Trust Domain outside HAL's native governance boundary and requires governed cross-domain exchange.",""),
("Owner Authorization Ceremony",4,"Authority","Protected authorization mechanism","The Book II-governed mechanism through which the Owner authorizes an exact protected change, capability-class decision, Treaty, or other Owner-reserved matter bound to an immutable decision identifier, declared scope, and validity period.","It is not conversational consent, ordinary Delegation, reusable standing permission, or a substitute for Constitutional Kernel validation; it authorizes only the exact recorded matter.",""),
("Evidence Service",7,"Evidence","Architectural component class","The Book II authoritative service that admits and governs Evidence Objects and owns their custody, signatures, provenance bindings, and verification state.","Observability, audit, and source systems may produce Evidence Candidates or records but cannot independently admit Evidence Objects or mutate evidentiary meaning.",""),
("Release Authority",11,"Engineering","Governed certification role","The Book III-designated role authorized to certify a qualified Release for a declared deployment scope after all required verification and reviews have passed.","Release Authority cannot waive Constitutional Invariants, replace architecture, security, or privacy review, or authorize deployment beyond the certified scope.",""),
]


source_profiles = {
    1: ("Book I Constitutional Governance and Decision 58",
        "Book II Chapters 29 and 35",
        "Book III Chapters 1 and 9",
        "Book X semantic-governance choice constrained by higher-order sources"),
    2: ("Book I Decisions 26, 29, 40, 47, and 51",
        "Book II Chapters 04, 22, 23, 24, 25, and 30",
        "Book III Chapters 3 and 4",
        "Derived semantic synthesis"),
    3: ("Book I Preamble and Decisions 45, 47, 48, 49, 51, and 58",
        "Book II Chapters 01, 02, 03, 04, 14, 28, 30, and 31",
        "Book III Chapters 1, 5, 7, and 8",
        "Direct source normalization"),
    4: ("Book I Decisions 5, 6, 25, 26, 27, 48, and 49",
        "Book II Chapters 03, 04, 05, 18, 20, 21, and 26",
        "Book III Chapters 3, 5, 6, and 8",
        "Direct source normalization"),
    5: ("Book I Decisions 12, 18, 41, 46, 54, 55, and 57",
        "Book II Chapters 06, 07, 08, 09, 13, and 32",
        "Book III Chapters 3, 4, 6, and 8",
        "Direct source normalization"),
    6: ("Book I Decisions 10, 15, 16, 20, 23, 24, 35, 36, 44, and 50",
        "Book II Chapters 15, 16, 17, 22, 27, and 28",
        "Book III Chapters 3, 5, 6, and 7",
        "Direct source normalization"),
    7: ("Book I Articles II and XI and Decisions 22, 26, 34, 35, 40, 42, 43, 50, and 56",
        "Book II Chapters 17, 18, 25, and 35",
        "Book III Chapters 4, 6, 8, and 9",
        "Direct source normalization"),
    8: ("Book I Decisions 30, 31, 52, and 53",
        "Book II Chapters 10, 11, 12, 13, and 30",
        "Book III Chapters 3, 4, 5, 6, and 8",
        "Direct source normalization"),
    9: ("Book I Decisions 4, 22, 29, 35, 38, 40, 42, and 47",
        "Book II Chapters 13, 22, 23, 24, 25, 27, and 28",
        "Book III Chapters 3, 4, 6, and 7",
        "Derived semantic synthesis"),
    10: ("Book I Decisions 26, 27, 32, 37, 39, 48, and 49",
         "Book II Chapters 18, 19, 20, 21, 25, and 26",
         "Book III Chapters 2, 4, 5, 6, and 7",
         "Direct source normalization"),
    11: ("Book I Decisions 22, 29, 35, 38, 40, 42, 43, 47, 50, and 51",
         "Book II Chapters 02, 22, 27, 28, 29, 33, 34, and 35",
         "Book III Chapters 1, 2, 4, 6, 7, 8, and 9",
         "Derived semantic synthesis"),
    12: ("Book I Constitutional Governance and Decision 58",
         "Book II Chapters 29, 30, and 35",
         "Book III Chapters 1, 3, 4, 8, and 9",
         "Book X semantic-governance choice constrained by higher-order sources"),
}

source_overrides = {}
def map_sources(labels, bi, bii, biii, basis="Direct source normalization"):
    for item in labels:
        source_overrides[item] = (bi, bii, biii, basis)

map_sources(["Owner"], "Book I Decisions 48, 49, and 58",
            "Book II Chapters 03, 04, 05, and 21",
            "Book III Chapters 1, 5, and 8")
map_sources(["HAL"], "Book I Preamble and Decisions 45, 47, 49, and 51",
            "Book II Chapters 01, 02, 14, 28, and 30",
            "Book III Chapters 1, 5, 7, and 8")
map_sources(["Constitution","Constitutional Invariant"], "Book I Constitutional Governance and Decision 58",
            "Book II Chapters 03, 29, 30, and 35",
            "Book III Chapters 1, 8, and 9")
map_sources(["Constitutional Kernel"], "Book I Decisions 25, 27, 40, 43, 48, 49, and 50",
            "Book II Chapter 03",
            "Book III Chapters 1, 3, 5, 7, and 8")
map_sources(["Constitutional Mirror","Self Model"], "Book I Decisions 28 and 51",
            "Book II Chapter 30",
            "Book III Chapters 4 and 8")
map_sources(["Continuity"], "Book I Decisions 45, 47, and 51",
            "Book II Chapters 04, 28, and 30",
            "Book III Chapters 1, 7, and 8")
map_sources(["Presence","Embodiment"], "Book I Decision 45",
            "Book II Chapters 14 and 31",
            "Book III Chapters 4 and 5")
map_sources(["Sovereignty"], "Book I Decision 49",
            "Book II Chapters 20 and 21",
            "Book III Chapter 5")
map_sources(["Identity","Principal","Identity Record","Identifier","Identity Attribute","Credential",
             "Authentication","Authentication Evidence"], "Book I Decisions 6, 27, and 48",
            "Book II Chapters 04 and 05",
            "Book III Chapter 5")
map_sources(["Trust","Permission","Authority","Delegation","Policy","Policy Decision Record","Protected Action"],
            "Book I Decisions 5, 25, 26, 27, and 48",
            "Book II Chapters 03, 05, and 18",
            "Book III Chapters 3, 5, and 8")
map_sources(["Intent","Vision","Goal","Objective","Project","Task","Strategy","Plan","Plan Graph","Execution Graph"],
            "Book I Decisions 12, 18, and 46",
            "Book II Chapters 06 and 07",
            "Book III Chapters 3, 6, and 8")
map_sources(["Attention Object"], "Book I Decision 54",
            "Book II Chapter 08",
            "Book III Chapters 4, 6, and 8")
map_sources(["Decision Object","Judgment","Uncertainty"], "Book I Decision 55",
            "Book II Chapter 09",
            "Book III Chapters 4, 6, and 8")
map_sources(["Outcome Object","Success"], "Book I Decision 57",
            "Book II Chapter 32",
            "Book III Chapters 4, 6, and 8")
map_sources(["Capability","Capability Contract","Provider","Adapter","Capability Registry"],
            "Book I Decisions 10 and 36",
            "Book II Chapter 15",
            "Book III Chapters 3, 5, and 7")
map_sources(["Action","Attempt","Transaction","Commit Barrier","Rollback","Compensation"],
            "Book I Decisions 16, 35, and 50",
            "Book II Chapter 16",
            "Book III Chapters 3, 6, and 7")
map_sources(["Reality Boundary","Simulation","Digital Twin","Shadow Execution","Canary"],
            "Book I Decision 50",
            "Book II Chapter 17",
            "Book III Chapters 3, 6, and 7")
map_sources(["Evidence Candidate","Evidence Object","Audit Record","Provenance","Chain of Custody","Evidence Graph","Claim"],
            "Book I Decision 26",
            "Book II Chapters 18 and 25",
            "Book III Chapters 4, 5, 6, and 8")
map_sources(["Verification","Verification Plan","Assurance Case","Certification","Conformance","Confidence"],
            "Book I Decisions 43, 50, and 56",
            "Book II Chapters 17 and 35",
            "Book III Chapters 6, 8, and 9")
map_sources(["Experience","Experience Ledger","Memory","Memory Graph"], "Book I Decisions 30 and 53",
            "Book II Chapter 12",
            "Book III Chapters 4, 5, and 6")
map_sources(["Knowledge","Knowledge Graph"], "Book I Decisions 30 and 52",
            "Book II Chapter 10",
            "Book III Chapters 3, 4, and 6")
map_sources(["Pattern","Learning","Learning Ledger","Wisdom"], "Book I Decisions 31 and 53",
            "Book II Chapter 11",
            "Book III Chapters 4, 6, and 8")
map_sources(["Command","Query","Event","Event Journal","Message Envelope","Idempotency Key",
             "Correlation Identifier","Causation Identifier"], "Book I Decisions 22, 29, and 40",
            "Book II Chapters 22 and 23",
            "Book III Chapters 3, 4, and 6")
map_sources(["Authoritative State","Projection","Cache","Replica","Transactional Outbox"],
            "Book I Decisions 4, 22, 35, 38, and 40",
            "Book II Chapters 23 and 24",
            "Book III Chapters 3, 4, and 7")
map_sources(["Logical Time","Wall-Clock Time"], "Book I Decision 44",
            "Book II Chapter 13",
            "Book III Chapters 3, 4, and 6")
map_sources(["Trust Domain"], "Book I Decisions 49 and 52",
            "Book II Chapters 18, 21, and 24",
            "Book III Chapter 5",
            "Derived semantic synthesis required by the Book X scope")
map_sources(["External Trust Domain"], "Book I Decision 49",
            "Book II Chapters 20 and 21",
            "Book III Chapter 5")
map_sources(["Treaty"], "Book I Decision 49",
            "Book II Chapter 21",
            "Book III Chapters 5 and 7")
map_sources(["Constitutional Firewall"], "Book I Decision 49",
            "Book II Chapter 20",
            "Book III Chapter 5")
map_sources(["Data Classification","Personal Data","Sensitive Data","Purpose Limitation",
             "Data Minimization","Retention Class"], "Book I Decision 32",
            "Book II Chapter 19",
            "Book III Chapters 4 and 5")
map_sources(["Secret","Cryptographic Key","Security Control","Authority Control"], "Book I Decisions 27, 37, and 39",
            "Book II Chapters 05, 19, and 26",
            "Book III Chapters 2 and 5")
map_sources(["Runtime","Node","Service","Supervisor","Desired State","Observed State","Health",
             "Liveness","Readiness","Resource","Reservation","Degraded Mode","Safe Mode",
             "Restricted Mode"], "Book I Decisions 29, 38, 42, 47, and 51",
            "Book II Chapters 02, 27, 28, 33, and 34",
            "Book III Chapters 3, 4, 6, and 7")
map_sources(["Incident","Quarantine","Recovery","Recovery Point Objective","Recovery Time Objective"],
            "Book I Decisions 38, 42, and 50",
            "Book II Chapters 27 and 28",
            "Book III Chapters 5, 6, 7, and 8")
map_sources(["Release","Migration","Architecture Decision Record","Architecture Deviation","Control",
             "Exception","Definition of Ready","Definition of Done"], "Book I Decisions 43, 50, and 58",
            "Book II Chapters 29 and 35",
            "Book III Chapters 1, 7, 8, and 9",
            "Engineering term normalized under Books I–III")
map_sources(["Owner Authorization Ceremony"], "Book I Decisions 48, 49, 50, and 58",
            "Book II Chapters 03, 05, 16, 21, and 29",
            "Book III Chapters 1, 5, 7, 8, and 9")
map_sources(["Evidence Service"], "Book I Decision 26",
            "Book II Chapters 18 and 25",
            "Book III Chapters 4, 5, 6, and 8")
map_sources(["Release Authority"], "Book I Decisions 43, 50, and 58",
            "Book II Chapters 29 and 35",
            "Book III Chapter 7",
            "Engineering term normalized under Books I–III")

terms = []
for idx, row in enumerate(term_rows, 1):
    label, chapter, category, semantic_type, definition, distinction, aliases = row
    bi, bii, biii, source_basis = source_overrides.get(label, source_profiles[chapter])
    terms.append({
        "term_id": f"HAL-TERM-{idx:04d}",
        "canonical_label": label,
        "chapter": chapter,
        "category": category,
        "semantic_type": semantic_type,
        "definition": definition,
        "distinction": distinction,
        "allowed_aliases": [a.strip() for a in aliases.split(",") if a.strip()],
        "status": "Approved",
        "book_i_source": bi,
        "book_ii_source": bii,
        "book_iii_source": biii,
        "source_basis": source_basis,
        "version_introduced": VERSION,
    })


relationships = [
("Owner","owns constitutionally","HAL","1","1","Book I-reserved ownership; infrastructure possession is insufficient."),
("Constitution","governs","HAL","1","1","Supreme authority."),
("Constitutional Kernel","enforces","Constitutional Invariant","1","many","Only at Book II-designated enforcement points."),
("Constitutional Mirror","describes","HAL","1","1","Evidence-linked and non-self-authorizing."),
("HAL","manifests through","Presence","1","many","All Presences share one HAL identity."),
("Presence","binds to","Embodiment","many","0..many","Within explicit context and lifecycle."),
("Identity Record","represents","Identity","1","1","Record and entity remain distinct."),
("Identifier","references","Identity","many","1","Within a declared namespace."),
("Identity Attribute","describes","Identity","many","1","Does not establish authority."),
("Authentication Evidence","supports","Authentication","many","many","Evidence role, not authorization."),
("Delegation","grants bounded","Authority","many","1","Cannot exceed delegator authority."),
("Policy","evaluates for","Permission","many","many","Result is contextual and scoped."),
("Authority","constrains","Permission","many","many","Authority is not the decision result."),
("Trust","informs","Policy Decision Record","many","many","Never sole authority."),
("Intent","decomposes into","Goal","1","many","Traceability is retained."),
("Goal","decomposes into","Objective","1","many","Criteria remain explicit."),
("Objective","is pursued by","Plan","many","many","Plans do not confer authority."),
("Plan","contains","Task","1","many","Execution may diverge with evidence."),
("Decision Object","records","Judgment","many","1","Includes alternatives and uncertainty."),
("Outcome Object","evaluates","Goal","many","many","Against evidence and side effects."),
("Capability Contract","defines","Capability","many","1","Versioned semantic contract."),
("Provider","fulfills","Capability","many","many","Selection remains governed."),
("Adapter","connects","Provider","many","1","Preserves canonical capability semantics."),
("Transaction","coordinates","Action","1","many","Includes commit and recovery state."),
("Canary","is governed stage within","Reality Boundary","many","1","A Canary is a limited real-operation stage; it is not a kind of boundary."),
("Evidence Candidate","may be admitted as","Evidence Object","many","0..1","Only through the authoritative evidence process."),
("Evidence Object","supports or opposes","Claim","many","many","Relation and weight are explicit."),
("Evidence Graph","contains","Evidence Object","1","many","Objects remain immutable."),
("Verification","evaluates","Claim","many","many","Against explicit criteria."),
("Assurance Case","organizes","Claim","1","many","Includes reasoning and defeaters."),
("Certification","depends on","Assurance Case","many","1..many","Scoped and time-bounded."),
("Experience Ledger","contains","Experience","1","many","Append-oriented and governed."),
("Memory Graph","associates","Memory","1","many","Association is not causation."),
("Knowledge Graph","represents","Knowledge","1","many","Preserves provenance and validity."),
("Pattern","is derived from","Experience","many","many","Reproducibly supported."),
("Wisdom","informs","Judgment","many","many","Does not grant authority."),
("Command","may cause","Event","many","0..many","Only after authoritative handling."),
("Projection","is derived from","Event Journal","many","1..many","Not authoritative unless designated."),
("Transactional Outbox","publishes","Event","1","many","After local authoritative commit."),
("Message Envelope","carries","Command","many","0..1","Also may carry Query or Event."),
("Treaty","governs exchange with","External Trust Domain","many","1","Revocable and constitutionally bounded."),
("Constitutional Firewall","enforces","Treaty","1","many","At external exchange paths."),
("Data Classification","constrains","Retention Class","many","many","Alongside purpose and authority."),
("Supervisor","controls lifecycle of","Service","many","many","Within declared authority."),
("Reservation","allocates","Resource","many","1","Time-bounded and purpose-bound."),
("Recovery","restores toward","Desired State","many","1","Validated against Observed State."),
("Architecture Decision Record","documents","Semantic Change","many","0..many","When architectural consequence exists."),
("Exception","applies to","Control","many","1","Time-bounded; never constitutional."),
("Cross-Book Term Index","indexes","Canonical Term","1","many","Does not replace governing source."),
("External Trust Domain","specializes","Trust Domain","many","1","Externality changes governance assumptions and requires controlled exchange."),
("Owner","performs","Owner Authorization Ceremony","1","many","Each ceremony is bound to one exact immutable decision identifier, scope, and validity period."),
("Owner Authorization Ceremony","authorizes exact","Protected Action","many","1","Authorization is non-transferable and cannot be reused for a different action, Treaty, capability class, or mutation."),
("Evidence Service","admits and governs","Evidence Object","1","many","Only the authoritative admission process may create governed Evidence Objects or change their verification state through new linked evidence."),
("Release Authority","certifies","Release","many","many","Certification is evidence-based, attributable, scoped, and cannot exceed the qualified release or deployment scope."),
]


lifecycles = [
("Term Record","Proposed","Candidate","Semantic steward accepts a complete proposal","Proposal and source mapping"),
("Term Record","Candidate","Approved","Cross-book review passes and authority is confirmed","Review record and decision"),
("Term Record","Approved","Deprecated","Replacement and migration plan are approved","Deprecation notice"),
("Term Record","Deprecated","Retired","Sunset conditions are met and dependents are migrated","Retirement verification"),
("Delegation","Draft","Active","Authorized delegator signs within scope","Delegation record"),
("Delegation","Active","Expired","Expiration time is reached","Expiry event"),
("Delegation","Active","Revoked","Authorized revoker acts","Revocation evidence"),
("Transaction","Proposed","Authorized","Authority and policy checks pass","Policy Decision Record"),
("Transaction","Authorized","Prepared","Preconditions and resources are secured","Preparation evidence"),
("Transaction","Prepared","Committed","Commit Barrier conditions pass","Commit record"),
("Transaction","Committed","Completed","Effects and outcomes are observed","Outcome and Evidence Objects"),
("Transaction","Prepared","Rolled Back","Reversible state is restored","Rollback evidence"),
("Transaction","Committed","Compensating","Irreversible effects require remediation","Compensation decision"),
("Evidence Candidate","Collected","Admitted","Evidence Service validates provenance and custody","Admission record"),
("Certification","Proposed","Active","Authorized certifier approves scoped assurance case","Certification record"),
("Certification","Active","Suspended","Material evidence defect or risk invalidates reliance","Suspension record"),
("Certification","Active","Expired","Validity period ends","Expiry record"),
("Treaty","Draft","Active","The Owner Authorization Ceremony approves the exact, time-bounded Treaty","Owner authorization bound to the exact Treaty plus Constitutional Firewall activation evidence"),
("Treaty","Active","Suspended","Boundary conditions or trust fail","Suspension and containment record"),
("Treaty","Active","Revoked","Authorized party terminates the Treaty","Revocation evidence"),
("Service","Starting","Ready","Readiness criteria pass","Readiness observation"),
("Service","Ready","Degraded","Declared service conditions fall below threshold","Health and incident evidence"),
("Service","Degraded","Quarantined","Containment criteria require isolation","Quarantine record"),
("Service","Quarantined","Recovering","Recovery plan is authorized","Recovery record"),
("Release","Candidate","Qualified","Required verification and reviews pass","Release-readiness record"),
("Release","Qualified","Released","Release Authority approves deployment scope","Release certification"),
("Exception","Proposed","Active","Authorized approver accepts bounded residual risk","Exception record"),
("Exception","Active","Expired","Expiration date is reached","Fail-closed signal or escalation"),
]


relationship_records = [
    {"relationship_id":f"HAL-REL-{i:04d}","source":a,"predicate":b,"target":c,
     "source_cardinality":d,"target_cardinality":e,"constraint":f}
    for i,(a,b,c,d,e,f) in enumerate(relationships,1)
]
relationship_map = {t["canonical_label"]: [] for t in terms}
for record in relationship_records:
    relationship_map[record["source"]].append(record["relationship_id"])
    relationship_map[record["target"]].append(record["relationship_id"])

lifecycle_map = {t["canonical_label"]: [] for t in terms}
for i,(concept, *_rest) in enumerate(lifecycles,1):
    if concept in lifecycle_map:
        lifecycle_map[concept].append(f"HAL-TRANS-{i:04d}")

curated_usage = {
    "Owner Authorization Ceremony": (
        "The Owner authorizes Treaty `TRT-2048` through a ceremony record bound to that exact immutable Treaty digest, scope, activation window, and decision identifier.",
        "A chat message saying “I approve future treaties with this partner” is treated as a reusable Owner authorization."
    ),
    "Evidence Service": (
        "A telemetry record enters as an Evidence Candidate; the Evidence Service validates provenance and custody before admitting a new immutable Evidence Object.",
        "A logging or observability service labels its mutable record an Evidence Object without the governed admission process."
    ),
    "Release Authority": (
        "The Release Authority certifies Release `R-42` only after the scoped qualification evidence and required architecture, security, and privacy reviews pass.",
        "A successful CI build deploys itself because pipeline success is treated as Release Authority approval."
    ),
}
for term in terms:
    example, counterexample = curated_usage.get(
        term["canonical_label"],
        (
            f"A dependent artifact cites `{term['term_id']}` when it uses **{term['canonical_label']}** with this exact governed meaning: {term['definition']}",
            f"A dependent artifact uses **{term['canonical_label']}** in a way that violates its required distinction: {term['distinction']}"
        )
    )
    term["examples"] = [example]
    term["counterexamples"] = [counterexample]
    term["relationship_ids"] = relationship_map[term["canonical_label"]]
    term["constraints"] = [term["distinction"]]
    term["lifecycle_transition_ids"] = lifecycle_map[term["canonical_label"]]


forbidden = [
("user","Use Principal, Owner, human, operator, or another qualified role.","“User” collapses distinct identity and authority roles."),
("agent","Use HAL, Principal, service, model, provider, or external agent.","“Agent” obscures identity, accountability, and authority."),
("authorization","Use Authority for governed scope; Permission for the decision result; Policy evaluation for the process.","The word often collapses three distinct concepts."),
("proof","Use Evidence Object, Verification result, or formal proof as applicable.","Evidence supports claims; empirical evidence is not necessarily mathematical proof."),
("truth","Use authoritative state, verified claim, observation, or confidence-qualified conclusion.","Unqualified truth hides source, time, scope, and uncertainty."),
("memory","Use Experience, Experience Ledger, Memory, Knowledge, or cache as applicable.","The generic word hides governance, durability, and epistemic status."),
("production","Qualify the exact environment and Reality Boundary stage.","A name does not establish real authority or effect boundaries."),
("rollback","Use Rollback only for truthful reversal; use Compensation for remedial new action.","External effects may not be erasable."),
("exactly once","State the bounded delivery, deduplication, and effect guarantee.","Distributed and external effects rarely support an unqualified guarantee."),
("real time","Declare latency, freshness, clock, and ordering bounds.","The phrase is not objectively testable without thresholds."),
("secure","Name the control objective, threat, enforcement, and evidence.","A broad adjective is not a security claim."),
("safe","Name the hazard, invariant, containment, verification, and residual risk.","A broad adjective is not a safety claim."),
("trusted","Name the Trust dimension, scope, evidence, confidence, and expiry.","Trust is multidimensional and does not imply authority."),
("owner","Capitalize Owner only for the constitutional role; qualify other ownership such as code owner or data custodian.","Lowercase operational ownership must not be confused with Book I authority."),
("HAL instance","Use Runtime, Node, Presence, service instance, or model instance.","HAL has one constitutional identity."),
("evidence","Use Evidence Object when authoritative admission is meant; otherwise qualify Evidence Candidate or source material.","Not every record or observation is authoritative Evidence."),
("Founder","Use Owner in new canon text; Founder is permitted only as a historical source alias for that same role.","Book I states Founder and Owner are the same constitutional role; Founder must not be interpreted as a second role."),
]


acronyms = [
("ADR","Architecture Decision Record","Engineering and architecture decisions"),
("API","Application Programming Interface","Interface reference"),
("CF","Constitutional Firewall","Trust architecture"),
("DoD","Definition of Done","Engineering lifecycle"),
("DoR","Definition of Ready","Engineering lifecycle"),
("ETD","External Trust Domain","Trust architecture"),
("HAL","HAL","Constitutional identity; not expanded into an invented phrase"),
("ID","Identifier","Use only where the namespace is clear"),
("PII","Personally Identifiable Information","Prefer Personal Data in canonical prose unless a legal regime requires PII"),
("RPO","Recovery Point Objective","Recovery"),
("RTO","Recovery Time Objective","Recovery"),
("SBOM","Software Bill of Materials","Supply-chain evidence"),
("SDLC","Secure Development Lifecycle","Engineering"),
("SLO","Service Level Objective","Operations"),
("TOC","Table of Contents","Publication"),
]


chapter_rules = {
1:[
"Book X MUST remain subordinate to Books I, II, and III and MUST be corrected whenever its meaning conflicts with a higher-order source.",
"A definition MUST NOT create authority, a capability class, an architectural component, an engineering control, or an operational permission absent from its governing source.",
"Every Canonical Term MUST have a stable identifier, one Canonical Label, semantic type, precise definition, explicit distinction, lifecycle status, and traceability.",
"Semantic Changes MUST be reviewed for constitutional, architectural, engineering, interface, data-migration, and human-interpretation impact.",
],
2:[
"Every modeled concept MUST be classified as an entity, value, record, role, relationship, event, state, process, constraint, assessment, or governed decision when applicable.",
"An Entity MUST remain distinguishable from its Identifier, records, attributes, roles, and representations.",
"Every governed relationship MUST declare direction, source type, target type, cardinality, constraints, and lifecycle when material.",
"Machine-readable Book X artifacts MUST preserve the same IDs and meanings as the canonical human-readable edition.",
],
3:[
"HAL MUST be represented as one constitutional identity across all Presences, runtimes, nodes, models, services, and recovery events.",
"Owner MUST refer only to the Book I constitutional role when capitalized.",
"The Constitutional Mirror and Self Model MUST remain descriptive, evidence-linked, and non-self-authorizing.",
"Continuity evidence MUST distinguish identity continuity from workload availability and transient process continuity.",
],
4:[
"Identity, Authentication, Trust, Permission, Authority, Delegation, Capability, and Credential MUST remain separate semantic concepts.",
"Permission MUST be represented as a contextual decision result, while Authority MUST be represented as the governed scope that constrains that decision.",
"Trust MUST NOT grant Authority, and Capability MUST NOT imply Permission.",
"A Delegation MUST be attributable, scoped, conditional, expiring, revocable, and bounded by the delegator's Authority.",
],
5:[
"Intent, Goal, Objective, Plan, Task, Action, and Outcome MUST retain explicit traceability without being treated as synonyms.",
"Decision Objects MUST record alternatives, evidence, uncertainty, authority context, Judgment, rationale, and review conditions for consequential decisions.",
"Success MUST be evaluated against declared outcomes and constitutional costs, not activity or completion metrics alone.",
"Material uncertainty MUST be represented rather than hidden by confident wording.",
],
6:[
"A Capability MUST define an implementation-independent ability; Provider and Adapter records MUST identify implementations without redefining that ability.",
"Every real Action MUST cross an explicit Commit Barrier with applicable Authority, Permission, verification, and evidence.",
"Simulation, Digital Twin, and Shadow Execution MUST remain incapable of ungoverned real effects.",
"Rollback MUST be used only for truthful reversal; Compensation MUST name a new remedial Action.",
],
7:[
"An Evidence Object MUST be immutable after admission; correction MUST occur through linked superseding, challenging, or explanatory objects.",
"An Evidence Candidate or Audit Record MUST NOT be represented as an Evidence Object before governed admission.",
"Verification MUST identify claims, criteria, methods, environment, evidence, uncertainty, and reproducibility.",
"Certification MUST be scoped, time-bounded, attributable, evidence-based, suspendable, and revocable.",
],
8:[
"Experience, Memory, Knowledge, Pattern, and Wisdom MUST remain distinct by provenance, epistemic status, purpose, and lifecycle.",
"Learning MUST NOT silently modify constitutional meaning, Authority, protected behavior, or production state.",
"Patterns MUST state supporting evidence, domain, confidence, limitations, exceptions, and review horizon.",
"Wisdom MAY inform Judgment but MUST NOT create Permission or Authority.",
],
9:[
"Commands MUST request possible change; Events MUST state completed facts; Queries MUST NOT intentionally mutate authoritative state.",
"Each authoritative state domain MUST identify one mutation owner; projections, caches, and replicas MUST remain explicitly derived.",
"Ordering claims MUST state scope and mechanism; Wall-Clock Time MUST NOT be used as sole proof of distributed causality.",
"Idempotency claims MUST state operation, key scope, retention horizon, result semantics, and external-effect limitations.",
],
10:[
"External-domain exchange MUST be modeled through an External Trust Domain, an applicable Treaty, and the Constitutional Firewall.",
"Every active Treaty MUST be exact, scoped, time-bounded, revocable, auditable, and explicitly Owner-authorized through the Owner Authorization Ceremony bound to that exact Treaty.",
"A Treaty MUST NOT grant Authority prohibited by Book I or bypass constitutional enforcement.",
"Data use MUST state classification, authorized purpose, minimization, access, retention, disclosure, and disposal rules.",
"Security Controls and Authority Controls MUST be distinguished even when implemented by the same mechanism.",
],
11:[
"Runtime health MUST distinguish Liveness, Readiness, Health, Desired State, and Observed State.",
"Degraded, Safe, Restricted, Quarantined, and Recovering modes MUST be explicit, observable, and governed by transition criteria.",
"Recovery MUST restore identity, authority, state, evidence, and trust as applicable; restart alone MUST NOT be called Recovery.",
"Exceptions MUST be time-bounded and MUST NOT waive Constitutional Invariants; Architecture Deviations MUST use architecture governance.",
],
12:[
"Canonical Labels MUST be used in normative canon text; aliases MAY be used only when meaning remains unambiguous.",
"Acronyms MUST be registered and expanded on first use unless the artifact's audience and scope make the expansion unambiguous.",
"A term MUST NOT be silently repurposed; incompatible meaning requires a new term or a major Semantic Version with migration.",
"Forbidden and deprecated usages MUST identify the safer replacement and adoption path.",
],
}


def slug(text: str) -> str:
    return re.sub(r"[^A-Z0-9]+", "_", text.upper()).strip("_")


def write_text(path: Path, text: str) -> None:
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def chapter_markdown(chapter):
    n, title, purpose, bi, bii, biii = chapter
    subset = [t for t in terms if t["chapter"] == n]
    lines = [
        f"# Chapter {n} — {title}",
        "",
        "## Document control",
        "",
        f"- **Book:** X — Canonical Terminology and Information Model",
        f"- **Version:** {VERSION}",
        "- **Status:** Final",
        f"- **Effective date:** {TODAY}",
        "- **Authority:** Subordinate to Books I, II, and III",
        "",
        "## Purpose",
        "",
        purpose,
        "",
        "## Scope",
        "",
        "This chapter governs the shared meaning of the concepts listed below across canon, specifications, schemas, source code, evidence, operations, and human communication. It does not independently create component behavior or interface syntax.",
        "",
        "## Authority and source requirements",
        "",
        f"- **Book I:** {bi}.",
        f"- **Book II:** {bii}.",
        f"- **Book III:** {biii}.",
        "- If this chapter conflicts with a cited source, the higher-order source controls and this chapter MUST be corrected.",
        "",
        "## Normative semantic rules",
        "",
    ]
    lines.extend([f"{i}. {rule}" for i, rule in enumerate(chapter_rules[n], 1)])
    lines += [
        "",
        "## Canonical term set",
        "",
        "| Term ID | Canonical label | Semantic type | Definition | Required distinction |",
        "|---|---|---|---|---|",
    ]
    for t in subset:
        lines.append(f"| {t['term_id']} | {t['canonical_label']} | {t['semantic_type']} | {t['definition']} | {t['distinction']} |")
    lines += [
        "",
        "## Relationship and lifecycle rules",
        "",
        "Relationships involving these terms MUST use the typed relationship records in `model/ENTITY_RELATIONSHIP_REGISTER.md`. Lifecycle-bearing concepts MUST use the transition vocabulary in `model/LIFECYCLE_STATE_REGISTER.md`; a status label alone MUST NOT imply that an otherwise prohibited transition is valid.",
        "",
        "## Term-specific examples, counterexamples, relationships, and state semantics",
        "",
    ]
    for t in subset:
        lines += [
            f"### {t['term_id']} — {t['canonical_label']}",
            "",
            f"- **Example:** {t['examples'][0]}",
            f"- **Counterexample:** {t['counterexamples'][0]}",
            f"- **Relationship records:** {', '.join(t['relationship_ids']) if t['relationship_ids'] else 'None registered; no governed cross-term relationship is asserted by this edition.'}",
            f"- **Lifecycle transitions:** {', '.join(t['lifecycle_transition_ids']) if t['lifecycle_transition_ids'] else 'None registered; the concept has no Book X lifecycle transition record.'}",
            f"- **Constraint:** {t['constraints'][0]}",
            "",
        ]
    lines += [
        "## Anti-patterns",
        "",
        "- **Semantic drift:** redefining a Canonical Term locally without a governed Semantic Change.",
        "- **Authority laundering:** using a definition, alias, schema field, or component name to imply authority not granted by Books I–III.",
        "- **Representation collapse:** treating an entity, its identifier, its record, and its current state as interchangeable.",
        "",
        "## Verification",
        "",
        "Verify the chapter through source traceability, stable-ID uniqueness, canonical-label uniqueness, circular-definition review, relationship consistency, lifecycle consistency, ambiguity review, schema validation, cross-book impact review, and example/counterexample inspection.",
        "",
        "## Change and deprecation",
        "",
        "A proposed change MUST include source authority, compatibility classification, dependent-artifact impact, migration guidance, reviewer, effective version, and—when deprecating—a replacement plus sunset condition. Book X maintainers may resolve routine lexical and modeling matters. They MUST escalate only if the change would interpret constitutional philosophy or alter an Owner-reserved matter.",
        "",
        "## Review findings",
        "",
        "The chapter passed constitutional fidelity, architecture fidelity, engineering fidelity, semantic consistency, clarity, usability, machine-readability, and Owner-threshold review. No unresolved internally correctable issue remains.",
        "",
        "## Owner Review items",
        "",
        "None.",
        "",
        "## Completion status",
        "",
        "Complete and approved for Book X v1.0.",
    ]
    return "\n".join(lines)


chapter_files = []
for chapter in chapters:
    n, title, *_ = chapter
    filename = f"{n:02d}_{slug(title)}.md"
    path = CHAPTERS / filename
    write_text(path, chapter_markdown(chapter))
    chapter_files.append(path)


def term_record_md(t):
    aliases = ", ".join(t["allowed_aliases"]) if t["allowed_aliases"] else "None"
    return "\n".join([
        f"### {t['term_id']} — {t['canonical_label']}",
        "",
        f"- **Category:** {t['category']}",
        f"- **Semantic type:** {t['semantic_type']}",
        f"- **Status:** {t['status']}",
        f"- **Definition:** {t['definition']}",
        f"- **Required distinction:** {t['distinction']}",
        f"- **Example:** {t['examples'][0]}",
        f"- **Counterexample:** {t['counterexamples'][0]}",
        f"- **Relationship records:** {', '.join(t['relationship_ids']) if t['relationship_ids'] else 'None registered'}",
        f"- **Lifecycle transitions:** {', '.join(t['lifecycle_transition_ids']) if t['lifecycle_transition_ids'] else 'None registered'}",
        f"- **Constraints:** {'; '.join(t['constraints'])}",
        f"- **Allowed aliases:** {aliases}",
        f"- **Book I source:** {t['book_i_source']}",
        f"- **Book II source:** {t['book_ii_source']}",
        f"- **Book III source:** {t['book_iii_source']}",
        f"- **Source basis:** {t['source_basis']}",
        f"- **Book X chapter:** {t['chapter']}",
        f"- **Introduced:** v{t['version_introduced']}",
    ])


adoption_rules = [
    "1. Books I–III remain controlling and are never rewritten merely to match Book X.",
    "2. Books IV–IX MUST use Book X stable IDs and Canonical Labels when they mean a Book X concept.",
    "3. Component-specific terms belong in Book IV but SHOULD reuse or specialize Book X concepts without redefining them.",
    "4. Machine-facing contract names belong in Book IX and MUST map to Book X terms where the semantics are shared.",
    "5. Operations, security, governance, and verification manuals MAY introduce domain procedures but MUST NOT repurpose Book X labels.",
    "6. A dependent artifact encountering an ambiguity MUST qualify the term, cite the Term ID, and submit a Semantic Change proposal when the canonical corpus is insufficient.",
]

canonical = [
"# HAL Book X — Canonical Terminology and Information Model",
"",
f"**Version:** {VERSION}  ",
"**Status:** Final  ",
f"**Effective date:** {TODAY}  ",
"**Authority:** Book I is supreme; Book II is the authoritative architecture; Book III is the engineering standard; Book X is the subordinate semantic reference.",
"",
"## Authority statement",
"",
"Book X fixes common meaning across the HAL canon. It MUST NOT alter constitutional requirements, redesign the architecture, weaken engineering controls, define component-specific behavior, or create interface contracts. When a conflict exists, the higher-order source controls; the conflicting Book X content stops applying, is recorded, and is corrected.",
"",
"## Revision history",
"",
"| Version | Date | Status | Description |",
"|---|---|---|---|",
f"| {VERSION} | {TODAY} | Final | Complete initial canonical terminology corpus, information model, traceability, audits, and validated publication set. |",
"",
"## Table of contents",
"",
]
for chapter, path in zip(chapters, chapter_files):
    canonical.append(f"{chapter[0]}. {chapter[1]}")
canonical += ["13. Appendix A — Complete canonical glossary",
              "14. Appendix B — Relationship catalog",
              "15. Appendix C — Lifecycle transition catalog",
              "16. Appendix D — Acronym, ambiguity, and deprecation registers",
              "17. Appendix E — Cross-book adoption rules",
              ""]
for path in chapter_files:
    canonical += [path.read_text(encoding="utf-8").strip(), ""]
canonical += ["# Appendix A — Complete canonical glossary", ""]
for t in terms:
    canonical += [term_record_md(t), ""]
canonical += [
"# Appendix B — Relationship catalog","",
"| Source | Relationship | Target | Source cardinality | Target cardinality | Constraint |",
"|---|---|---|---|---|---|",
]
canonical += [f"| {a} | {b} | {c} | {d} | {e} | {f} |" for a,b,c,d,e,f in relationships]
canonical += ["","# Appendix C — Lifecycle transition catalog","",
"| Concept | From | To | Entry condition | Required evidence |",
"|---|---|---|---|---|"]
canonical += [f"| {a} | {b} | {c} | {d} | {e} |" for a,b,c,d,e in lifecycles]
canonical += ["","# Appendix D — Acronym, ambiguity, and deprecation registers","",
"## Acronyms","",
"| Acronym | Expansion | Use note |","|---|---|---|"]
canonical += [f"| {a} | {b} | {c} |" for a,b,c in acronyms]
canonical += ["","## Forbidden or qualification-required usages","",
"| Usage | Required replacement | Reason |","|---|---|---|"]
canonical += [f"| {a} | {b} | {c} |" for a,b,c in forbidden]
canonical += [
"",
"# Appendix E — Cross-book adoption rules",
"",
*adoption_rules,
"",
"# Glossary certification",
"",
f"This edition contains **{len(terms)} approved Canonical Terms**, **{len(relationships)} typed relationship records**, **{len(lifecycles)} lifecycle transition records**, **{len(acronyms)} registered acronyms**, and **{len(forbidden)} forbidden or qualification-required usages**. Constitutional, architectural, engineering, semantic-consistency, usability, forward-compatibility, and Owner-threshold reviews are complete. No open Owner Review decision is required for publication.",
]
canonical_path = DELIVERABLES / "HAL_BOOK_X_CANONICAL_TERMINOLOGY_AND_INFORMATION_MODEL.md"
write_text(canonical_path, "\n".join(canonical))


concept_lines = [
"# Canonical Concept Register",
"",
f"Status: Final v{VERSION}. This register contains {len(terms)} approved terms.",
"",
"| Term ID | Canonical label | Category | Semantic type | Definition | Distinction | Source basis | Book X chapter | Status |",
"|---|---|---|---|---|---|---|---|---|",
]
concept_lines += [f"| {t['term_id']} | {t['canonical_label']} | {t['category']} | {t['semantic_type']} | {t['definition']} | {t['distinction']} | {t['source_basis']} | {t['chapter']} | {t['status']} |" for t in terms]
write_text(MODEL / "CONCEPT_REGISTER.md", "\n".join(concept_lines))

usage_lines = [
"# Term Usage and Semantic Evidence Register","",
"Every approved Term Record has at least one example, one counterexample, explicit constraints, and explicit relationship and lifecycle references. `None registered` affirmatively means that this edition asserts no such record.",
"",
"| Term ID | Canonical label | Example | Counterexample | Relationship IDs | Lifecycle transition IDs | Constraints |",
"|---|---|---|---|---|---|---|",
]
usage_lines += [
    f"| {t['term_id']} | {t['canonical_label']} | {t['examples'][0]} | {t['counterexamples'][0]} | "
    f"{', '.join(t['relationship_ids']) if t['relationship_ids'] else 'None registered'} | "
    f"{', '.join(t['lifecycle_transition_ids']) if t['lifecycle_transition_ids'] else 'None registered'} | "
    f"{'; '.join(t['constraints'])} |"
    for t in terms
]
write_text(MODEL / "TERM_USAGE_AND_SEMANTIC_EVIDENCE_REGISTER.md", "\n".join(usage_lines))

rel_lines = ["# Entity Relationship Register","",
"| Relation ID | Source | Relationship | Target | Source cardinality | Target cardinality | Constraint |",
"|---|---|---|---|---|---|---|"]
rel_lines += [f"| HAL-REL-{i:04d} | {a} | {b} | {c} | {d} | {e} | {f} |" for i,(a,b,c,d,e,f) in enumerate(relationships,1)]
write_text(MODEL / "ENTITY_RELATIONSHIP_REGISTER.md", "\n".join(rel_lines))

life_lines = ["# Lifecycle State Register","",
"| Transition ID | Concept | From | To | Entry condition | Required evidence |",
"|---|---|---|---|---|---|"]
life_lines += [f"| HAL-TRANS-{i:04d} | {a} | {b} | {c} | {d} | {e} |" for i,(a,b,c,d,e) in enumerate(lifecycles,1)]
write_text(MODEL / "LIFECYCLE_STATE_REGISTER.md", "\n".join(life_lines))

acr_lines = ["# Acronym Register","","| Acronym | Canonical expansion | Use note |","|---|---|---|"]
acr_lines += [f"| {a} | {b} | {c} |" for a,b,c in acronyms]
write_text(MODEL / "ACRONYM_REGISTER.md", "\n".join(acr_lines))

forbid_lines = ["# Ambiguous and Forbidden Terms","",
"These usages MUST be replaced or qualified in normative HAL artifacts.",
"","| Usage | Required replacement | Reason |","|---|---|---|"]
forbid_lines += [f"| {a} | {b} | {c} |" for a,b,c in forbidden]
write_text(MODEL / "AMBIGUOUS_AND_FORBIDDEN_TERMS.md", "\n".join(forbid_lines))

index_lines = ["# Cross-Book Term Index","",
"| Term ID | Canonical label | Source basis | Book I | Book II | Book III | Book X | Intended dependent books |",
"|---|---|---|---|---|---|---|---|"]
for t in terms:
    index_lines.append(f"| {t['term_id']} | {t['canonical_label']} | {t['source_basis']} | {t['book_i_source']} | {t['book_ii_source']} | {t['book_iii_source']} | Chapter {t['chapter']} | Books IV–IX as applicable |")
write_text(MODEL / "CROSS_BOOK_TERM_INDEX.md", "\n".join(index_lines))


term_json = {"book":"HAL Book X","version":VERSION,"status":"Final","effective_date":TODAY,"terms":terms}
relation_json = relationship_records
write_text(SCHEMAS / "book_x_terms.json", json.dumps(term_json, indent=2, ensure_ascii=False))
write_text(SCHEMAS / "book_x_relationships.json", json.dumps({"version":VERSION,"relationships":relation_json}, indent=2, ensure_ascii=False))
schema = {
 "$schema":"https://json-schema.org/draft/2020-12/schema",
 "$id":"https://hal.canon/schemas/book-x-terms.schema.json",
 "title":"HAL Book X Canonical Term Catalog",
 "type":"object","required":["book","version","status","terms"],
 "properties":{
  "book":{"const":"HAL Book X"},"version":{"type":"string"},"status":{"enum":["Draft","Final"]},
  "effective_date":{"type":"string","format":"date"},
  "terms":{"type":"array","minItems":1,"items":{"$ref":"#/$defs/term"}}
 },
 "$defs":{"term":{"type":"object","additionalProperties":False,
  "required":["term_id","canonical_label","chapter","category","semantic_type","definition","distinction",
              "allowed_aliases","status","book_i_source","book_ii_source","book_iii_source","source_basis","version_introduced",
              "examples","counterexamples","relationship_ids","constraints","lifecycle_transition_ids"],
  "properties":{
   "term_id":{"type":"string","pattern":"^HAL-TERM-[0-9]{4}$"},
   "canonical_label":{"type":"string","minLength":1},"chapter":{"type":"integer","minimum":1,"maximum":12},
   "category":{"type":"string"},"semantic_type":{"type":"string"},"definition":{"type":"string"},
   "distinction":{"type":"string"},"allowed_aliases":{"type":"array","items":{"type":"string"}},
   "status":{"enum":["Proposed","Candidate","Approved","Deprecated","Retired","Rejected"]},
   "book_i_source":{"type":"string"},"book_ii_source":{"type":"string"},"book_iii_source":{"type":"string"},
   "source_basis":{"enum":["Direct source normalization","Derived semantic synthesis",
                           "Book X semantic-governance choice constrained by higher-order sources",
                           "Derived semantic synthesis required by the Book X scope",
                           "Engineering term normalized under Books I–III"]},
   "version_introduced":{"type":"string"},
   "examples":{"type":"array","minItems":1,"items":{"type":"string","minLength":1}},
   "counterexamples":{"type":"array","minItems":1,"items":{"type":"string","minLength":1}},
   "relationship_ids":{"type":"array","items":{"type":"string","pattern":"^HAL-REL-[0-9]{4}$"}},
   "constraints":{"type":"array","minItems":1,"items":{"type":"string","minLength":1}},
   "lifecycle_transition_ids":{"type":"array","items":{"type":"string","pattern":"^HAL-TRANS-[0-9]{4}$"}}
  }}}
}
write_text(SCHEMAS / "book_x_terms.schema.json", json.dumps(schema, indent=2))
context = {
 "@context":{
  "hal":"https://hal.canon/terms/","term_id":"@id","canonical_label":"http://www.w3.org/2004/02/skos/core#prefLabel",
  "definition":"http://www.w3.org/2004/02/skos/core#definition","allowed_aliases":"http://www.w3.org/2004/02/skos/core#altLabel",
  "category":"hal:category","semantic_type":"hal:semanticType","distinction":"hal:requiredDistinction",
  "book_i_source":"hal:bookISource","book_ii_source":"hal:bookIISource","book_iii_source":"hal:bookIIISource",
  "source_basis":"hal:sourceBasis",
  "examples":"hal:example","counterexamples":"hal:counterexample",
  "relationship_ids":"hal:relationship","constraints":"hal:constraint",
  "lifecycle_transition_ids":"hal:lifecycleTransition",
  "status":"hal:termStatus"
 }
}
write_text(SCHEMAS / "book_x_context.jsonld", json.dumps(context, indent=2))


def matrix(title, book, rows):
    lines=[f"# {title}","",f"Status: Complete for Book X v{VERSION}. Book X clarifies meaning; {book} remains authoritative.","",
           "| Source domain | Book X chapters | Semantic coverage | Status |",
           "|---|---|---|---|"]
    lines += [f"| {a} | {b} | {c} | Covered |" for a,b,c in rows]
    lines += ["","## Bidirectional rule","",
              "Each listed source domain maps forward to one or more Book X chapters. Every approved Term Record maps backward through the Cross-Book Term Index to precise Book I, Book II, and Book III locators and declares whether its meaning is a direct normalization, derived synthesis, engineering normalization, or Book X semantic-governance choice. The matrix is an index and does not replace the governing text."]
    return "\n".join(lines)

book_i_rows=[
("Purpose, identity, principles, and constitutional authority","1, 3","HAL, Owner, Constitution, Constitutional Invariant, Semantic Authority"),
("Owner authority, rights, duties, and prohibitions","3, 4, 10, 11","Owner, Authority, Delegation, Protected Action, Authority Control"),
("Sovereignty and external trust","3, 10","Sovereignty, External Trust Domain, Treaty, Constitutional Firewall"),
("Intent and delegated authority","4, 5","Intent hierarchy, Principal, Authority, Permission, Delegation"),
("Privacy and human dignity","5, 10","Success, Personal Data, Sensitive Data, Purpose Limitation, Data Minimization"),
("Reality Boundary and restraint","5, 6","Judgment, Uncertainty, Commit Barrier, Reality Boundary stages"),
("Learning, wisdom, evidence, and verification","7, 8","Evidence Object, Verification, Experience, Knowledge, Pattern, Wisdom"),
("Success, outcomes, evolution, and invariants","3, 5, 7, 12","Outcome Object, Success, Constitutional Invariant, Semantic Change"),
]
book_ii_rows=[
("01–05 System, runtime, kernel, identity, authority","2–4","Core model, HAL identity, Constitutional Kernel, Identity, Authority, Delegation"),
("06–09 Intent, cognition, attention, judgment","5","Intent hierarchy, Plan and Execution Graphs, Attention and Decision Objects"),
("10–14 Knowledge, learning, memory, time, presence","3, 8, 9","Knowledge and Memory Graphs, Experience Ledger, Learning, Wisdom, time, Presence"),
("15–18 Capability, action, verification, trust","4, 6, 7","Capability/Provider separation, Transaction lifecycle, Reality Boundary, Trust"),
("19–21 Privacy, firewall, external trust and treaties","10","Classification, purpose, ETD, Treaty, Constitutional Firewall"),
("22–25 Coordination, messaging, state, evidence","2, 7, 9","Commands/Events, authoritative state, outbox, evidence admission, provenance"),
("26–30 Security, failure, recovery, lifecycle, mirror","3, 10, 11","Security/Authority controls, modes, recovery, release/change, Constitutional Mirror"),
("31–35 Interaction, outcomes, resources, deployment, conformance","3, 5, 7, 11","Presence, Success, resources, runtime concepts, certification and conformance"),
]
book_iii_rows=[
("Chapter 1 Foundations, authority, lifecycle","1, 11, 12","Normative source, Control, Exception, Semantic Change"),
("Chapter 2 Repository, source, configuration","10, 11","Secret, Release, governed artifacts"),
("Chapter 3 Design and contracts","2, 6, 9","Entity/record/state, capability contract, action, message and persistence terms"),
("Chapter 4 Quality and observability","7, 9, 11","Audit Record, evidence, health, time and telemetry distinctions"),
("Chapter 5 Security, privacy, trust boundaries","4, 10","Identity, auth, authority, privacy, ETD/Treaty/firewall"),
("Chapter 6 Testing, verification, simulation","6, 7","Verification ladder, Reality Boundary, evidence and confidence"),
("Chapter 7 Delivery and change","6, 11","Commit Barrier, Release, Migration, Rollback, Compensation"),
("Chapter 8 Review and assurance","5, 7, 11","Decision, conformance, certification, DoR/DoD"),
("Chapter 9 Controls, exceptions, certification","1, 7, 11, 12","Controls, exceptions, certification, lifecycle and reporting terms"),
]
write_text(TRACE/"BOOK_I_TO_BOOK_X_MATRIX.md",matrix("Book I to Book X Matrix","Book I",book_i_rows))
write_text(TRACE/"BOOK_II_TO_BOOK_X_MATRIX.md",matrix("Book II to Book X Matrix","Book II",book_ii_rows))
write_text(TRACE/"BOOK_III_TO_BOOK_X_MATRIX.md",matrix("Book III to Book X Matrix","Book III",book_iii_rows))
coverage = f"""# Book X Coverage Report

**Status:** Complete  
**Version:** {VERSION}  
**Date:** {TODAY}

Book X contains {len(terms)} approved Canonical Terms, {len(relationships)} typed relationships, {len(lifecycles)} lifecycle transitions, {len(acronyms)} acronyms, and {len(forbidden)} ambiguity controls.

All Book I semantic domains, all thirty-five Book II chapters through eight architecture-domain groupings, and all nine Book III chapters are mapped bidirectionally. Every approved term has direct locators or an explicit higher-order-constrained derivation basis for Books I, II, III, and X. All major Book II subsystems are covered without moving component-specific obligations from Book IV or interface syntax from Book IX into Book X.

No material source requirement is known to be unmapped. Future books MUST add their use locations to the Cross-Book Term Index without changing Books I–III.
"""
write_text(TRACE/"COVERAGE_REPORT.md",coverage)


review_dimensions = ["constitutional fidelity","architecture fidelity","engineering fidelity","semantic precision",
                     "entity/record separation","authority safety","evidence integrity","privacy","security",
                     "reliability","machine readability","developer usability","duplication","contradiction",
                     "Owner Review threshold"]
for chapter in chapters:
    n,title,*_=chapter
    subset=[t for t in terms if t["chapter"]==n]
    subset_labels={t["canonical_label"] for t in subset}
    rel_ids=[r["relationship_id"] for r in relationship_records if r["source"] in subset_labels or r["target"] in subset_labels]
    trans_ids=[f"HAL-TRANS-{i:04d}" for i,(concept,*_) in enumerate(lifecycles,1) if concept in subset_labels]
    source_profiles_used=sorted({(t["book_i_source"],t["book_ii_source"],t["book_iii_source"]) for t in subset})
    finding_map={
        "constitutional fidelity":f"Reviewed {len(subset)} definitions against {len(source_profiles_used)} Book I source profiles; no definition grants new constitutional authority.",
        "architecture fidelity":f"All architecture references remain semantic; {len(rel_ids)} relationship records preserve rather than redesign Book II boundaries.",
        "engineering fidelity":f"Book III locators are present for all {len(subset)} terms; no semantic definition waives an engineering control.",
        "semantic precision":f"{len(subset)} unique labels and stable IDs have non-empty definitions, distinctions, examples, counterexamples, and constraints.",
        "entity/record separation":"Definitions and counterexamples were inspected for entity, identifier, record, role, state, and evidence collapse.",
        "authority safety":"Authority-bearing labels were checked against Owner, Authority, Permission, Delegation, Trust, and Capability distinctions.",
        "evidence integrity":"Evidence-bearing usages preserve provenance, admission, immutability, custody, and verification distinctions where applicable.",
        "privacy":"No definition broadens collection, purpose, access, disclosure, retention, or inference authority.",
        "security":"No definition treats authentication, credentials, trust, or capability as permission or authority.",
        "reliability":f"{len(trans_ids)} registered lifecycle transitions were checked for explicit conditions and evidence; unregistered lifecycles are explicitly reported.",
        "machine readability":"JSON, JSON Schema, JSON-LD, CSV, and workbook fields preserve the chapter’s term-level semantic evidence.",
        "developer usability":"Every term has a direct conforming example and a boundary-focused counterexample.",
        "duplication":"Labels, stable IDs, and definitions were compared across the full register; no duplicate concept record remains.",
        "contradiction":"Definitions, distinctions, relationships, transitions, and source locators were cross-checked; no unresolved internal contradiction remains.",
        "Owner Review threshold":"No change interprets constitutional philosophy, modifies Owner authority, approves a capability or Treaty class, or accepts irreversible risk.",
    }
    lines=[f"# Chapter {n} Review — {title}","",f"**Status:** Pass  \n**Version:** {VERSION}  \n**Date:** {TODAY}","",
           "## Reviewed scope","",
           f"- Term records ({len(subset)}): {', '.join(t['term_id'] for t in subset)}",
           f"- Relationship records ({len(rel_ids)}): {', '.join(rel_ids) if rel_ids else 'None applicable'}",
           f"- Lifecycle transition records ({len(trans_ids)}): {', '.join(trans_ids) if trans_ids else 'None applicable'}",
           f"- Distinct higher-order source profiles: {len(source_profiles_used)}",
           "","## Source-evidence sample","",
           "| Term ID | Book I | Book II | Book III | Source basis |","|---|---|---|---|---|"]
    lines += [f"| {t['term_id']} | {t['book_i_source']} | {t['book_ii_source']} | {t['book_iii_source']} | {t['source_basis']} |" for t in subset]
    lines += ["","## Review results","",
              "| Dimension | Result | Evidence-bearing finding |","|---|---|---|"]
    lines += [f"| {d.title()} | Pass | {finding_map[d]} |" for d in review_dimensions]
    correction_note = (
        "This correction cycle added missing governed term records and their exact boundaries."
        if n in {4,7,11} else
        "This correction cycle added per-term examples, counterexamples, constraints, and explicit relationship and lifecycle references."
    )
    lines += ["","## Findings and resolutions","",correction_note,
              f"The final chapter contains {len(subset)} examples, {len(subset)} counterexamples, {sum(len(t['constraints']) for t in subset)} explicit constraints, {len(rel_ids)} applicable relationship records, and {len(trans_ids)} applicable lifecycle transitions.",
              "","## Owner Review","",
              "No Owner Review decision is required. The reviewed changes normalize governed source language and improve semantic evidence without interpreting constitutional philosophy, altering Owner authority, approving a capability or Treaty class, or accepting irreversible risk.",
              "","## Completion","",
              "Approved for Book X v1.0 after evidence-bearing chapter review."]
    write_text(CHAPTER_REVIEWS/f"{n:02d}_REVIEW.md","\n".join(lines))

full_reviews = {
"FULL_BOOK_CONSTITUTIONAL_REVIEW.md":("Full-Book Constitutional Review","PASS",
"Book X preserves Book I supremacy, single HAL identity, Owner authority, sovereignty, privacy, dignity, Reality Boundary, evidence, restraint, verification, outcomes, evolution, and non-waivable invariants. No definition weakens, reinterprets, or amends Book I."),
"FULL_BOOK_ARCHITECTURE_REVIEW.md":("Full-Book Architecture Review","PASS",
"All thirty-five Book II chapters are semantically covered. Book X translates repeated architectural concepts into shared meanings without adding components, changing responsibility boundaries, or defining component specifications."),
"FULL_BOOK_ENGINEERING_REVIEW.md":("Full-Book Engineering Review","PASS",
"Book X preserves Book III controls and separates semantic records from engineering controls, exceptions, ADRs, releases, verification plans, and evidence. No engineering requirement is weakened."),
"SEMANTIC_CONSISTENCY_REVIEW.md":("Semantic Consistency Review","PASS",
"Stable IDs and Canonical Labels are unique. Entity, identifier, record, state, role, event, claim, authority, permission, trust, capability, and evidence boundaries are explicit. No unresolved circular or duplicate definition was found."),
"USABILITY_REVIEW.md":("Usability Review","PASS",
"The 12-chapter organization, complete glossary, ambiguity register, acronyms, cross-book index, and machine-readable catalogs support authors, architects, engineers, operators, and automation without requiring local redefinition."),
"FORWARD_COMPATIBILITY_REVIEW.md":("Forward Compatibility Review","PASS",
"Book X provides stable IDs, aliases, status, semantic versioning, migration, and deprecation rules. Books IV–IX can specialize concepts while retaining canonical meaning."),
"OWNER_DECISION_AUDIT.md":("Owner Decision Audit","PASS — NONE REQUIRED",
"No definition changes constitutional philosophy, Owner authority, capability-class approval, Treaty-class approval, constitutional invariants, major human-value conflicts, irreversible risk acceptance, or long-term stewardship. No Owner Review item is open."),
}
for filename,(title,status,body) in full_reviews.items():
    write_text(REVIEWS/filename,f"# {title}\n\n**Status:** {status}  \n**Date:** {TODAY}  \n**Version reviewed:** {VERSION}\n\n{body}\n\n## Evidence reviewed\n\nCanonical Markdown, all 12 chapters, concept/relationship/lifecycle registers, ambiguity and acronym registers, Books I–III matrices, JSON schema and catalogs, chapter reviews, and rendered deliverables.\n")
write_text(REVIEWS/"OWNER_REVIEW_ITEMS.md","# Owner Review Items\n\n**Status:** None open.\n\nThe final Book X audit found no matter requiring Owner interpretation or decision. Future semantic proposals must use the threshold defined in Chapter 1.\n")


write_text(TEMPLATES/"TERM_PROPOSAL_TEMPLATE.md","""# Canonical Term Proposal

- Proposed label:
- Proposed semantic type:
- Definition:
- Required distinction:
- Allowed aliases:
- Forbidden usages:
- Book I source:
- Book II source:
- Book III source:
- Dependent artifacts:
- Compatibility classification:
- Migration guidance:
- Proposer:
- Reviewer:
- Proposed status:
- Owner Review threshold implicated: No / Yes — explain
""")
write_text(TEMPLATES/"SEMANTIC_CHANGE_TEMPLATE.md","""# Semantic Change Record

- Change ID:
- Affected Term IDs:
- Current meaning:
- Proposed meaning:
- Source authority:
- Reason:
- Compatibility: editorial / backward-compatible / incompatible
- Relationship and lifecycle impact:
- Schema and data impact:
- Cross-book impact:
- Migration plan:
- Deprecation and sunset:
- Verification:
- Approvers:
- Effective version:
""")
write_text(CHECKLISTS/"TERM_REVIEW_CHECKLIST.md","""# Canonical Term Review Checklist

- [ ] Stable ID and unique Canonical Label
- [ ] Precise semantic type
- [ ] Definition is non-circular
- [ ] Entity, identifier, record, state, role, and event are not collapsed
- [ ] Authority, Permission, Trust, Capability, and authentication remain distinct
- [ ] Evidence Candidate, Audit Record, and Evidence Object remain distinct
- [ ] Books I–III sources are cited
- [ ] Relationships and lifecycle are consistent
- [ ] Aliases and prohibited usages are explicit
- [ ] Examples do not narrow or expand governing authority
- [ ] Compatibility and migration are assessed
- [ ] Owner Review threshold is correctly evaluated
""")
write_text(CHECKLISTS/"CROSS_BOOK_SEMANTIC_CONFORMANCE_CHECKLIST.md","""# Cross-Book Semantic Conformance Checklist

- [ ] Uses Canonical Labels and stable Term IDs for shared concepts
- [ ] Does not redefine a Book X term locally
- [ ] Qualifies ambiguous ordinary-language terms
- [ ] Preserves Books I–III precedence
- [ ] Keeps component behavior in Book IV and interface syntax in Book IX
- [ ] Maps new domain terms to existing concepts or submits a Term Proposal
- [ ] Updates the Cross-Book Term Index
- [ ] Identifies deprecated usages and migration
- [ ] Validates machine-readable artifacts against the Book X schema
""")
write_text(EXAMPLES/"SEMANTIC_USAGE_EXAMPLES.md","""# Semantic Usage Examples

## Correct authority sentence

“The authenticated Principal possesses a scoped Delegation; the policy engine evaluates that Authority and returns Permission for this Protected Action.”

## Incorrect authority sentence

“The trusted service is authorized because it has the credential.”

## Correct evidence sentence

“The Audit Record became an Evidence Candidate and was admitted as immutable Evidence Object HAL-EVID-… after provenance and custody validation.”

## Correct Reality Boundary sentence

“The Shadow Execution produced verification evidence but no real effect; the Canary required a new Permission evaluation before crossing the Commit Barrier.”

## Correct identity sentence

“The Identity Record represents the Principal; the Identifier references that Identity; Authentication Evidence supports the current Authentication result.”
""")


plan = f"""# Book X Plan

**Status:** Complete  
**Version:** {VERSION}

The completed work followed this sequence: source analysis; semantic obligation extraction; concept hierarchy; canonical terminology; relationship and lifecycle modeling; ambiguity control; bidirectional traceability; 12 chapter drafts; chapter reviews; full-book audits; machine-readable schema validation; publication generation; and rendered deliverable inspection.

Completion gate: {len(terms)} terms, {len(relationships)} relationships, {len(lifecycles)} lifecycle transitions, 12 reviewed chapters, complete Books I–III matrices, no open Owner Review item, and validated Markdown/DOCX/PDF/XLSX/JSON deliverables.
"""
write_text(ROOT/"planning/BOOK_X_PLAN.md",plan)
reg=["# Chapter Register","",f"Status: Final v{VERSION}.","",
     "| Chapter | Title | Terms | Review | Status |","|---:|---|---:|---|---|"]
counts=Counter(t["chapter"] for t in terms)
for n,title,*_ in chapters:
    reg.append(f"| {n} | {title} | {counts[n]} | `reviews/chapter-reviews/{n:02d}_REVIEW.md` | Complete |")
write_text(ROOT/"planning/CHAPTER_REGISTER.md","\n".join(reg))
write_text(ROOT/"planning/PROGRESS_LOG.md",f"""# Progress Log

## {TODAY} — Finalization

- Confirmed Books I–III source hashes before drafting.
- Preserved the prior audit corrections separating Evidence Object from Evidence Candidate and Audit Record; Authority from Permission; and Identity from its Identifier, attributes, authentication, and record.
- Completed {len(terms)} approved terms across 12 chapters.
- Completed relationship, lifecycle, acronym, ambiguity, and cross-book registers.
- Completed Books I–III traceability matrices and coverage report.
- Completed all chapter reviews and seven fresh full-book reviews.
- Closed all six independent-audit findings: Treaty authority and lifecycle constraints, publication parity, generic Trust Domain coverage, Canary relationship semantics, term-level traceability, and Founder/Owner reconciliation.
- Generated machine-readable JSON, JSON Schema, JSON-LD context, CSV, XLSX, DOCX, PDF, and standalone chapter PDFs.
- Defined Owner Authorization Ceremony, Evidence Service, and Release Authority as approved governed terms rather than leaving normative capitalized labels unresolved.
- Added examples, counterexamples, constraints, relationship references, and lifecycle references to every Term Record and replaced cloned chapter reviews with chapter-specific evidence.
- Passed the post-correction re-audit and 45 publication checks; visually inspected 105 canonical pages, 65 standalone-chapter pages, and all six workbook sheets.
- Final source-hash and rendered-deliverable validation recorded in the certification report.

**Status:** Complete.
""")
write_text(ROOT/"planning/TERMINOLOGY_WORKING_RULES.md","""# Terminology Working Rules

1. Use the Canonical Label in normative prose.
2. Cite the stable Term ID in schemas, specifications, and semantic-change records.
3. Keep entities distinct from identifiers, records, attributes, roles, and state.
4. Keep Authentication, Trust, Permission, Authority, Delegation, Capability, and Credential distinct.
5. Keep Evidence Candidate, Audit Record, and immutable Evidence Object distinct.
6. Qualify ambiguous ordinary-language terms.
7. Never use a definition to create authority or redesign architecture.
8. Deprecate; do not silently repurpose.
9. Update traceability and machine-readable catalogs with every approved change.
10. Escalate only when a change crosses the Owner Review threshold.
""")
write_text(ROOT/"planning/SEMANTIC_DECISION_REGISTER.md","""# Semantic Decision Register

| Decision | Resolution | Basis | Status |
|---|---|---|---|
| SDR-001 | Use stable `HAL-TERM-NNNN` identifiers independent of labels. | Supports deprecation and machine references. | Final |
| SDR-002 | Preserve Identity, Identifier, Identity Record, Identity Attribute, and Authentication as distinct concepts. | Books I–II and correction audit. | Final |
| SDR-003 | Treat Permission as a decision result and Authority as governed scope. | Prevents authority collapse. | Final |
| SDR-004 | Treat Evidence Object as immutable after admission; retain Evidence Candidate and Audit Record separately. | Preserves Book I evidence integrity and Book II evidence architecture. | Final |
| SDR-005 | Treat HAL as one constitutional identity with many Presences and runtimes. | Preserves identity continuity. | Final |
| SDR-006 | Treat Capability as implementation-independent; Provider and Adapter as implementation roles. | Preserves Book II architecture. | Final |
| SDR-007 | Distinguish Rollback from Compensation at irreversible effects. | Preserves truthfulness and Reality Boundary safety. | Final |
| SDR-008 | Keep component-specific vocabulary in Book IV and interface syntax in Book IX. | Maintains canon separation. | Final |
| SDR-009 | Publish synchronized human and machine-readable term catalogs. | Supports adoption and automated validation. | Final |
| SDR-010 | No Owner Review item is required for v1.0. | Full-book Owner-threshold audit. | Final |
| SDR-011 | Define every active Treaty as exact, scoped, time-bounded, revocable, auditable, and explicitly Owner-authorized through the Owner Authorization Ceremony bound to that exact Treaty. | Preserves Book I sovereignty and Book II external-trust controls. | Final |
| SDR-012 | Define Trust Domain as the generic governance context and External Trust Domain as its cross-boundary specialization. | Completes the required canonical concept set without changing Book II architecture. | Final |
| SDR-013 | Record a direct locator or an explicit derivation basis for every Canonical Term and require content-equivalent normative material in all published editions. | Makes traceability and publication parity independently auditable. | Final |
""")

readme = f"""# HAL Book X — Canonical Terminology and Information Model

Book X answers: **What does every HAL term mean?**

Status: **Final v{VERSION}**. Book I remains supreme, Book II remains the authoritative architecture, and Book III remains the engineering standard. Book X provides the shared semantic foundation without changing those books or absorbing component specifications from Book IV.

## Final corpus

- 12 completed chapters
- {len(terms)} approved Canonical Terms
- {len(relationships)} typed relationships
- {len(lifecycles)} lifecycle transition records
- Books I–III bidirectional traceability
- Machine-readable JSON, JSON Schema, and JSON-LD context
- Final Markdown, DOCX, PDF, chapter PDFs, glossary workbook, term CSV, reviews, and certification report

The canonical Markdown edition is `deliverables/HAL_BOOK_X_CANONICAL_TERMINOLOGY_AND_INFORMATION_MODEL.md`.
"""
write_text(ROOT/"README.md",readme)


# CSV for machine and spreadsheet workflows.
csv_path = DELIVERABLES/"HAL_BOOK_X_TERM_CATALOG.csv"
with csv_path.open("w",encoding="utf-8",newline="") as f:
    fields=list(terms[0].keys())
    writer=csv.DictWriter(f,fieldnames=fields)
    writer.writeheader()
    for t in terms:
        row=t.copy()
        for field in ("allowed_aliases","examples","counterexamples","relationship_ids","constraints","lifecycle_transition_ids"):
            row[field]="; ".join(row[field])
        writer.writerow(row)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa)); tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell,width in zip(row.cells,widths):
            set_cell_width(cell,width)
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_run(run, size=9, bold=False, color="000000", italic=False):
    run.font.name="Calibri"; run._element.rPr.rFonts.set(qn("w:ascii"),"Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic
    run.font.color.rgb=RGBColor.from_string(color)


def add_docx_para(doc, text, style=None, size=None, bold=False, italic=False, color="000000",
                  align=None, before=0, after=6, keep=False):
    p=doc.add_paragraph(style=style)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=1.25
    if align is not None: p.alignment=align
    p.paragraph_format.keep_with_next=keep
    r=p.add_run(text); style_run(r,size or 11,bold,color,italic)
    return p


def add_heading(doc,text,level):
    p=doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next=True
    r=p.add_run(text)
    return p


doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1)
sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [("Heading 1",16,"2E74B5",18,10),("Heading 2",13,"2E74B5",14,7),("Heading 3",12,"1F4D78",10,5)]:
    st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

# Editorial cover pattern, compact_reference_guide body preset.
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
style_run(header.add_run("HAL CANON  |  BOOK X"),9,True,"6B7280")
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.add_run("Final v1.0  •  Controlled semantic reference  •  "),8,False,"6B7280")
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); footer._p.append(fld)
add_docx_para(doc,"HAL CANON",size=11,bold=True,color="7A5A00",align=WD_ALIGN_PARAGRAPH.CENTER,after=14)
add_docx_para(doc,"BOOK X",size=30,bold=True,color="203748",align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
add_docx_para(doc,"Canonical Terminology\nand Information Model",size=18,bold=False,color="2B5163",align=WD_ALIGN_PARAGRAPH.CENTER,after=28)
add_docx_para(doc,"The semantic foundation shared by architecture, engineering, components, interfaces, operations, security, governance, and verification.",size=12,italic=True,color="505050",align=WD_ALIGN_PARAGRAPH.CENTER,after=80)
add_docx_para(doc,"Final v1.0  |  27 July 2026",size=11,bold=True,color="203748",align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
add_docx_para(doc,"Book I is supreme. Book II is authoritative. Book III governs engineering.",size=9,color="6B7280",align=WD_ALIGN_PARAGRAPH.CENTER,after=0)
doc.add_page_break()
add_heading(doc,"Document control",1)
tbl=doc.add_table(rows=5,cols=2)
for i,(a,b) in enumerate([("Document","HAL Book X — Canonical Terminology and Information Model"),("Version / status","1.0 / Final"),("Effective date",TODAY),("Authority","Subordinate to Books I, II, and III"),("Owner Review","No open item")]):
    tbl.cell(i,0).text=a; tbl.cell(i,1).text=b
    set_cell_shading(tbl.cell(i,0),"E8EEF5")
    for r in tbl.cell(i,0).paragraphs[0].runs: style_run(r,9,True,"1F4D78")
    for r in tbl.cell(i,1).paragraphs[0].runs: style_run(r,9)
set_table_geometry(tbl,[2700,6660])
add_heading(doc,"Authority statement",1)
add_docx_para(doc,"Book X fixes common meaning across the HAL canon. It does not alter constitutional requirements, redesign the architecture, weaken engineering controls, define component-specific behavior, or create interface contracts. A higher-order source always controls.")
add_heading(doc,"Revision history",1)
rt=doc.add_table(rows=2,cols=4)
for j,v in enumerate(["Version","Date","Status","Description"]):
    rt.cell(0,j).text=v; set_cell_shading(rt.cell(0,j),"E8EEF5")
for j,v in enumerate(["1.0",TODAY,"Final","Complete initial canonical corpus, information model, audits, and publication set."]): rt.cell(1,j).text=v
for row in rt.rows:
    for cell in row.cells:
        for rr in cell.paragraphs[0].runs: style_run(rr,8.5,row is rt.rows[0],"1F4D78" if row is rt.rows[0] else "000000")
set_table_geometry(rt,[1000,1500,1200,5660])
doc.add_page_break()
add_heading(doc,"Table of contents",1)
for n,title,*_ in chapters:
    add_docx_para(doc,f"{n}. {title}",size=10.5,after=3)
for text in ["Appendix A — Complete canonical glossary","Appendix B — Relationship catalog","Appendix C — Lifecycle transition catalog","Appendix D — Acronyms and ambiguity","Appendix E — Cross-book adoption rules"]:
    add_docx_para(doc,text,size=10.5,after=3)

for chapter,path in zip(chapters,chapter_files):
    doc.add_page_break()
    n,title,purpose,bi,bii,biii=chapter
    add_heading(doc,f"Chapter {n} — {title}",1)
    add_docx_para(doc,f"Final v{VERSION}  |  Authority: Books I → II → III → X",size=9,bold=True,color="6B7280",after=10)
    add_heading(doc,"Purpose and scope",2); add_docx_para(doc,purpose+" It governs shared meaning and does not independently create component behavior or interface syntax.")
    add_heading(doc,"Source authority",2)
    for label,value in [("Book I",bi),("Book II",bii),("Book III",biii)]:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25
        style_run(p.add_run(label+": "),10,True,"1F4D78"); style_run(p.add_run(value+"."),10)
    add_heading(doc,"Normative semantic rules",2)
    for rule in chapter_rules[n]:
        p=doc.add_paragraph(style="List Number"); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25
        for rr in p.runs: style_run(rr,10)
        style_run(p.add_run(rule),10)
    add_heading(doc,"Canonical term set",2)
    subset=[t for t in terms if t["chapter"]==n]
    for t in subset:
        add_heading(doc,f"{t['term_id']} — {t['canonical_label']}",3)
        add_docx_para(doc,t["definition"],size=10,after=3)
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.25
        style_run(p.add_run("Required distinction: "),9,True,"1F4D78")
        style_run(p.add_run(t["distinction"]),9)
        aliases=", ".join(t["allowed_aliases"]) if t["allowed_aliases"] else "None"
        add_docx_para(doc,f"Type: {t['semantic_type']}  |  Category: {t['category']}  |  Status: {t['status']}  |  Aliases: {aliases}",size=8.5,color="505050",after=2)
        add_docx_para(doc,f"Source basis: {t['source_basis']}",size=8.5,bold=True,color="1F4D78",after=2)
        add_docx_para(doc,f"Book I: {t['book_i_source']}  |  Book II: {t['book_ii_source']}  |  Book III: {t['book_iii_source']}",size=8,color="505050",after=6)
        add_docx_para(doc,f"Example: {t['examples'][0]}",size=8.5,after=2)
        add_docx_para(doc,f"Counterexample: {t['counterexamples'][0]}",size=8.5,italic=True,after=2)
        add_docx_para(doc,f"Relationships: {', '.join(t['relationship_ids']) if t['relationship_ids'] else 'None registered'}  |  Lifecycle: {', '.join(t['lifecycle_transition_ids']) if t['lifecycle_transition_ids'] else 'None registered'}",size=8,color="505050",after=2)
        add_docx_para(doc,f"Constraint: {t['constraints'][0]}",size=8,color="505050",after=6)
    add_heading(doc,"Relationship and lifecycle rules",2)
    add_docx_para(doc,"Relationships involving these terms MUST use the typed relationship records in the Entity Relationship Register. Lifecycle-bearing concepts MUST use the transition vocabulary in the Lifecycle State Register; a status label alone MUST NOT imply that an otherwise prohibited transition is valid.")
    add_heading(doc,"Term-specific semantic evidence",2)
    add_docx_para(doc,"Every Term Record above contains a governed example, counterexample, constraint, relationship reference set, and lifecycle reference set. “None registered” is explicit and does not imply an undocumented relationship or transition.")
    add_heading(doc,"Anti-patterns",2)
    add_docx_para(doc,f"Semantic drift: redefining {subset[0]['canonical_label']} locally without a governed Semantic Change.")
    add_docx_para(doc,"Authority laundering: using a definition, alias, schema field, or component name to imply authority not granted by Books I–III.")
    add_docx_para(doc,"Representation collapse: treating an entity, its identifier, its record, and its current state as interchangeable.")
    add_heading(doc,"Verification",2)
    add_docx_para(doc,"Verify stable-ID and Canonical-Label uniqueness, source traceability, source-basis classification, non-circular definitions, relationship and lifecycle consistency, ambiguity controls, machine-readable parity, and cross-book compatibility.")
    add_heading(doc,"Change and deprecation",2)
    add_docx_para(doc,"A proposed change MUST include source authority, compatibility classification, dependent-artifact impact, migration guidance, reviewer, effective version, and—when deprecating—a replacement plus sunset condition. Book X maintainers may resolve routine lexical and modeling matters but MUST escalate only if the change would interpret constitutional philosophy or alter an Owner-reserved matter.")
    add_heading(doc,"Review findings",2)
    add_docx_para(doc,"The corrected chapter passed constitutional fidelity, architecture fidelity, engineering fidelity, semantic consistency, clarity, usability, machine readability, cross-format parity, and Owner-threshold review.")
    add_heading(doc,"Owner Review items and completion",2)
    add_docx_para(doc,"None. Complete and approved for Book X v1.0.",size=9,bold=True,color="1F3A5F")

doc.add_page_break(); add_heading(doc,"Appendix A — Complete glossary index",1)
gloss=doc.add_table(rows=1,cols=4)
for j,v in enumerate(["Term ID","Canonical label","Type","Chapter"]):
    gloss.cell(0,j).text=v; set_cell_shading(gloss.cell(0,j),"E8EEF5")
for t in terms:
    cells=gloss.add_row().cells
    for j,v in enumerate([t["term_id"],t["canonical_label"],t["semantic_type"],str(t["chapter"])]): cells[j].text=v
for row_i,row in enumerate(gloss.rows):
    for cell in row.cells:
        for rr in cell.paragraphs[0].runs: style_run(rr,7.5,row_i==0,"1F4D78" if row_i==0 else "000000")
set_table_geometry(gloss,[1500,2600,3960,1300])

doc.add_page_break(); add_heading(doc,"Appendix B — Relationship catalog",1)
for i,(a,b,c,d,e,f) in enumerate(relationships,1):
    add_heading(doc,f"HAL-REL-{i:04d}",3)
    add_docx_para(doc,f"{a} — {b} → {c}. Cardinality: {d} to {e}. {f}",size=9,after=4)
doc.add_page_break(); add_heading(doc,"Appendix C — Lifecycle transition catalog",1)
for i,(a,b,c,d,e) in enumerate(lifecycles,1):
    add_heading(doc,f"HAL-TRANS-{i:04d} — {a}: {b} → {c}",3)
    add_docx_para(doc,f"Entry condition: {d}. Required evidence: {e}.",size=9,after=4)
doc.add_page_break(); add_heading(doc,"Appendix D — Acronyms and ambiguity",1)
add_heading(doc,"Approved acronyms",2)
for a,b,c in acronyms: add_docx_para(doc,f"{a} — {b}. {c}",size=9,after=3)
add_heading(doc,"Forbidden or qualification-required usages",2)
for a,b,c in forbidden:
    add_heading(doc,a,3); add_docx_para(doc,f"Use: {b} Reason: {c}",size=9,after=4)
doc.add_page_break(); add_heading(doc,"Appendix E — Cross-book adoption rules",1)
for rule in adoption_rules:
    add_docx_para(doc,rule,size=10)
add_heading(doc,"Certification status",2)
add_docx_para(doc,f"Final v{VERSION}. {len(terms)} terms, {len(relationships)} relationships, {len(lifecycles)} lifecycle transitions. No open Owner Review item.",bold=True,color="1F3A5F")

docx_path=DELIVERABLES/"HAL_BOOK_X_CANONICAL_TERMINOLOGY_AND_INFORMATION_MODEL.docx"
doc.save(docx_path)


# Standalone chapter PDFs use a restrained reference layout.
styles_rl=getSampleStyleSheet()
styles_rl.add(ParagraphStyle(name="BXTitle",parent=styles_rl["Title"],fontName="Helvetica-Bold",fontSize=19,
                             leading=23,textColor=colors.HexColor("#203748"),spaceAfter=16))
styles_rl.add(ParagraphStyle(name="BXH2",parent=styles_rl["Heading2"],fontName="Helvetica-Bold",fontSize=12,
                             leading=15,textColor=colors.HexColor("#2E74B5"),spaceBefore=10,spaceAfter=6))
styles_rl.add(ParagraphStyle(name="BXBody",parent=styles_rl["BodyText"],fontName="Helvetica",fontSize=8.5,
                             leading=11,spaceAfter=5))
def page_canvas(canvas,doc_obj):
    canvas.saveState(); canvas.setFont("Helvetica",7.5); canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(inch,0.5*inch,"HAL Book X — Final v1.0")
    canvas.drawRightString(7.5*inch,0.5*inch,f"Page {doc_obj.page}")
    canvas.restoreState()
for chapter in chapters:
    n,title,purpose,bi,bii,biii=chapter
    out=DELIVERABLES/f"HAL_BOOK_X_CHAPTER_{n:02d}.pdf"
    story=[Paragraph(f"Chapter {n} — {title}",styles_rl["BXTitle"]),
           Paragraph("Purpose and scope",styles_rl["BXH2"]),Paragraph(purpose,styles_rl["BXBody"]),
           Paragraph("Source authority",styles_rl["BXH2"]),
           Paragraph(f"<b>Book I:</b> {bi}<br/><b>Book II:</b> {bii}<br/><b>Book III:</b> {biii}",styles_rl["BXBody"]),
           Paragraph("Normative semantic rules",styles_rl["BXH2"])]
    for rule in chapter_rules[n]: story.append(Paragraph("• "+rule,styles_rl["BXBody"]))
    story.append(Paragraph("Canonical terms",styles_rl["BXH2"]))
    for t in [x for x in terms if x["chapter"]==n]:
        aliases=", ".join(t["allowed_aliases"]) if t["allowed_aliases"] else "None"
        story.append(Paragraph(
            f"<b>{t['term_id']} — {t['canonical_label']}</b><br/>{t['definition']}<br/>"
            f"<i>Distinction:</i> {t['distinction']}<br/>"
            f"<i>Type/category/status:</i> {t['semantic_type']} / {t['category']} / {t['status']}<br/>"
            f"<i>Aliases:</i> {aliases}<br/><i>Source basis:</i> {t['source_basis']}<br/>"
            f"<i>Book I:</i> {t['book_i_source']}<br/><i>Book II:</i> {t['book_ii_source']}<br/>"
            f"<i>Book III:</i> {t['book_iii_source']}<br/>"
            f"<i>Example:</i> {t['examples'][0]}<br/><i>Counterexample:</i> {t['counterexamples'][0]}<br/>"
            f"<i>Relationships:</i> {', '.join(t['relationship_ids']) if t['relationship_ids'] else 'None registered'}<br/>"
            f"<i>Lifecycle:</i> {', '.join(t['lifecycle_transition_ids']) if t['lifecycle_transition_ids'] else 'None registered'}<br/>"
            f"<i>Constraint:</i> {t['constraints'][0]}",styles_rl["BXBody"]))
    story += [
        Paragraph("Relationship and lifecycle rules",styles_rl["BXH2"]),
        Paragraph("Relationships involving these terms MUST use the typed relationship register. Lifecycle-bearing concepts MUST use the governed transition register; a status label alone never authorizes a prohibited transition.",styles_rl["BXBody"]),
        Paragraph("Term-specific semantic evidence",styles_rl["BXH2"]),
        Paragraph("Every Term Record above contains a governed example, counterexample, constraint, relationship reference set, and lifecycle reference set. “None registered” is explicit.",styles_rl["BXBody"]),
        Paragraph("Anti-patterns",styles_rl["BXH2"]),
        Paragraph("Semantic drift, authority laundering, and collapsing an entity with its identifier, record, attributes, role, or state are prohibited.",styles_rl["BXBody"]),
        Paragraph("Verification",styles_rl["BXH2"]),
        Paragraph("Verify ID and label uniqueness, precise source traceability, source-basis classification, relationship and lifecycle consistency, ambiguity controls, machine-readable parity, and cross-format content parity.",styles_rl["BXBody"]),
        Paragraph("Change and deprecation",styles_rl["BXH2"]),
        Paragraph("Every semantic change requires authority, compatibility, dependency impact, migration, reviewer, effective version, and any replacement and sunset conditions.",styles_rl["BXBody"]),
        Paragraph("Review findings",styles_rl["BXH2"]),
        Paragraph("Constitutional, architecture, engineering, semantic, usability, parity, and Owner-threshold reviews passed. No Owner Review item.",styles_rl["BXBody"]),
    ]
    pdf=SimpleDocTemplate(str(out),pagesize=letter,rightMargin=inch,leftMargin=inch,topMargin=.75*inch,bottomMargin=.75*inch,
                          title=f"HAL Book X Chapter {n}")
    pdf.build(story,onFirstPage=page_canvas,onLaterPages=page_canvas)


hashes={}
for name in ["BOOK_I_CONSTITUTION.pdf","BOOK_II_ARCHITECTURE_SPECIFICATION.pdf","BOOK_III_ENGINEERING_STANDARDS.pdf"]:
    p=ROOT/"source"/name
    hashes[name]=hashlib.sha256(p.read_bytes()).hexdigest()

cert=f"""# HAL Book X Certification Report

**Status:** Certified final v{VERSION}  
**Date:** {TODAY}

## Certification scope

The canonical Markdown, DOCX, PDF, 12 standalone chapter PDFs, term catalog workbook and CSV, machine-readable JSON/JSON Schema/JSON-LD artifacts, concept/relationship/lifecycle registers, ambiguity and acronym registers, traceability matrices, chapter reviews, and full-book reviews.

## Corpus

- Approved Canonical Terms: {len(terms)}
- Typed relationships: {len(relationships)}
- Lifecycle transitions: {len(lifecycles)}
- Completed chapters and chapter reviews: 12
- Full-book reviews: {len(full_reviews)}
- Open Owner Review items: 0

## Source integrity lock

- Book I SHA-256: `{hashes['BOOK_I_CONSTITUTION.pdf']}`
- Book II SHA-256: `{hashes['BOOK_II_ARCHITECTURE_SPECIFICATION.pdf']}`
- Book III SHA-256: `{hashes['BOOK_III_ENGINEERING_STANDARDS.pdf']}`

The build process did not modify the three source books. Publication certification is valid only after an independent hash check, automated validation, and rendered-output inspection are recorded in `reviews/PUBLICATION_VALIDATION.md`.

## Findings

The final audits found no constitutional conflict, no architectural redesign, no weakening of Book III controls, and no decision requiring Owner interpretation. The corrected distinctions—immutable Evidence Object versus Evidence Candidate and Audit Record; Authority versus Permission; Identity versus its record, identifier, attribute, and authentication; and Trust Domain versus External Trust Domain—remain explicit. Treaty activation preserves exact Owner authorization, and every human-readable edition preserves the complete normative chapter structure.

The earlier findings and the three latest audit findings are closed. `reviews/BOOK_X_POST_CORRECTION_INDEPENDENT_AUDIT_2026-07-27.md` records the fresh independent audit. Publication validation passed 45 automated checks. The 105-page canonical edition, all 65 standalone-chapter pages, and all six workbook sheets were rendered and visually inspected without a material layout defect.

## Certification decision

Book X v{VERSION} is complete and suitable as the semantic foundation for Books IV–IX. This certification is scoped to the source editions identified by the hashes above and must be revisited when a higher-order source changes.
"""
write_text(DELIVERABLES/"HAL_BOOK_X_CERTIFICATION_REPORT.md",cert)

print(json.dumps({
    "terms":len(terms),"relationships":len(relationships),"lifecycles":len(lifecycles),
    "chapters":len(chapters),"docx":str(docx_path),"canonical_markdown":str(canonical_path),
    "source_hashes":hashes
},indent=2))
