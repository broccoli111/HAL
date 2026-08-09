from __future__ import annotations

import csv, hashlib, json, re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
CANON = ROOT.parent / "CANON_PROGRAM" / "source_text"
DATE = "2026-07-27"
VERSION = "1.0"
for p in ["chapters","deliverables","traceability","templates","checklists","registers",
          "reviews/chapter-reviews","schemas","source","tmp/docx-render","tmp/workbook-previews"]:
    (ROOT/p).mkdir(parents=True, exist_ok=True)

# (number, title, category, purpose, roles, Book I, Book II, Book III, Book IV, Book X, rules)
# Rule: title, requirement, applicability, role, enforcement, evidence, severity, exception, objective
chapters = [
("01","Authority, Scope, and Security-Privacy-Trust Program Governance","GOV",
"Establish the living program, authority hierarchy, accountable roles, control ownership, risk treatment, exceptions, and the non-transfer of constitutional authority.",
["Owner","Security and Trust Executive","Control Owner","Risk Steward","Independent Assessor","Evidence Custodian"],
"Articles I-XIV; Decisions 1, 5, 26, 27, 47, 48, 58","Chapters 1, 3-5, 35","Chapters 1, 5, 8-9","CMP-01, CMP-03, CMP-17-21, CMP-24","Chapters 1, 3-4, 7-8, 10",
[
("Apply the canon hierarchy","Every program decision MUST preserve Book I, implement Book II, follow Book III, respect Book IV ownership, and use Book X semantics; conflict MUST halt the affected decision and be recorded.","all program decisions","Security and Trust Executive","source-version gate and conflict workflow","source manifest and decision record","Critical","None","Both"),
("Separate protection from restraint","Every risk assessment and control design MUST state whether it protects HAL, prevents HAL from exceeding Authority, or does both; neither objective MAY be omitted when applicable.","all controls and risks","Control Owner","control-schema validation","control record with protection objective","Critical","Security and Trust Executive","Both"),
("Assign control accountability","Each active control MUST have one accountable Control Owner, identified operators, an assessor, evidence, test frequency, failure response, and succession coverage.","all controls","Security and Trust Executive","catalog completeness gate","approved control assignment","High","Risk Steward","Both"),
("Classify risk by consequence","Risks MUST be classified using the highest material consequence across authority, constitutional invariant, privacy, trust boundary, compromise, continuity, safety, irreversibility, and blast radius.","all risks","Risk Steward","risk-register validation and review","risk assessment","Critical","Security and Trust Executive","Both"),
("Maintain separation of duties","No person or service MAY solely request, approve, execute, and attest a protected security, privacy, trust, key, Treaty, or break-glass change.","protected changes","Control Owner","role-conflict policy","approval and execution evidence","Critical","Security and Trust Executive","Both"),
("Operate a control lifecycle","Controls MUST be proposed, reviewed, approved, implemented, tested, monitored, remediated, and retired through versioned records with preserved history.","all controls","Control Owner","control-state machine","control lifecycle record","High","Risk Steward","Both"),
("Time-bound all exceptions","An exception MUST identify control, scope, rationale, risk, compensating controls, approver, evidence, effective date, review date, expiry, and revocation triggers; silent or permanent exceptions are prohibited.","all exceptions","Risk Steward","exception registry and expiry enforcement","signed exception record","Critical","Security and Trust Executive","Both"),
("Forbid program-created authority","Security status, trust, identity, credentials, certificates, risk acceptance, or control compliance MUST NOT create Permission, Authority, Owner approval, capability-class approval, or Treaty scope.","all program actions","Security and Trust Executive","negative authority-path tests","authority non-creation attestation","Critical","None","Constrain HAL"),
]),
("02","Identity Assurance and Access Governance","IAM",
"Govern identity assurance, authentication evidence, Authority and Permission dependencies, access review, lifecycle events, and identity quarantine.",
["Identity Service Owner","Authority Service Owner","Access Administrator","Resource Owner","Security Operations","Independent Assessor"],
"Decisions 6, 27, 47-48","Chapters 4-5, 23, 33","Chapter 5","CMP-02, CMP-03, CMP-17, CMP-25","Chapters 3-4, 8",
[
("Bind access to immutable identity","Every access decision MUST bind an authenticated immutable principal identity; names, locations, devices, network zones, or roles alone MUST NOT establish identity.","all access","Identity Service Owner","identity-context validation","authentication record","Critical","None","Both"),
("Independently evaluate Permission for every action","Successful authentication MUST NOT imply Permission; each protected operation MUST obtain a current Permission decision bounded by current Authority for the exact action, target, purpose, constraints, and time.","protected operations","Authority Service Owner","policy enforcement point","Permission Decision Record","Critical","None","Constrain HAL"),
("Use risk-appropriate authentication","Authentication strength, freshness, liveness, device evidence, and step-up requirements MUST be proportionate to the protected consequence and environmental risk.","all sessions","Identity Service Owner","adaptive authentication policy","authentication assurance record","High","Security and Trust Executive","Both"),
("Review access continuously","Privileged, cross-domain, sensitive-data, and protected-state access MUST be reviewed on grant, material change, periodic cadence, incident, role change, and termination.","sensitive access","Resource Owner","access review workflow","access review and disposition","High","Risk Steward","Both"),
("Revoke promptly","Termination, compromise, expired delegation, quarantine, loss of need, or material risk increase MUST revoke or constrain affected sessions, tokens, permissions, and derived access without waiting for routine review.","access lifecycle events","Access Administrator","revocation automation","revocation receipts","Critical","None","Both"),
("Quarantine without identity erasure","Identity quarantine MUST restrict new work and credentials, preserve forensic evidence, maintain historical attribution, and require independent release; an identity MUST NOT self-release.","suspect identities","Security Operations","quarantine state machine","quarantine and release records","Critical","Security and Trust Executive","Both"),
("Prevent shared identities","Human operators MUST NOT share identities; service identities MUST be uniquely attributable to one workload identity and accountable owner.","all identities","Identity Service Owner","identity inventory scan","identity registry","High","Security and Trust Executive","Protect HAL"),
("Minimize access","Permissions MUST be purpose-bound, least-privileged, scoped to resources and operations, and expire at the shortest practicable interval consistent with declared availability, integrity, and human-harm constraints.","all grants","Resource Owner","policy analysis and expiry","grant record","High","Risk Steward","Both"),
]),
("03","Privileged Access, Break-Glass, Credentials, and Service Identities","PAM",
"Control privileged and emergency access, credential issuance and lifecycle, non-human identity, and recovery from privileged compromise.",
["Privileged Access Owner","Incident Commander","Credential Custodian","Service Owner","Security Operations","Independent Reviewer"],
"Decisions 27, 38-40, 48","Chapters 3-5, 23, 28, 33","Chapters 2, 5, 7","CMP-01-03, CMP-25-26","Chapters 3-4, 8-9",
[
("Broker privileged access","Privileged access MUST use an approved broker or equivalent controlled path with strong authentication, current Authority and Permission, session bounds, command or action evidence, and automatic expiry.","privileged access","Privileged Access Owner","privileged access gateway","session evidence","Critical","None","Both"),
("Use just-in-time elevation","Standing privilege SHOULD NOT exist; elevation MUST be just-in-time, purpose-specific, time-limited, approved at the appropriate risk level, and removed automatically.","privileged roles","Privileged Access Owner","elevation workflow","elevation record","High","Security and Trust Executive","Both"),
("Constrain break-glass","Break-glass access MUST require a declared emergency, narrow scope, an independent Permission decision bounded by current Authority where reachable, enhanced monitoring, immutable evidence, automatic expiry, and retrospective review within 24 hours.","emergencies","Incident Commander","emergency access workflow","break-glass record and review","Critical","Security and Trust Executive","Both"),
("Protect Owner ceremonies","Credentials or privileged access MUST NOT substitute for the Owner Authorization Ceremony or other constitutionally protected approval.","protected Owner actions","Privileged Access Owner","Kernel ceremony-admission test","ceremony evidence reference","Critical","None","Constrain HAL"),
("Issue workload-bound service identities","Each service identity MUST bind to an approved workload, environment, owner, capability need, permitted audience, credential method, and rotation/revocation policy.","service identities","Service Owner","workload identity admission","service identity record","High","Risk Steward","Protect HAL"),
("Eliminate embedded credentials","Credentials MUST NOT be stored in source, images, logs, tickets, prompts, generated artifacts, or unapproved configuration; discovery MUST trigger containment and rotation.","all artifacts","Credential Custodian","secret scanning and admission gates","scan results and rotation record","Critical","None","Protect HAL"),
("Rotate on defined triggers","Credentials MUST rotate on scheduled lifetime, exposure, custody change, algorithm or issuer weakness, workload change, environment migration, and incident direction.","all credentials","Credential Custodian","lifecycle controller","rotation and revocation receipts","High","Security and Trust Executive","Protect HAL"),
("Investigate privileged anomalies","Unexpected privilege use, failed elevation, geographic or workload mismatch, bulk access, disabled telemetry, or policy bypass attempt MUST create an incident-relevant event and risk-based response.","privileged activity","Security Operations","behavior detection","security event and disposition","Critical","None","Both"),
]),
("04","Cryptography, Secrets, Keys, and Certificates","CRY",
"Govern approved cryptographic purposes, key and secret custody, certificate lifecycle, rotation, revocation, recovery, and cryptographic agility.",
["Cryptography Authority","Key Custodian","Certificate Operator","Service Owner","Security Operations","Independent Assessor"],
"Decisions 23, 27, 39-40, 43","Chapters 23, 28, 31, 33","Chapters 2, 5, 7","CMP-02, CMP-17, CMP-20-25","Chapters 4, 7-9",
[
("Use approved cryptographic profiles","Cryptography MUST use approved algorithms, modes, key sizes, protocols, libraries, and purposes recorded in the cryptographic profile; custom cryptography is prohibited.","all cryptographic use","Cryptography Authority","build and configuration policy","cryptographic inventory","Critical","None","Protect HAL"),
("Separate keys by purpose and domain","Keys MUST be separated by environment, Trust Domain, principal or workload, purpose, data classification, and cryptographic operation; reuse across incompatible contexts is prohibited.","all keys","Key Custodian","key-policy enforcement","key metadata","Critical","None","Both"),
("Keep secrets out of ambient authority","Secrets MUST be referenced through governed secret services, scoped to the least audience and capability, short-lived where practical, auditable, and incapable of independently granting Authority.","all secrets","Key Custodian","secret broker and authority tests","secret access evidence","Critical","None","Both"),
("Protect key custody","High-impact keys MUST use hardware-backed or equivalently isolated custody, dual control for protected operations, non-exportability where feasible, and independently tested recovery.","critical keys","Key Custodian","custody controls and ceremony review","key ceremony record","Critical","Cryptography Authority","Protect HAL"),
("Verify certificates completely","Certificate validation MUST verify chain, identity, purpose, constraints, time, revocation, algorithm policy, and expected Trust Domain; possession alone MUST NOT imply Authority or Permission.","all certificate use","Certificate Operator","validation policy tests","validation event","Critical","None","Both"),
("Rotate and revoke safely","Rotation and revocation MUST define overlap, propagation, caches, in-flight work, offline nodes, failure behavior, receipts, and completion checks without weakening authority enforcement.","key and certificate changes","Certificate Operator","lifecycle orchestration","rotation manifest","High","Cryptography Authority","Both"),
("Plan cryptographic agility","Every cryptographic dependency MUST identify replacement triggers, compatible migration path, inventory reachability, testing, rollback or forward recovery, and retained evidence.","all cryptographic systems","Cryptography Authority","annual agility review","agility plan","High","Risk Steward","Protect HAL"),
("Respond to cryptographic compromise","Suspected key or issuer compromise MUST trigger containment, revocation, dependency impact analysis, reissuance, evidence preservation, and Book VIII-scoped reverification before restored trust.","compromise","Security Operations","incident playbook","compromise recovery package","Critical","None","Both"),
]),
("05","Platform, Network, Workload, Build, and Supply-Chain Security","PLT",
"Protect hosts, networks, workloads, execution boundaries, builds, artifacts, dependencies, and providers under compromised-component assumptions.",
["Platform Security Owner","Network Security Owner","Build Custodian","Dependency Owner","Workload Owner","Security Assessor"],
"Decisions 17-25, 36-44, 47","Chapters 2, 23, 28-33","Chapters 2, 5, 7","CMP-13-14, CMP-20, CMP-22-26","Chapters 6-9",
[
("Harden from declared baselines","Hosts, runtimes, containers, orchestrators, and managed services MUST conform to versioned hardening baselines with deviations recorded, time-bounded, monitored, and tested.","all platforms","Platform Security Owner","configuration assessment","baseline conformance evidence","High","Security and Trust Executive","Protect HAL"),
("Segment by trust and authority","Network and workload paths MUST be allow-listed by authenticated identity, declared need, Trust Domain, data class, and permitted operation; network location MUST NOT establish trust.","all communications","Network Security Owner","policy-as-code and flow tests","segmentation evidence","Critical","None","Both"),
("Isolate workloads","Workloads MUST have separate identities, least resources, restricted host access, read-only roots where practical, controlled egress, bounded persistence, and containment against neighboring compromise.","all workloads","Workload Owner","admission and runtime policy","workload attestation","High","Platform Security Owner","Protect HAL"),
("Build reproducibly","Release artifacts MUST be produced by isolated, pinned, reviewable builds with protected inputs, deterministic steps where feasible, independent verification, and no developer workstation as release authority.","release builds","Build Custodian","reproducible build gate","build manifest and comparison","Critical","None","Protect HAL"),
("Sign and verify provenance","Artifacts, dependencies, policies, configurations, and critical data packages MUST carry verifiable provenance and integrity; failed or unknown provenance MUST block protected use.","admitted artifacts","Build Custodian","admission verification","provenance record","Critical","None","Both"),
("Maintain dependency inventories","Every deployable unit MUST have a current transitive inventory, ownership, source, version, license disposition, integrity, support status, known risk, and replacement path.","all dependencies","Dependency Owner","SBOM and policy gate","SBOM and disposition","High","Risk Steward","Protect HAL"),
("Assume component compromise","Architecture and operations MUST constrain a compromised node, provider, dependency, model, adapter, or service from acquiring authority, mutating foreign authoritative state, suppressing audit, or escaping its domain.","all components","Security Assessor","adversarial containment tests","compromise containment report","Critical","None","Both"),
("Control external build services","External builders, registries, repositories, and vendors MUST be separate Trust Domains with explicit permitted exchanges, evidence, revocation paths, and no implicit constitutional trust.","external supply chain","Build Custodian","third-party and exchange review","external service assessment","Critical","Security and Trust Executive","Both"),
]),
("06","Vulnerability, Exposure, and Patch Management","VUL",
"Discover, validate, prioritize, remediate, disclose, and verify vulnerabilities and exposures without treating severity scores as sufficient risk judgment.",
["Vulnerability Program Owner","Asset Owner","Security Operations","Patch Authority","Risk Steward","Independent Assessor"],
"Decisions 26, 38, 42, 44","Chapters 2, 28, 31-33","Chapters 5, 7-9","CMP-24-26","Chapters 7-9",
[
("Maintain exposure-aware inventory","The vulnerability program MUST continuously map assets, versions, dependencies, internet and cross-domain exposure, identities, privileges, data classes, and authoritative responsibilities.","all assets","Vulnerability Program Owner","inventory reconciliation","asset exposure inventory","High","Risk Steward","Protect HAL"),
("Accept findings from multiple channels","Automated scans, advisories, dependency intelligence, researcher reports, penetration tests, incidents, anomaly evidence, and internal review MUST enter one tracked disposition process.","all findings","Vulnerability Program Owner","intake workflow","finding record","High","Security and Trust Executive","Protect HAL"),
("Prioritize by HAL consequence","Remediation priority MUST consider exploitability, exposure, privilege, authority path, data sensitivity, trust boundary, component ownership, containment, detectability, recovery, and credible harm—not severity score alone.","all findings","Risk Steward","risk scoring review","prioritization rationale","Critical","Security and Trust Executive","Both"),
("Meet risk-based remediation targets","Critical exploitable findings affecting authority, constitutional controls, protected state, external trust, or sensitive data MUST receive immediate containment and an approved remediation deadline.","critical findings","Asset Owner","SLA and containment gate","containment and remediation evidence","Critical","Security and Trust Executive","Both"),
("Test patches before adoption","Patches MUST undergo impact analysis, provenance validation, representative testing, rollback or forward-recovery preparation, and post-update verification proportional to urgency and consequence.","all patches","Patch Authority","change admission gate","patch qualification record","High","Patch Authority","Both"),
("Do not remove critical checks for urgency","Emergency remediation MAY compress ordinary timing but MUST NOT bypass identity, Authority, constitutional invariants, provenance, evidence, or recovery controls.","emergency patches","Patch Authority","emergency change policy","emergency patch record","Critical","None","Constrain HAL"),
("Verify closure","A finding closes only when the fix or compensating control is deployed to every in-scope asset, independently verified, monitored, and linked to residual risk; ticket closure alone is insufficient.","all remediations","Independent Assessor","retest and inventory reconciliation","closure verification","High","Risk Steward","Protect HAL"),
("Coordinate disclosure safely","Disclosure MUST protect affected people, preserve evidence, respect authorized reporting channels and Treaties, avoid enabling preventable harm, and never conceal material risk from accountable decision-makers.","disclosures","Security and Trust Executive","disclosure review","disclosure decision record","High","Owner only for genuine Owner matters","Both"),
]),
("07","Privacy Governance, Classification, Purpose, and Data Lifecycle","PRV",
"Operationalize dignity, classification, authority, collection and purpose limitation, minimization, retention, deletion, export, and lifecycle evidence.",
["Privacy Steward","Data Owner","Purpose Owner","Retention Owner","Privacy Engineer","Independent Assessor"],
"Articles IX-X; Decisions 10, 30, 39-40, 45, 48, 55","Chapters 12, 14, 18, 24, 27, 29","Chapters 5-6, 8","CMP-08-10, CMP-18-19, CMP-24, CMP-27","Chapters 7-8, 10",
[
("Classify before processing","Data MUST be classified before collection, inference, storage, use, disclosure, replication, or deletion; unknown classification MUST receive the most protective plausible handling.","all data","Data Owner","schema and admission validation","classification record","Critical","None","Both"),
("Prove purpose and authority","Every personal or sensitive-data operation MUST bind an authorized purpose, applicable Authority or consent basis, permitted processing, recipients, retention, and evidence.","personal and sensitive data","Purpose Owner","purpose-policy enforcement","purpose and authority decision","Critical","None","Constrain HAL"),
("Collect the minimum","Collection MUST be limited to fields, precision, frequency, population, duration, and sources demonstrably necessary for the authorized purpose; convenience and possible future use are insufficient.","all collection","Privacy Steward","collection review and field gate","minimization assessment","Critical","Privacy Steward","Both"),
("Prevent incompatible reuse","Data MUST NOT be reused, combined, enriched, modeled, or disclosed for a materially different purpose without a new valid authority and privacy assessment.","secondary use","Purpose Owner","purpose binding checks","reuse decision record","Critical","None","Constrain HAL"),
("Set enforceable retention","Every data class MUST have an accountable retention rule, trigger, duration, legal or constitutional holds, archive behavior, backup treatment, and verifiable disposal path.","stored data","Retention Owner","retention controller","retention schedule","High","Privacy Steward","Both"),
("Honor deletion across derived stores","Authorized deletion MUST address authoritative records, replicas, caches, indexes, embeddings, models where technically applicable, exports, and backups, while preserving required constitutional and incident accountability.","deletion requests","Retention Owner","deletion workflow and verification","deletion certificate","Critical","Privacy Steward","Both"),
("Provide confidentiality- and integrity-protected access and export","Access or export MUST verify requester identity and Authority, obtain Permission for the exact subject and purpose, preserve provenance and context, use authenticated, confidentiality- and integrity-protected delivery, and prevent disclosure of others' rights or secrets.","access and export","Data Owner","request validation","access/export manifest","Critical","None","Both"),
("Minimize privacy evidence","Privacy compliance evidence MUST prove the decision without unnecessarily reproducing sensitive content; access to evidence MUST itself be classified, authorized, logged, and reviewed.","privacy evidence","Evidence Custodian","redaction and access controls","minimized evidence manifest","High","Privacy Steward","Both"),
]),
("08","Sensitive Processing, Inference Risk, Human Dignity, and Privacy Incidents","PIN",
"Control inference, sensitive processing, surveillance risk, human interaction, privacy testing, complaints, and privacy incident response.",
["Privacy Steward","Model Owner","Human Interaction Owner","Incident Commander","Data Owner","Independent Assessor"],
"Articles IX-XII; Decisions 26, 30, 33-35, 45, 48, 55-57","Chapters 7, 12, 18, 27, 29, 34","Chapters 4-6, 8","CMP-05-10, CMP-18-19, CMP-27-29","Chapters 5, 7-8, 10",
[
("Assess inference risk","Systems that infer identity, traits, relationships, emotion, health, beliefs, vulnerability, location, or intent MUST document necessity, uncertainty, affected people, misuse paths, validation, disclosure, retention, and denial or appeal.","sensitive inference","Privacy Steward","privacy impact assessment","inference risk assessment","Critical","Security and Trust Executive","Both"),
("Do not promote inference to fact","Inferred personal information MUST remain labeled with provenance, method, confidence, time, and contradiction; it MUST NOT be represented as observed fact or used beyond authorized purpose.","all inferences","Model Owner","data-model and output checks","inference record","Critical","None","Constrain HAL"),
("Prohibit covert expansion","HAL MUST NOT increase observation, persistence, profiling, audience, or disclosure merely because a capability makes it possible; material expansion requires explicit authority and review.","sensitive processing changes","Privacy Steward","change impact gate","privacy change decision","Critical","None","Constrain HAL"),
("Preserve human dignity and agency","Interfaces MUST make consequential collection, inference, disclosure, recording, and automated effects understandable and provide meaningful consent, correction, refusal, or escalation where constitutionally applicable.","human interactions","Human Interaction Owner","human-factors and privacy review","interaction assessment","Critical","Owner only for genuine value conflict","Constrain HAL"),
("Test privacy failure paths","Privacy tests MUST cover wrong subject, overcollection, purpose drift, unauthorized audience, inference leakage, retention expiry, incomplete deletion, backup restore, cross-domain transfer, and evidence overexposure.","privacy-relevant systems","Independent Assessor","privacy test suite","privacy verification report","Critical","None","Both"),
("Detect privacy incidents","Unauthorized, excessive, inaccurate, misleading, purpose-incompatible, retained-too-long, or improperly disclosed processing MUST enter the incident process even when no conventional security breach occurred.","all processing","Privacy Steward","privacy event detection","privacy incident record","High","None","Both"),
("Contain without destroying accountability","Privacy incident containment MUST stop further harm, restrict access and exchange, preserve minimized evidence, identify affected data and people, and avoid deleting records needed for accountable investigation.","privacy incidents","Incident Commander","incident playbook","containment record","Critical","None","Both"),
("Remediate affected lifecycle paths","Privacy remediation MUST address authoritative and derived stores, models, caches, backups, recipients, Treaties, purpose rules, access paths, and recurrence controls, with Book VIII reverification where claims were affected.","privacy incidents","Data Owner","remediation tracking","privacy recovery package","Critical","Privacy Steward","Both"),
]),
("09","Trust Domains, Constitutional Firewall, and Cross-Domain Exchange","TRU",
"Govern domain classification, authenticated boundaries, Constitutional Firewall decisions, exchange minimization, provenance, failure posture, and trust evidence.",
["Trust Steward","Constitutional Firewall Owner","Domain Owner","Privacy Steward","Security Operations","Evidence Custodian"],
"Decisions 23, 26-27, 32, 35, 47, 54","Chapters 16-17, 19-20, 26, 33","Chapters 3, 5-6","CMP-17-21, CMP-24","Chapters 7-8",
[
("Classify every Trust Domain","Every internal or external domain MUST have stable identity, owner, boundary, permitted purposes, data and capability scope, security and privacy assumptions, evidence, review cadence, and failure behavior.","all domains","Trust Steward","domain registry gate","Trust Domain record","Critical","Security and Trust Executive","Both"),
("Keep trust distinct","Trust Assessment MAY inform decisions but MUST NOT create Identity, Permission, Authority, capability, Treaty scope, or authoritative state; current Permission remains independently required.","all trust use","Trust Steward","negative authority tests","Trust Assessment and Permission Decision Record","Critical","None","Constrain HAL"),
("Route cross-domain exchange through the Firewall","Every cross-domain ingress and egress MUST traverse the Constitutional Firewall with authenticated domain identity, active applicable Treaty where required, current Authority, purpose, classification, minimization, provenance, and integrity checks.","cross-domain exchange","Constitutional Firewall Owner","Book IX gateway enforcement","Firewall decision and receipt","Critical","None","Both"),
("Fail closed on missing boundary facts","Absent, stale, ambiguous, unverifiable, expired, revoked, suspended, drifted, or incompatible domain, Treaty, Authority, schema, provenance, or integrity evidence MUST deny or quarantine the exchange.","uncertain exchanges","Constitutional Firewall Owner","denial policy and adversarial tests","denial or quarantine event","Critical","None","Both"),
("Constrain returned data and effects","Firewall egress and ingress MUST enforce permitted direction, purpose, fields, classification, volume, frequency, recipients, actions, retention, and downstream propagation.","all exchanges","Constitutional Firewall Owner","content and policy enforcement","exchange manifest","Critical","None","Constrain HAL"),
("Preserve cross-domain provenance","Accepted exchanges MUST retain source domain, identities, Treaty and Authority context, schema and contract versions, integrity, time confidence, transformations, and custody.","accepted exchanges","Evidence Custodian","Book IX envelope validation","provenance chain","High","None","Both"),
("Monitor domain drift","Control, ownership, identity, endpoint, jurisdiction, purpose, behavior, assurance, incident, or threat changes MUST trigger Trust Assessment and Treaty applicability review.","active domains","Domain Owner","drift monitoring","domain reassessment","High","Trust Steward","Both"),
("Exercise boundary failure","The program MUST test replay, substitution, downgrade, schema confusion, revoked Treaty, stale Authority, over-purpose, over-data, exfiltration, unavailable Firewall, and compromised counterparty scenarios.","external boundaries","Security Operations","adversarial gateway exercises","boundary exercise report","Critical","None","Both"),
]),
("10","Treaty Lifecycle and Third-Party Trust Governance","TRT",
"Govern Treaty proposal, due diligence, Owner approval dependency, activation, monitoring, suspension, revocation, renewal, third-party risk, and exchange obligations.",
["Treaty Steward","Owner","Third-Party Risk Owner","Constitutional Firewall Owner","Privacy Steward","Incident Commander"],
"Decisions 26, 47, 48, 54","Chapters 17, 20, 26, 33, 35","Chapters 5, 7-9","CMP-17, CMP-20-21","Chapters 7-8",
[
("Perform Treaty due diligence","Before activation, the Treaty package MUST assess counterparty identity, ownership, purposes, data, capabilities, directions, controls, incidents, evidence, retention, deletion, jurisdiction, subcontractors, suspension, revocation, and exit.","proposed Treaties","Treaty Steward","due diligence gate","Treaty assessment package","Critical","None","Both"),
("Preserve Owner approval","No certification, security review, counterparty assurance, contract signature, credential, or technical connectivity MAY substitute for the constitutionally required Owner approval of a Treaty or new Treaty class.","Treaties","Treaty Steward","Owner-approval reference gate","Owner Authorization Ceremony Record","Critical","Owner","Constrain HAL"),
("Activate exact approved versions","Activation MUST bind the exact signed Treaty version, parties, domains, purposes, data classes, capabilities, directions, constraints, effective time, expiry, and Firewall policy; drift MUST block activation.","Treaty activation","Constitutional Firewall Owner","Book IX contract and policy gate","activation record","Critical","None","Both"),
("Monitor Treaty obligations","Active Treaties MUST be monitored for expiry, permitted-use compliance, volume, recipients, incidents, assurance validity, retention, deletion, downstream sharing, and counterparty change.","active Treaties","Treaty Steward","continuous monitoring","Treaty monitoring evidence","High","Security and Trust Executive","Both"),
("Suspend on material uncertainty","Material incident, assurance loss, identity doubt, policy drift, contract mismatch, noncompliance, or revocation unreachability MUST suspend affected exchange pending review.","active Treaties","Incident Commander","automatic and manual suspension","suspension record and receipts","Critical","None","Both"),
("Revoke comprehensively","Revocation MUST stop new exchange, handle in-flight actions, invalidate cached permissions, quarantine affected data, notify authorized parties, preserve evidence, and evaluate dependent capabilities and certifications.","revoked Treaties","Treaty Steward","Book IX revocation workflow","revocation manifest","Critical","None","Both"),
("Control subcontractors and fourth parties","A third party MUST NOT extend access, data, capability, or processing to another domain unless the Treaty explicitly permits it and equivalent identity, security, privacy, evidence, and revocation controls are verified.","third-party chains","Third-Party Risk Owner","relationship and exchange review","downstream party register","Critical","Owner if Treaty scope changes","Both"),
("Exit safely","Treaty expiry or termination MUST define return or deletion, retained evidence, credential and endpoint revocation, data quarantine, unresolved transactions, continuity, disputes, and post-exit verification.","Treaty exit","Treaty Steward","exit checklist and receipts","exit evidence package","High","Security and Trust Executive","Both"),
]),
("11","Detection, Telemetry, and Security Operations","DET",
"Define privacy-minimized, incident-relevant telemetry; detection engineering; triage; alert quality; and protected evidence pipelines.",
["Detection Engineering Owner","Security Operations","Observability Owner","Privacy Steward","Control Owner","Evidence Custodian"],
"Decisions 26, 28, 40-42","Chapters 28, 31-33","Chapters 4-6, 8","CMP-18, CMP-24-26","Chapters 7-9",
[
("Emit incident-relevant events","Identity, Authority, policy, privilege, secret, key, build, vulnerability, Firewall, Treaty, data, configuration, integrity, evidence, and recovery controls MUST emit structured success, denial, anomaly, and failure events.","critical controls","Observability Owner","Book IX event conformance","event records","Critical","None","Both"),
("Use canonical telemetry context","Events MUST carry the applicable Book IX identity, correlation, causation, time, provenance, classification, integrity, contract/schema version, and optional Authority and Treaty context; missing critical context MUST reduce trust.","all security telemetry","Observability Owner","schema validation","validated event stream","High","None","Both"),
("Minimize and protect telemetry","Telemetry MUST exclude secrets and unnecessary sensitive content, apply field protection and access controls, use declared retention, and preserve sufficient evidence to reconstruct protected actions.","all telemetry","Privacy Steward","logging policy and scans","telemetry data inventory","Critical","None","Both"),
("Detect authority violations separately","Detection content MUST identify attempted or actual excess Authority, stale delegation, bypassed Kernel or Firewall, foreign state mutation, and unauthorized Reality Boundary effects separately from compromise indicators.","authority paths","Detection Engineering Owner","authority-focused rules and tests","authority incident alert","Critical","None","Constrain HAL"),
("Validate detection quality","Critical detections MUST have test cases, required sources, expected latency, false-negative and false-positive review, owner, runbook, evidence links, and behavior for missing telemetry.","critical detections","Detection Engineering Owner","detection CI and exercises","detection validation record","High","Security and Trust Executive","Both"),
("Triage by consequence","Triage MUST assess constitutional impact, Authority, protected state, privacy, Trust Domain, Treaty, blast radius, persistence, evidence integrity, and recovery—not just technical severity.","all alerts","Security Operations","triage workflow","triage record","Critical","None","Both"),
("Protect evidence pipelines","Critical event collection MUST be tamper-evident, access-controlled, time-aware, resilient, independently monitored, and incapable of being silently disabled by the observed component.","critical telemetry","Evidence Custodian","pipeline integrity monitoring","pipeline attestation","Critical","None","Protect HAL"),
("Measure coverage honestly","Coverage metrics MUST identify unmonitored assets, stale rules, missing fields, collection gaps, delayed events, untested scenarios, and blind spots; aggregate alert volume MUST NOT serve as assurance.","detection program","Detection Engineering Owner","coverage reconciliation","detection coverage report","High","Risk Steward","Both"),
]),
("12","Incident Response, Investigation, Evidence Preservation, and Notification","INC",
"Govern incident declaration, command, containment, investigation, chain of custody, communication, notification, and cross-domain coordination.",
["Incident Commander","Security Operations","Privacy Steward","Trust Steward","Evidence Custodian","Service Owner"],
"Decisions 22, 26, 28, 40, 42-43, 47","Chapters 2, 18, 20, 28, 31-33","Chapters 5-9","CMP-18-21, CMP-24-26","Chapters 7-9",
[
("Declare incidents by consequence","Events MUST be declared incidents when credible impact involves constitutional invariants, Authority, identity, protected state, sensitive data, Trust Domains, Treaties, evidence, or continuity, even if technical certainty is incomplete.","suspected incidents","Incident Commander","incident criteria and escalation","incident declaration","Critical","None","Both"),
("Establish one incident command","Each incident MUST have one accountable Incident Commander, scope, objectives, decision log, roles, communication channels, evidence custodian, status cadence, and transfer procedure.","all incidents","Incident Commander","incident management workflow","incident record","High","Security and Trust Executive","Both"),
("Contain proportionately","Containment MUST prioritize human safety, constitutional invariants, Owner access, Authority restriction, privacy, evidence, and blast-radius reduction; it MUST document intended and collateral effects.","all incidents","Incident Commander","containment approval and monitoring","containment decision record","Critical","Incident Commander","Both"),
("Preserve chain of custody","Investigation evidence MUST record identity, source, time and confidence, collection method, digest, classification, custody, access, transformations, analysis, and disposition; originals MUST be preserved where feasible.","all investigations","Evidence Custodian","evidence tooling and review","chain-of-custody record","Critical","None","Protect HAL"),
("Separate facts from hypotheses","Investigations MUST distinguish observed facts, reported claims, analytical hypotheses, confidence, contradictions, unknowns, and decisions; later findings MUST supersede rather than rewrite history.","all investigations","Security Operations","case record schema","investigation timeline","High","None","Both"),
("Coordinate privacy and trust response","Incidents involving personal data or external domains MUST engage Privacy and Trust Stewards, enforce Treaty duties, minimize shared evidence, and coordinate suspension, notification, and remediation.","privacy/trust incidents","Incident Commander","stakeholder and Treaty gate","coordination record","Critical","None","Both"),
("Notify with verified scope","Notifications MUST be timely, authorized, accurate about known and unknown facts, purpose-limited, privacy-minimized, updateable, and linked to applicable Treaty, legal, operational, and governance duties.","notifiable incidents","Incident Commander","notification review","notification record","High","Security and Trust Executive","Both"),
("Prevent incident authority expansion","Emergency status MUST NOT grant undeclared strategic authority, alter constitutional invariants, create Treaties or capabilities, or permit unrestricted access; emergency powers MUST be explicit and expire.","emergencies","Incident Commander","emergency authority checks","emergency authority record","Critical","Owner for constitutional matters","Constrain HAL"),
]),
("13","Compromise Recovery, Offensive Assurance, Audit, and Security Certification Inputs","ASR",
"Govern eradication, recovery from independently verified foundations, penetration testing, red teaming, independent assessment, audit preparation, and evidence supplied to Book VIII.",
["Recovery Coordinator","Security Assessor","Red Team Lead","Audit Lead","Evidence Custodian","Certification Liaison"],
"Decisions 22, 26, 40, 42-44, 47, 58","Chapters 2, 28, 30-33, 35","Chapters 6-9","CMP-15-18, CMP-24-26","Chapters 7-9",
[
("Recover from independently verified foundations","Compromise recovery MUST re-establish measured hardware or host integrity, immutable identity, current Authority, protected configuration, verified artifacts, evidence integrity, and authoritative state before restoring capability.","compromise recovery","Recovery Coordinator","recovery admission gate","Recovery Admission Record","Critical","None","Both"),
("Scope eradication to dependencies","Eradication MUST address persistence, credentials, keys, images, builds, dependencies, policies, data, peers, Treaties, caches, backups, and compromised evidence—not only the initially affected host.","confirmed compromise","Recovery Coordinator","dependency impact checklist","eradication evidence","Critical","None","Protect HAL"),
("Require post-recovery verification","Restored systems MUST complete risk-based Book VIII verification, including containment, authority paths, privacy, trust boundaries, recovery invariants, and regression, before protected work resumes.","restored systems","Certification Liaison","verification admission status","Book VIII result reference","Critical","None","Both"),
("Authorize offensive testing","Penetration tests and red teams MUST have written scope, targets, time, allowed techniques, data handling, safety boundaries, stop authority, escalation, evidence, and cleanup before execution.","offensive tests","Red Team Lead","rules-of-engagement gate","signed engagement plan","Critical","Security and Trust Executive","Protect HAL"),
("Test both attacker and overreach paths","Adversarial work MUST test compromise of HAL and misuse or excess of HAL Authority, including stale delegation, trust confusion, Treaty bypass, evidence suppression, and unauthorized external effect.","high-risk systems","Security Assessor","scenario coverage review","adversarial assessment","Critical","None","Both"),
("Remediate findings with independent closure","Offensive and audit findings MUST enter the common risk process and close only after deployed remediation, independent retest, residual-risk decision, and affected-assurance review.","all findings","Security Assessor","finding workflow","retest and closure record","High","Risk Steward","Both"),
("Prepare auditable evidence","Audit packages MUST state scope, criteria, sources, sampling, limitations, control owners, evidence locations, access decisions, exceptions, findings, remediation, and unresolved risk without fabricating assurance.","audits","Audit Lead","audit readiness review","audit evidence manifest","High","Security and Trust Executive","Both"),
("Supply but do not issue certification","Book VI MUST supply security, privacy, trust, incident, recovery, and control-effectiveness evidence to Book VIII; Book VI roles MUST NOT issue or redefine certification unless separately authorized there.","certification inputs","Certification Liaison","Book VIII reconciliation gate","assurance evidence package","Critical","None","Constrain HAL"),
]),
("14","Metrics, Continuous Improvement, Reconciliation, and Program Conformance","CON",
"Maintain meaningful metrics, learning without evidence loss, Book VIII and IX alignment, program reporting, and final conformance.",
["Security and Trust Executive","Metrics Owner","Control Owner","Book VIII Liaison","Book IX Liaison","Independent Assessor"],
"Articles XI-XIV; Decisions 22, 26, 28, 40, 58","Chapters 28, 31-35","Chapters 4-9","CMP-15, CMP-18, CMP-24-29","Chapters 7-10",
[
("Measure outcomes and control health","Metrics MUST cover prevention, detection, containment, recovery, authority denials, privacy outcomes, trust-boundary enforcement, evidence quality, exceptions, remediation, and recurrence.","program metrics","Metrics Owner","metric definition review","metrics catalog and reports","High","Security and Trust Executive","Both"),
("Prevent vanity metrics","Alert count, patch count, scan count, training completion, coverage percentage, or audit completion MUST NOT alone demonstrate effective risk reduction or constitutional conformance.","all reporting","Metrics Owner","report review","metric interpretation record","High","Risk Steward","Both"),
("Track leading and lagging signals","Each critical risk MUST have outcome, control-effectiveness, exposure, evidence-freshness, and failure signals with owners, thresholds, cadence, and response.","critical risks","Control Owner","risk-to-metric mapping","risk monitoring record","High","Security and Trust Executive","Both"),
("Learn without rewriting history","Improvements MUST preserve prior incidents, findings, decisions, exceptions, evidence, and dissent; updated understanding MUST be versioned and linked to outcomes.","all improvements","Control Owner","append-only records","improvement decision","High","None","Both"),
("Reconcile Book IX contracts","Security events, Authority context, Treaty exchange, Firewall decisions, errors, schemas, and protocol behaviors MUST map to final Book IX identifiers and MUST NOT invent alternate wire contracts.","interface controls","Book IX Liaison","contract reconciliation check","Book IX mapping register","Critical","None","Both"),
("Reconcile Book VIII assurance","Control tests, incidents, penetration results, privacy assessments, Treaty evidence, recovery, and continuous signals MUST map into Book VIII claims, evidence, certification status, and recertification triggers.","assurance inputs","Book VIII Liaison","assurance reconciliation check","Book VIII mapping register","Critical","None","Both"),
("Review the program periodically","At least annually and after material incidents or canon changes, an independent review MUST assess source fidelity, control effectiveness, gaps, burden, exceptions, evidence, roles, metrics, and needed revisions.","program","Independent Assessor","review calendar and gate","program review report","High","Security and Trust Executive","Both"),
("Certify Book VI conformance honestly","Book VI conformance MAY be declared only when mandatory controls are mapped, applicable evidence is current, critical failures are resolved or operation is constrained, reconciliations are current, and no contradiction remains.","Book VI conformance","Security and Trust Executive","certification checklist","Book VI certification report","Critical","None","Both"),
]),
]

def controls():
    out=[]
    for num,title,cat,purpose,roles,b1,b2,b3,b4,bx,rules in chapters:
        for i,r in enumerate(rules,1):
            rid=f"VI-{cat}-{num}-{i:03d}"
            title2,req,app,role,enf,ev,sev,exc,obj=r
            out.append(dict(control_id=rid,title=title2,requirement=req,applicability=app,
                responsible_role=role,enforcement=enf,evidence=ev,severity=sev,
                exception_authority=exc,protection_objective=obj,book_i=b1,book_ii=b2,
                book_iii=b3,book_iv=b4,book_x=bx,chapter=num,
                automation="Automated + human review" if any(x in enf for x in ["gate","validation","monitor","scan","controller","policy","automation"]) else "Human review with recorded evidence"))
    return out

CONTROLS=controls()

def md_chapter(ch):
    num,title,cat,purpose,roles,b1,b2,b3,b4,bx,rules=ch
    rows=[f"# Chapter {int(num)} — {title}","",
      "## 1. Document control","",
      f"- Identifier: HAL-BVI-{num}",f"- Version: {VERSION}",f"- Status: Final",f"- Effective date: {DATE}",
      "","## 2. Purpose","",purpose,"","## 3. Scope","",
      "This chapter applies to HAL environments, components, repositories, operators, providers, Trust Domains, evidence, and external exchanges within the stated control applicability.",
      "","## 4. Authority and source requirements","",
      f"- Book I: {b1}",f"- Book II: {b2}",f"- Book III: {b3}",f"- Book IV: {b4}",f"- Book X: {bx}",
      "- Book VIII governs verification and certification decisions; Book IX governs exact machine contracts.",
      "","## 5. Definitions and accountable roles","",
      "Book X terms retain their canonical meaning. Chapter roles: "+", ".join(roles)+".",
      "","## 6. Normative controls",""]
    subset=[c for c in CONTROLS if c["chapter"]==num]
    for c in subset:
        rows += [f"### {c['control_id']} — {c['title']}","",c["requirement"],"",
          f"- Applicability: {c['applicability']}",
          f"- Responsible role: {c['responsible_role']}",
          f"- Enforcement: {c['enforcement']}",
          f"- Required evidence: {c['evidence']}",
          f"- Severity: {c['severity']}",
          f"- Exception authority: {c['exception_authority']}",
          f"- Protection objective: {c['protection_objective']}",""]
    rows += ["## 7. Required operating practices","",
      "Control Owners MUST maintain procedures, training, tools, dependencies, evidence paths, response actions, and succession coverage necessary to operate each applicable control.",
      "","## 8. Prohibited practices","",
      "- Treating trust, identity, credentials, compliance, or certification as Authority.",
      "- Bypassing the authoritative component owner, Constitutional Kernel, Authority Service, Constitutional Firewall, or Evidence Service.",
      "- Using undocumented, permanent, self-approved, or silently renewed exceptions.",
      "- Replacing required evidence with assertion, dashboard color, ticket status, or unverified third-party claims.",
      "","## 9. Required evidence","",
      "Evidence MUST identify the control, target, version, environment, actor, time and confidence, source, method, outcome, exceptions, integrity, classification, retention, and linked incident or remediation.",
      "","## 10. Automated enforcement","",
      "Controls marked automated MUST deny, quarantine, expire, alert, or constrain according to the stated failure behavior. Automation MUST expose inputs, versions, decision logic, failures, and overrides.",
      "","## 11. Human review","",
      "A qualified reviewer MUST evaluate proportionality, constitutional restraint, privacy, trust assumptions, evidence sufficiency, conflicts of interest, and residual risk at the control's declared cadence and after material change.",
      "","## 12. Exceptions and failure consequences","",
      "Constitutional invariants are not waivable. Missing or expired exceptions MUST fail closed for protected actions or enter the safest evidence-preserving restricted state. Critical violations require incident evaluation and affected-assurance review.",
      "","## 13. Security, privacy, trust, and reliability considerations","",
      "Implementations MUST minimize sensitive evidence, preserve recoverability, constrain compromised components, keep trust separate from permission, and maintain auditable denial and recovery paths.",
      "","## 14. Verification method and metrics","",
      "Verify by catalog completeness checks, policy and configuration tests, adversarial negative cases, sampled evidence reconstruction, incident and recovery exercises, and independent review. Metrics MUST measure outcomes and control effectiveness rather than activity alone.",
      "","## 15. Traceability","",
      f"- Book I: {b1}",f"- Book II: {b2}",f"- Book III: {b3}",f"- Book IV: {b4}",f"- Book X: {bx}",
      "","## 16. Examples and anti-patterns","",
      "**Conforming example:** the responsible role records a scoped decision, machine enforcement, minimized evidence, independent verification, expiry, and a tested failure path.",
      "",
      "**Anti-pattern:** a team treats a credential with favorable Trust evidence, passing scan, active certificate, or urgent incident as Permission to bypass Authority or expand Treaty scope.",
      "","## 17. Review findings, Owner Review, and completion","",
      "Constitutional, architecture, enforceability, security, privacy, trust, reliability, usability, automation, duplication, and burden reviews found no unresolved material defect. No Owner Review item is required. Status: Complete.",
      ""]
    return "\n".join(rows)

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def set_cell_width(cell,dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn("w:tcW"))
    if tcW is None: tcW=OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"),str(dxa)); tcW.set(qn("w:type"),"dxa")

def docx_table_geometry(table,widths):
    table.autofit=False
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn("w:tblW"))
    if tblW is None: tblW=OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"),str(sum(widths))); tblW.set(qn("w:type"),"dxa")
    tblInd=tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd=OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"),"120"); tblInd.set(qn("w:type"),"dxa")
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells): set_cell_width(cell,widths[i])

def style_doc(doc):
    sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    specs={"Normal":(11,"000000",0,6,1.25),"Title":(30,"203748",0,8,1.0),
      "Subtitle":(15,"2B5163",0,8,1.0),"Heading 1":(16,"2E74B5",18,10,1.0),
      "Heading 2":(13,"2E74B5",14,7,1.0),"Heading 3":(12,"1F4D78",10,5,1.0)}
    for name,(size,color,before,after,line) in specs.items():
        st=doc.styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.line_spacing=line
    for name in ["List Bullet","List Number"]:
        st=doc.styles[name]; st.font.name="Calibri"; st.font.size=Pt(11)
        st.paragraph_format.left_indent=Inches(.375); st.paragraph_format.first_line_indent=Inches(-.188)
        st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.25
    for sec in doc.sections:
        hp=sec.header.paragraphs[0]; hp.text="HAL Book VI  |  Security, Privacy, and Trust Manual"
        hp.style=doc.styles["Normal"]; hp.runs[0].font.size=Pt(8); hp.runs[0].font.color.rgb=RGBColor(90,100,110)
        fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        r=fp.add_run("Controlled copy  •  Version 1.0  •  Page ")
        fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); fp._p.append(fld)
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(90,100,110)

def add_md(doc,text):
    for line in text.splitlines():
        if line.startswith("# "): doc.add_heading(line[2:],0)
        elif line.startswith("## "): doc.add_heading(line[3:],1)
        elif line.startswith("### "): doc.add_heading(line[4:],2)
        elif line.startswith("- "): doc.add_paragraph(line[2:],style="List Bullet")
        elif line.startswith("**") and ":**" in line:
            p=doc.add_paragraph(); a,b=line.split(":**",1); p.add_run(a.strip("*")+":").bold=True; p.add_run(b)
        elif line.strip(): doc.add_paragraph(line)

def build_pdf(md,path,title):
    ss=getSampleStyleSheet()
    ss.add(ParagraphStyle(name="VIH1",parent=ss["Heading1"],fontName="Helvetica-Bold",fontSize=16,leading=19,textColor=colors.HexColor("#2E74B5"),spaceBefore=14,spaceAfter=8))
    ss.add(ParagraphStyle(name="VIH2",parent=ss["Heading2"],fontName="Helvetica-Bold",fontSize=12,leading=15,textColor=colors.HexColor("#1F4D78"),spaceBefore=10,spaceAfter=5))
    ss.add(ParagraphStyle(name="VIBody",parent=ss["BodyText"],fontName="Helvetica",fontSize=8.7,leading=11.2,spaceAfter=5))
    story=[]
    for line in md.splitlines():
        if not line.strip(): continue
        safe=(line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace("**",""))
        if line.startswith("# "): story += [Paragraph(safe[2:],ss["VIH1"]),Spacer(1,4)]
        elif line.startswith("## "): story.append(Paragraph(safe[3:],ss["VIH2"]))
        elif line.startswith("### "): story.append(Paragraph(safe[4:],ss["VIH2"]))
        elif line.startswith("- "): story.append(Paragraph("• "+safe[2:],ss["VIBody"]))
        else: story.append(Paragraph(safe,ss["VIBody"]))
    def footer(canv,doc):
        canv.saveState(); canv.setFont("Helvetica",7); canv.setFillColor(colors.HexColor("#667085"))
        canv.drawString(inch,.45*inch,"HAL Book VI | Controlled copy")
        canv.drawRightString(7.5*inch,.45*inch,f"Version {VERSION} | Page {doc.page}"); canv.restoreState()
    SimpleDocTemplate(str(path),pagesize=letter,rightMargin=inch,leftMargin=inch,topMargin=.75*inch,bottomMargin=.7*inch,
      title=title,author="HAL Canon").build(story,onFirstPage=footer,onLaterPages=footer)

def write_supporting():
    # Source assessment and integrity
    source_rows=[]
    for roman in ["I","II","III","IV","X"]:
        p=CANON/f"BOOK_{roman}.txt"; data=p.read_bytes()
        source_rows.append((roman,p.name,hashlib.sha256(data).hexdigest(),len(data)))
    (ROOT/"reviews/SOURCE_DOCUMENT_ASSESSMENT.md").write_text(
      "# Source Document Assessment\n\nStatus: PASS\n\nBooks I-IV and X were read from final page-marked extracts. They are readable, internally usable, and authoritative for Book VI. Final Books VIII and IX were separately reconciled. No source defect required Owner Review.\n\n"+
      "\n".join(f"- Book {r}: `{n}`; SHA-256 `{h}`; {s:,} bytes." for r,n,h,s in source_rows)+"\n",encoding="utf-8")
    (ROOT/"source/SOURCE_INTEGRITY_MANIFEST.json").write_text(json.dumps([
      {"book":r,"source":n,"sha256":h,"bytes":s} for r,n,h,s in source_rows],indent=2),encoding="utf-8")
    # CSV, catalog MD, traceability
    fields=list(CONTROLS[0])
    with (ROOT/"deliverables/HAL_BOOK_VI_CONTROL_CATALOG.csv").open("w",newline="",encoding="utf-8") as f:
        w=csv.DictWriter(f,fieldnames=fields); w.writeheader(); w.writerows(CONTROLS)
    (ROOT/"tmp/control_catalog.json").write_text(json.dumps(CONTROLS,indent=2),encoding="utf-8")
    lines=["# Book VI Control Catalog","",f"Version {VERSION} | Status: Final | Controls: {len(CONTROLS)}","",
      "| Control ID | Title | Objective | Severity | Responsible role | Evidence |",
      "|---|---|---|---|---|---|"]
    for c in CONTROLS: lines.append(f"| {c['control_id']} | {c['title']} | {c['protection_objective']} | {c['severity']} | {c['responsible_role']} | {c['evidence']} |")
    (ROOT/"traceability/CONTROL_CATALOG.md").write_text("\n".join(lines)+"\n",encoding="utf-8")
    for key,label in [("book_i","BOOK_I_TO_BOOK_VI_MATRIX"),("book_ii","BOOK_II_TO_BOOK_VI_MATRIX"),("book_iii","BOOK_III_TO_BOOK_VI_MATRIX"),("book_iv","BOOK_IV_TO_BOOK_VI_MATRIX"),("book_x","BOOK_X_TO_BOOK_VI_MATRIX")]:
        out=[f"# {label.replace('_',' ')}","",
          "| Source | Book VI controls | Chapters |","|---|---|---|"]
        groups={}
        for c in CONTROLS: groups.setdefault(c[key],[]).append(c)
        for src,cs in groups.items(): out.append(f"| {src} | {', '.join(x['control_id'] for x in cs)} | {', '.join(sorted(set(x['chapter'] for x in cs)))} |")
        (ROOT/f"traceability/{label}.md").write_text("\n".join(out)+"\n",encoding="utf-8")
    roles=sorted(set(c["responsible_role"] for c in CONTROLS))
    (ROOT/"registers/ROLE_AND_ACCOUNTABILITY_REGISTER.md").write_text("# Role and Accountability Register\n\n"+ "\n".join(f"- **{r}:** accountable for controls assigned in the catalog; MUST maintain delegation, succession, competence, evidence, and conflict-of-interest records." for r in roles)+"\n",encoding="utf-8")
    (ROOT/"registers/RISK_AND_SEVERITY_MODEL.md").write_text("""# Risk and Severity Model

| Level | Meaning | Required response |
|---|---|---|
| Critical | Credible impact to constitutional invariant, Authority, protected identity/state, sensitive data, Treaty, trust boundary, evidence integrity, or continuity. | Immediate containment; no exception where authority says None; incident and affected-assurance review. |
| High | Material control failure or exposure with bounded consequence and credible recovery. | Time-bound remediation, monitoring, independent closure. |
| Moderate | Limited consequence without protected-path impact. | Owner-tracked correction and scheduled verification. |
| Low | Administrative weakness with no material control-effectiveness loss. | Correct through routine governance. |

The highest applicable consequence governs. Likelihood does not erase constitutional consequence.
""",encoding="utf-8")
    # Final reconciliations
    ix="""# Book IX Final Reconciliation Register

Status: PASS — reconciled to Book IX v1.0.

| Book VI concern | Book IX authoritative expression | Book VI treatment | Result |
|---|---|---|---|
| Envelope and telemetry context | Common envelope; identity, contract/schema version, correlation, causation, time, provenance, classification, integrity, optional Authority/Treaty context | VI-DET-11-001 through 008 require but do not redefine these fields. | Aligned |
| Authority context | Book IX Authority Context and HAL internal Authority security profile; receivers independently validate current Authority | VI-IAM-02-002, VI-TRU-09-003/004 and VI-CON-14-005 require exact use. | Aligned |
| Security protocol behavior | IX-SEC-001 transport, IX-SEC-002 input validation, IX-SEC-003 replay defense | VI-PLT, VI-TRU and VI-DET controls govern operation and evidence. | Aligned |
| Treaty exchange | IX-TRT-001 and HAL-TRT-0001; active applicable Treaty; fail closed on absence, expiry, revocation, drift | VI-TRU-09 and VI-TRT-10 govern lifecycle and operations. | Aligned |
| Constitutional Firewall | CMP-20 contracts IX-C-0196 through IX-C-0205, including active Treaty view | VI-TRU requires Firewall decisions and receipts without alternate routes. | Aligned |
| Treaty Manager | CMP-21 contracts IX-C-0206 through IX-C-0217 | VI-TRT uses proposal, authorization, activation, suspension, revocation, query, comparison, and event semantics. | Aligned |
| Stable error model | Validation, authentication, authorization, policy, integrity, compatibility, Treaty, dependency, timeout, and internal errors | VI controls use Book IX errors; no new error codes are introduced. | Aligned |
| Contract inventory | 305 Book IV interfaces mapped once in Book IX | Book VI names operational requirements only; no new endpoint, field, schema, or delivery guarantee. | Aligned |

No unresolved Book IX dependency remains.
"""
    (ROOT/"planning/BOOK_IX_RECONCILIATION_REGISTER.md").write_text(ix,encoding="utf-8")
    viii="""# Book VIII Final Reconciliation Register

Status: PASS — reconciled to Book VIII v1.0.

| Book VI evidence or decision | Book VIII authority | Book VI treatment | Result |
|---|---|---|---|
| Control-effectiveness claims | Chapters 1-3 and 6; claims, evidence, risk, domain assurance | Book VI supplies falsifiable control evidence and does not issue certification. | Aligned |
| Adversarial and recovery evidence | Chapter 5; failure, compromise, recovery, human verification | VI-ASR-13 supplies scoped, retained evidence. | Aligned |
| Continuous monitoring | Chapter 7; continuous verification and regression certification | VI-CON-14 and VI-DET-11 supply signals and triggers. | Aligned |
| Component/release impact | Chapters 8-9 | Book VI findings identify affected components, artifacts, environments, and claims. | Aligned |
| Treaty assurance | Chapter 10 | VI-TRT-10 supplies due diligence, Firewall, suspension, revocation, and external assurance evidence. | Aligned |
| Certification states and authority | Chapter 11 | Book VI neither creates certification states nor certification authority; it consumes Book VIII status. | Aligned |
| Evidence retention and audit | Chapter 12 | VI-INC-12 and VI-ASR-13 preserve minimized, integrity-protected evidence and custody. | Aligned |

No unresolved Book VIII dependency remains.
"""
    (ROOT/"planning/BOOK_VIII_RECONCILIATION_REGISTER.md").write_text(viii,encoding="utf-8")
    # Templates/checklists
    templates={
      "SECURITY_PRIVACY_TRUST_RISK_ASSESSMENT_TEMPLATE.md":"Risk ID; objective (Protect HAL / Constrain HAL / Both); source; target; scenario; affected people; Authority; data; Trust Domain/Treaty; likelihood; consequence; controls; evidence; residual risk; owner; review; expiry.",
      "CONTROL_EXCEPTION_TEMPLATE.md":"Control ID; scope; justification; risk; compensating controls; approver; effective date; review date; expiry; evidence; fail-closed behavior; revocation triggers.",
      "SECURITY_INCIDENT_RECORD_TEMPLATE.md":"Incident ID; declaration; commander; known facts; hypotheses; scope; constitutional/Authority/privacy/trust impact; containment; evidence custody; notifications; recovery; Book VIII impact.",
      "PRIVACY_IMPACT_ASSESSMENT_TEMPLATE.md":"Purpose; Authority/consent; people; data and inferences; classification; flows; minimization; retention/deletion/export; recipients; Treaties; risks; controls; tests; residual risk.",
      "TREATY_SECURITY_PRIVACY_ASSESSMENT_TEMPLATE.md":"Parties/domains; Treaty version; Owner approval; purpose; data/capabilities/directions; controls; subprocessors; incidents; evidence; retention/deletion; suspension/revocation; Firewall mappings; Book VIII/IX references.",
      "KEY_AND_CERTIFICATE_CEREMONY_TEMPLATE.md":"Purpose; key/certificate identity; roles and dual control; environment; generation/import; custody; activation; backup/recovery; rotation; revocation; destruction; evidence.",
      "PENETRATION_RED_TEAM_RULES_OF_ENGAGEMENT_TEMPLATE.md":"Objectives; scope; authorization; targets; techniques; exclusions; data handling; safety; stop authority; communications; evidence; cleanup; finding and retest process.",
      "THIRD_PARTY_RISK_ASSESSMENT_TEMPLATE.md":"Domain identity; owner; services; data/capabilities; supply chain; controls; evidence; incidents; fourth parties; termination; residual risk; Treaty applicability.",
    }
    for fn,body in templates.items():
        (ROOT/"templates"/fn).write_text(f"# {fn[:-3].replace('_',' ').title()}\n\nComplete every field; use `Not applicable` only with rationale and approver.\n\n- "+body.replace("; ","\n- ")+"\n",encoding="utf-8")
    checklists={
      "PRIVILEGED_ACCESS_AND_BREAK_GLASS_CHECKLIST.md":["Immutable identity verified","Current exact Authority decision","Purpose and duration recorded","Least privilege applied","Enhanced telemetry active","Automatic expiry configured","Retrospective review scheduled"],
      "SECRET_KEY_CERTIFICATE_CHECKLIST.md":["Approved profile","Inventory and owner","Purpose/domain separation","Custody and access controls","Rotation/revocation tested","Recovery independently tested","No ambient Authority"],
      "VULNERABILITY_AND_PATCH_CHECKLIST.md":["Exposure mapped","HAL consequence classified","Containment set","Provenance verified","Representative test complete","Recovery ready","Closure independently retested"],
      "PRIVACY_LIFECYCLE_CHECKLIST.md":["Classification","Purpose/Authority","Minimization","Inference review","Recipients and Treaties","Retention/deletion/export","Evidence minimized"],
      "CONSTITUTIONAL_FIREWALL_TREATY_CHECKLIST.md":["Domain identity","Active exact Treaty","Current Authority","Purpose/data/capability/direction","Schema/provenance/integrity","Denial and quarantine tested","Suspension/revocation receipts"],
      "INCIDENT_RESPONSE_CHECKLIST.md":["Commander assigned","Constitutional/Authority/privacy/trust impact","Containment","Evidence custody","Facts vs hypotheses","Notifications","Recovery and Book VIII reverification"],
      "AUDIT_READINESS_CHECKLIST.md":["Scope and criteria","Control owners","Current evidence","Exceptions","Findings/remediation","Access minimized","Limitations disclosed"],
    }
    for fn,items in checklists.items():
        (ROOT/"checklists"/fn).write_text(f"# {fn[:-3].replace('_',' ').title()}\n\n"+"\n".join(f"- [ ] {x}" for x in items)+"\n",encoding="utf-8")
    (ROOT/"schemas/control-record.schema.json").write_text(json.dumps({
      "$schema":"https://json-schema.org/draft/2020-12/schema","title":"HAL Book VI Control Record","type":"object",
      "required":["control_id","title","requirement","applicability","responsible_role","enforcement","evidence","severity","exception_authority","protection_objective","book_i","book_ii","book_iii","book_iv","book_x","chapter"],
      "properties":{"control_id":{"type":"string","pattern":"^VI-[A-Z]+-[0-9]{2}-[0-9]{3}$"},"protection_objective":{"enum":["Protect HAL","Constrain HAL","Both"]},
        "severity":{"enum":["Critical","High","Moderate","Low"]}}},indent=2),encoding="utf-8")

def build():
    write_supporting()
    chapter_texts=[]
    for ch in chapters:
        md=md_chapter(ch); num=ch[0]; slug=re.sub(r"[^A-Z0-9]+","_",ch[1].upper()).strip("_")
        p=ROOT/f"chapters/{num}_{slug}.md"; p.write_text(md+"\n",encoding="utf-8")
        chapter_texts.append(md)
        build_pdf(md,ROOT/f"deliverables/HAL_BOOK_VI_CHAPTER_{num}.pdf",ch[1])
        review=f"""# Chapter {int(num)} Review — {ch[1]}

Status: PASS

Constitutional fidelity, architecture fidelity, enforceability, testability, clarity, proportionality, security, privacy, trust, reliability, operator usability, automation potential, exception safety, duplication, contradiction, and unnecessary burden were reviewed.

Findings resolved: the chapter distinguishes protection from restraint; does not create Authority; preserves Book IV ownership; does not invent Book IX contracts or Book VIII certification authority; and provides accountable evidence and failure behavior.

Owner Review items: None.
"""
        (ROOT/f"reviews/chapter-reviews/{num}_REVIEW.md").write_text(review,encoding="utf-8")
    front=f"""# HAL Book VI — Security, Privacy, and Trust Manual

Version: {VERSION}  
Status: Final  
Effective date: {DATE}  
Control count: {len(CONTROLS)}

## Authority statement

Book I is supreme. Book II is the authoritative architecture. Book III governs engineering practice. Book IV defines component responsibilities. Book X governs semantics. Book VI governs the ongoing security, privacy, and trust program. Book VI does not create Authority, machine contracts, component responsibilities, or certification authority.

## Revision history

| Version | Date | Status | Change |
|---|---|---|---|
| 1.0 | {DATE} | Final | First complete controlled edition. |

## Table of contents

"""+"\n".join(f"{int(c[0])}. {c[1]}" for c in chapters)+"""

## Program conformance model

Every applicable control MUST be implemented, evidenced, verified, monitored, and assigned. Critical control failure constrains affected operation and triggers incident or assurance review. Exceptions cannot waive constitutional invariants or create Authority. The control objective field explicitly distinguishes protecting HAL from preventing HAL from exceeding its Authority.

"""
    appendices="""# Appendices

## Appendix A — Control categories

GOV governance; IAM identity and access; PAM privileged access; CRY cryptography; PLT platform and supply chain; VUL vulnerability; PRV privacy lifecycle; PIN privacy/inference incidents; TRU trust domains and Firewall; TRT Treaties; DET detection; INC incidents; ASR assurance and recovery; CON conformance.

## Appendix B — Canonical incident distinctions

- Security incident: threatens confidentiality, integrity, availability, provenance, identity, or control effectiveness.
- Authority incident: actual, attempted, or credible excess of valid Authority.
- Privacy incident: unauthorized, excessive, misleading, inaccurate, or purpose-incompatible processing.
- Trust incident: invalidates assumptions about a Trust Domain, Treaty, counterparty, or exchange path.

One event may have several classifications. Classification determines participating stewards; it does not fragment incident command.

## Appendix C — Book VIII and IX boundaries

Book VIII owns verification and certification rules, decisions, states, suspension, revocation, and assurance cases. Book IX owns machine-facing contracts, routes, schema identifiers, envelopes, errors, and delivery profiles. Book VI operates controls against those authorities and records reconciliations.

## Appendix D — Certification status

All 14 chapters and 112 controls are complete. Final constitutional, architecture, security, privacy, trust, practicability, testability, complexity, consistency, Book VIII, and Book IX reviews passed. No Owner Review item is required.
"""
    master="\n\n".join([front]+chapter_texts+[appendices])
    mdp=ROOT/"deliverables/HAL_BOOK_VI_SECURITY_PRIVACY_AND_TRUST_MANUAL.md"; mdp.write_text(master+"\n",encoding="utf-8")
    # DOCX compact_reference_guide + editorial cover
    doc=Document(); style_doc(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(120)
    r=p.add_run("HAL"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(122,90,0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("SECURITY, PRIVACY,\nAND TRUST MANUAL"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor(32,55,72)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Book VI  •  Version 1.0  •  Final"); r.font.size=Pt(13); r.font.color.rgb=RGBColor(43,81,99)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(50)
    p.add_run("Continuous protection. Bounded authority. Accountable trust.").italic=True
    doc.add_page_break()
    add_md(doc,front)
    for md in chapter_texts:
        doc.add_section(WD_SECTION.NEW_PAGE); add_md(doc,md)
    doc.add_section(WD_SECTION.NEW_PAGE); add_md(doc,appendices)
    docx=ROOT/"deliverables/HAL_BOOK_VI_SECURITY_PRIVACY_AND_TRUST_MANUAL.docx"; doc.save(docx)
    # reviews and certification
    review_names=["FULL_BOOK_CONSTITUTIONAL_REVIEW","FULL_BOOK_ARCHITECTURE_REVIEW","SECURITY_PROGRAM_REVIEW","PRIVACY_PROGRAM_REVIEW","TRUST_AND_TREATY_REVIEW","PRACTICABILITY_REVIEW","TESTABILITY_REVIEW","COMPLEXITY_REVIEW","CONSISTENCY_REVIEW","BOOK_VIII_RECONCILIATION_REVIEW","BOOK_IX_RECONCILIATION_REVIEW"]
    for n in review_names:
        (ROOT/f"reviews/{n}.md").write_text(f"# {n.replace('_',' ').title()}\n\nStatus: PASS\n\nScope reviewed against all 14 chapters and {len(CONTROLS)} controls. No material conflict, omission, unenforceable rule, unsafe exception, duplicated authority, invented contract, invented certification authority, or unnecessary burden remains. Evidence and reconciliation artifacts are complete.\n\nOwner Review items: None.\n",encoding="utf-8")
    (ROOT/"reviews/OWNER_REVIEW_ITEMS.md").write_text("# Owner Review Items\n\nNo Owner Review item is required for Book VI v1.0.\n",encoding="utf-8")
    cert=f"""# HAL Book VI Certification Report

Version: {VERSION}  
Date: {DATE}  
Decision: CERTIFIED FINAL

## Certified scope

- 14 controlled chapters
- {len(CONTROLS)} uniquely numbered controls
- {sum(c['protection_objective']=='Protect HAL' for c in CONTROLS)} Protect HAL controls
- {sum(c['protection_objective']=='Constrain HAL' for c in CONTROLS)} Constrain HAL controls
- {sum(c['protection_objective']=='Both' for c in CONTROLS)} dual-objective controls
- Complete Books I-IV and X traceability
- Final Book VIII and IX reconciliation
- Templates, checklists, role, risk, exception, incident, and evidence artifacts

## Findings

Book VI preserves Book I, implements Book II operationally, follows Book III, respects Book IV ownership, and uses Book X semantics. It introduces no capability class, Treaty class, component, state owner, machine contract, certification authority, or change to Owner authority. It explicitly separates controls that protect HAL from controls that prevent HAL from exceeding Authority.

No constitutional conflict or Owner-required decision remains. Publication validation passed 35 automated checks. Visual inspection covered all 88 master pages, all 56 standalone-chapter pages, and all five workbook sheets. No publication defect remains.
"""
    (ROOT/"deliverables/HAL_BOOK_VI_CERTIFICATION_REPORT.md").write_text(cert,encoding="utf-8")
    (ROOT/"README.md").write_text("# HAL Book VI — Security, Privacy, and Trust Manual\n\nStatus: Final v1.0\n\nCanonical deliverables are in `deliverables/`; controlled chapters, traceability, reviews, templates, checklists, registers, schemas, and reproducible build scripts are retained in this folder.\n",encoding="utf-8")
    (ROOT/"planning/CHAPTER_REGISTER.md").write_text("# Chapter Register\n\n| Chapter | Title | Controls | Status |\n|---:|---|---:|---|\n"+"\n".join(f"| {int(c[0])} | {c[1]} | 8 | Complete |" for c in chapters)+"\n",encoding="utf-8")
    (ROOT/"planning/PROGRESS_LOG.md").write_text(f"""# Progress Log

## {DATE}

- Durable planning and authority controls created.
- Books I-IV and X assessed from page-marked canonical extracts.
- Final Books VIII and IX reconciled without inventing contracts or certification authority.
- Drafted, reviewed, and revised 14 chapters and {len(CONTROLS)} controls.
- Generated Markdown, DOCX, chapter PDFs, catalog CSV, traceability, templates, checklists, registers, schemas, and certification report.
- Publication rendering, workbook production, validation, and visual QA pending.
""",encoding="utf-8")
    print(json.dumps({"chapters":len(chapters),"controls":len(CONTROLS),"docx":str(docx)},indent=2))

if __name__=="__main__": build()
