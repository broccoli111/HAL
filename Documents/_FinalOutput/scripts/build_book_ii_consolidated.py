from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOCUMENTS = Path(__file__).resolve().parents[2]
BOOK = DOCUMENTS / "Book II"
OUTPUT = DOCUMENTS / "_FinalOutput"
EDITION_LABEL = "PROVISIONAL_2026-08-09"
DOC_OUT = OUTPUT / "DOC" / f"HAL_BOOK_II_ARCHITECTURE_SPECIFICATION_{EDITION_LABEL}.docx"
MD_OUT = OUTPUT / "Markdown" / f"HAL_BOOK_II_ARCHITECTURE_SPECIFICATION_{EDITION_LABEL}.md"


def edition(chapter: int) -> str:
    if chapter <= 5:
        return "v0.2"
    if chapter <= 20:
        return "v0.3"
    return "v0.2"


def find_one(folder: Path, pattern: str) -> Path:
    matches = sorted(folder.glob(pattern))
    if len(matches) != 1:
        raise RuntimeError(f"Expected one match for {pattern}, found {matches}")
    return matches[0]


chapter_docs = []
chapter_markdown = []
chapter_titles = []
for number in range(1, 36):
    version = edition(number)
    chapter_docs.append(
        find_one(
            BOOK / "DOC",
            f"HAL_Book_II_Chapter_{number:02d}_*_{version}.docx",
        )
    )
    chapter_markdown.append(
        find_one(
            BOOK / "markdown",
            f"HAL_Book_II_Chapter_{number:02d}_*_{version}.md",
        )
    )
    title_match = re.search(
        rf"Chapter_{number:02d}_(.+)_{re.escape(version)}\.md$",
        chapter_markdown[-1].name,
    )
    chapter_titles.append(title_match.group(1).replace("_", " "))


front_markdown = [
    "# HAL Book II — The Architecture Specification",
    "",
    "**Status:** Provisional working edition — not independently recertified  ",
    "**Scope:** Chapters 1–35  ",
    "**Compilation date:** 2026-08-09  ",
    "**Constitutional authority:** Book I — The Constitution",
    "",
    "Book I remains the supreme authority. This provisional edition contains the current Markdown working chapter editions, including the runtime-sovereignty clarification. It is published under the time-bounded Engineering Exception 0010 and does not replace the certified 2026-07-27 baseline or issue a certification claim.",
    "",
    "## Authoritative editions",
    "",
    "- Chapters 1–5: v0.2",
    "- Chapters 6–20: v0.3",
    "- Chapters 21–35: v0.2",
    "",
    "## Contents",
    "",
]
front_markdown.extend(
    f"{number}. {title} — {edition(number)}"
    for number, title in enumerate(chapter_titles, 1)
)
front_markdown.extend(["", "---", ""])
parts = ["\n".join(front_markdown)]
parts.extend(path.read_text(encoding="utf-8").strip() for path in chapter_markdown)
MD_OUT.write_text("\n\n---\n\n".join(parts) + "\n", encoding="utf-8")


document = Document(chapter_docs[0])
body = document._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

section = document.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("HAL Book II\nThe Architecture Specification")
run.bold = True
run.font.size = Pt(26)

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run(
    "Provisional working edition — not independently recertified\n"
    "Chapters 1–35\n"
    "Compiled 9 August 2026"
)

authority = document.add_paragraph()
authority.alignment = WD_ALIGN_PARAGRAPH.CENTER
authority_run = authority.add_run("Governed by Book I — The Constitution")
authority_run.bold = True

document.add_paragraph()
document.add_heading("Document Control", level=1)
control = document.add_table(rows=0, cols=2)
control.style = "Table Grid"
for key, value in [
    ("Document", "BOOK_II_ARCHITECTURE_SPECIFICATION"),
    ("Status", "Provisional working edition — not independently recertified"),
    ("Scope", "Thirty-five chapters"),
    ("Constitutional authority", "Book I Constitution v1.0"),
    ("Compilation date", "9 August 2026"),
    ("Independent review", "Pending; not waived by this publication"),
    ("Engineering exception", "0010; expires 16 August 2026"),
]:
    cells = control.add_row().cells
    cells[0].text = key
    cells[1].text = value

document.add_paragraph(
    "This edition includes the current Markdown working edition for each numbered "
    "chapter. It is a provisional publication for test-only documentation continuity. "
    "It must not be presented as independently recertified or as authorization for "
    "production runtime integration. Earlier certified editions and audit working papers "
    "remain outside this publication bundle."
)
statement = document.add_paragraph()
statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
statement_run = statement.add_run(
    "BOOK I REMAINS THE SUPREME AUTHORITY — THIS EDITION IS NOT RECERTIFIED"
)
statement_run.bold = True

document.add_page_break()
document.add_heading("Contents", level=1)
for number, title in enumerate(chapter_titles, 1):
    document.add_paragraph(
        f"{number:02d}  {title}  {edition(number)}",
        style="List Number",
    )
document.add_page_break()

for index, source in enumerate(chapter_docs):
    if index:
        document.add_page_break()
    source_document = Document(source)
    for child in source_document._element.body:
        if child.tag != qn("w:sectPr"):
            body.insert(len(body) - 1, deepcopy(child))

header = section.header.paragraphs[0]
header.text = "HAL  •  BOOK II  •  ARCHITECTURE SPECIFICATION"
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.clear()
footer.add_run("Provisional working edition — not independently recertified  •  ")
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
footer._p.append(field)

document.save(DOC_OUT)
print(
    {
        "chapters": len(chapter_docs),
        "markdown": str(MD_OUT),
        "docx": str(DOC_OUT),
    }
)
