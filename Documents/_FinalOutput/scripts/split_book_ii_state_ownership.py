#!/usr/bin/env python3
"""Give every Book II Chapter 1 §6 state domain one authoritative owner."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import _Cell


ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "DOC/HAL_BOOK_2_ARCHITECTURE_SPECIFICATION.docx"

ROWS = [
    ("Identity", "Identity Service", "Session caches, UI identity claims, identity projections"),
    ("Delegation", "Authority Service", "Delegation caches, effective-scope projections, UI delegation views"),
    ("Authentication", "Identity Service", "Session assertions, authentication caches, freshness projections"),
    ("Policy", "Policy System", "Signed evaluation bundles, local policy evaluators"),
    ("Exception", "Policy System", "Exception views, expiry queues, compliance projections"),
    ("Approval", "Policy System", "Approval views, workflow queues, notification projections"),
    ("Experience", "Experience Ledger", "Experience indexes, summaries, retention projections"),
    ("Audit", "Audit Service", "Forensic indexes, audit summaries, investigation projections"),
    ("Knowledge", "Knowledge Service", "Embeddings, caches, search indexes"),
    ("Pattern", "Knowledge Service", "Pattern indexes, confidence projections, retrieval caches"),
    ("Intent", "Intent Manager", "Dashboards, work queues, and intent projections"),
    ("Plan", "Planner", "Plan views, execution projections, and scheduling representations"),
    ("Transaction", "Transaction Coordinator", "Provider-specific task representations and transaction projections"),
    ("Outcome", "Outcome Service", "Outcome dashboards, summaries, and evaluation projections"),
    ("Configuration", "Configuration Plane", "Node-local verified configuration bundles, configuration projections"),
    ("Secret reference", "Secrets Service", "Short-lived credentials, secret-reference caches, rotation projections"),
    ("Node observation", "Node Registry", "Health views, capacity projections, scheduling inputs"),
    ("Provider observation", "Provider Registry", "Provider health views, benchmark projections, policy-fitness views"),
]


def set_cell_text(cell: _Cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.style = "Normal"


document = Document(PATH)
table = next(
    table
    for table in document.tables
    if [cell.text.strip() for cell in table.rows[0].cells]
    == ["State domain", "Authoritative owner", "Derived / replaceable forms"]
)

template = deepcopy(table.rows[1]._tr)
for row in list(table.rows[1:]):
    row._tr.getparent().remove(row._tr)

for domain, owner, derived in ROWS:
    clone = deepcopy(template)
    table._tbl.append(clone)
    cells = [_Cell(tc, table) for tc in clone.tc_lst]
    assert len(cells) == 3
    set_cell_text(cells[0], domain)
    set_cell_text(cells[1], owner)
    set_cell_text(cells[2], derived)

actual = [
    tuple(cell.text.strip() for cell in row.cells[:2])
    for row in table.rows[1:]
]
assert actual == [(domain, owner) for domain, owner, _ in ROWS]
assert len({domain for domain, _ in actual}) == len(actual)
assert all("/" not in domain and "," not in domain for domain, _ in actual)
assert all("/" not in owner and "," not in owner for _, owner in actual)

temporary = PATH.with_suffix(".updated.docx")
document.save(temporary)
temporary.replace(PATH)
print({"file": str(PATH), "state_domains": len(ROWS), "status": "updated"})
