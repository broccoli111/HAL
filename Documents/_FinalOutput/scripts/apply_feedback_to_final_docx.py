#!/usr/bin/env python3
"""Apply the 2026-07-28 feedback only to compiled DOCX files in _FinalOutput."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "DOC"


def all_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    text = paragraph.text.replace(old, new)
    nodes = paragraph._p.xpath(".//w:t")
    if nodes:
        nodes[0].text = text
        for node in nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(text)
    return True


def replace_everywhere(document, old: str, new: str) -> int:
    count = 0
    for paragraph in all_paragraphs(document):
        count += int(replace_in_paragraph(paragraph, old, new))
    return count


def save_atomic(document, path: Path) -> None:
    temporary = path.with_suffix(".updated.docx")
    document.save(temporary)
    temporary.replace(path)


def edit_book_vi() -> None:
    path = DOC_DIR / "HAL_BOOK_6_SECURITY_PRIVACY_AND_TRUST_MANUAL.docx"
    document = Document(path)
    assert replace_everywhere(document, "Articles I-XIV", "Articles I-XII") == 2
    assert replace_everywhere(document, "Articles XI-XIV", "Articles XI-XII") == 2
    save_atomic(document, path)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.style = "Normal"


def edit_book_ii() -> None:
    path = DOC_DIR / "HAL_BOOK_2_ARCHITECTURE_SPECIFICATION.docx"
    document = Document(path)
    target = None
    for table in document.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == "Intent, plan, transaction, outcome":
                target = row
                break
        if target is not None:
            break
    if target is None:
        existing_domains = {
            row.cells[0].text.strip()
            for table in document.tables
            for row in table.rows
            if row.cells
        }
        assert {"Intent", "Plan", "Transaction", "Outcome"} <= existing_domains
        return

    replacements = [
        (
            "Intent",
            "Intent Manager",
            "Dashboards, work queues, and intent projections",
        ),
        (
            "Plan",
            "Planner",
            "Plan views, execution projections, and scheduling representations",
        ),
        (
            "Transaction",
            "Transaction Coordinator",
            "Provider-specific task representations and transaction projections",
        ),
        (
            "Outcome",
            "Outcome Service",
            "Outcome dashboards, summaries, and evaluation projections",
        ),
    ]
    for domain, owner, derived in replacements:
        clone = deepcopy(target._tr)
        target._tr.addprevious(clone)
        cells = list(clone.tc_lst)
        assert len(cells) == 3
        from docx.table import _Cell

        set_cell_text(_Cell(cells[0], target._parent), domain)
        set_cell_text(_Cell(cells[1], target._parent), owner)
        set_cell_text(_Cell(cells[2], target._parent), derived)
    target._tr.getparent().remove(target._tr)
    save_atomic(document, path)


def edit_book_iv() -> None:
    path = DOC_DIR / "HAL_BOOK_4_COMPONENT_SPECIFICATIONS.docx"
    document = Document(path)
    assert replace_everywhere(document, "Forget Memory", "Restrict Memory Association") >= 2
    assert replace_everywhere(document, "Memory Forgotten", "Memory Association Restricted") >= 2

    command_base = (
        "Command `Restrict Memory Association` MUST carry identity, authority context "
        "where applicable, schema version, correlation, causation, time/freshness, "
        "provenance, and explicit success or denial semantics."
    )
    command_full = (
        command_base
        + " It MUST use Book II's protected-deletion and tombstone model: it may "
        "remove or limit accessibility and derived associations, but MUST NOT erase, "
        "rewrite, or make unauditable historical experience or immutable Experience "
        "Ledger records."
    )
    assert replace_everywhere(document, command_base, command_full) == 1

    event_base = (
        "Event `Memory Association Restricted` MUST carry identity, authority context "
        "where applicable, schema version, correlation, causation, time/freshness, "
        "provenance, and explicit success or denial semantics."
    )
    event_full = (
        event_base
        + " It records the governed restriction or tombstone without asserting "
        "deletion of historical experience or immutable archive evidence."
    )
    assert replace_everywhere(document, event_base, event_full) == 1

    save_atomic(document, path)


def replace_later_control_block(
    document,
    chapter_heading: str,
    first_control_heading: str,
    stop_heading: str,
    cross_reference: str,
) -> None:
    paragraphs = document.paragraphs
    chapter_index = next(
        index for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip() == chapter_heading
    )
    start_index = next(
        index for index in range(chapter_index + 1, len(paragraphs))
        if paragraphs[index].text.strip() == first_control_heading
    )
    stop_index = next(
        index for index in range(start_index + 1, len(paragraphs))
        if paragraphs[index].text.strip() == stop_heading
    )
    first = paragraphs[start_index]
    first.text = cross_reference
    first.style = "Normal"
    for paragraph in paragraphs[start_index + 1 : stop_index]:
        paragraph._element.getparent().remove(paragraph._element)


def move_control_definition_to_earlier_chapter(
    document,
    earlier_chapter: str,
    later_chapter: str,
) -> None:
    """Move a control block back when a TOC match caused the earlier block to be replaced."""

    def actual_heading_index(text: str) -> int:
        return next(
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip() == text
            and paragraph.style.name.lower() in {"title", "heading 1"}
        )

    def block_after(chapter_index: int):
        paragraphs = document.paragraphs
        normative_index = next(
            index
            for index in range(chapter_index + 1, len(paragraphs))
            if paragraphs[index].text.strip() == "Normative controls"
        )
        required_index = next(
            index
            for index in range(normative_index + 1, len(paragraphs))
            if paragraphs[index].text.strip() == "Required practices"
        )
        return paragraphs[normative_index + 1 : required_index], paragraphs[required_index]

    earlier_block, earlier_required = block_after(actual_heading_index(earlier_chapter))
    later_block, later_required = block_after(actual_heading_index(later_chapter))
    assert len(earlier_block) == 1
    assert len(later_block) > 1

    cross_reference = deepcopy(earlier_block[0]._p)
    control_definition = [deepcopy(paragraph._p) for paragraph in later_block]

    for paragraph in earlier_block:
        paragraph._element.getparent().remove(paragraph._element)
    for element in control_definition:
        earlier_required._p.addprevious(element)

    for paragraph in later_block:
        paragraph._element.getparent().remove(paragraph._element)
    later_required._p.addprevious(cross_reference)


def repair_book_ix_control_locations() -> None:
    path = DOC_DIR / "HAL_BOOK_9_INTERFACE_AND_PROTOCOL_REFERENCE.docx"
    document = Document(path)
    move_control_definition_to_earlier_chapter(
        document,
        "1. Purpose, Scope, Authority, and Conformance",
        "2. Contract Taxonomy and Catalog",
    )
    move_control_definition_to_earlier_chapter(
        document,
        "14. OpenAPI, AsyncAPI, Protocol Buffers, and JSON Schema Profiles",
        "15. Conformance, Compatibility Testing, and Certification Evidence",
    )
    save_atomic(document, path)


def edit_book_ix() -> None:
    path = DOC_DIR / "HAL_BOOK_9_INTERFACE_AND_PROTOCOL_REFERENCE.docx"
    document = Document(path)

    replace_later_control_block(
        document,
        "2. Contract Taxonomy and Catalog",
        "IX-GOV-001 — Authority hierarchy",
        "Required practices",
        (
            "This chapter applies controls IX-GOV-001 through IX-GOV-003 as defined "
            "once in Chapter 1. The identifiers are cross-references here and are not "
            "reissued or independently waivable."
        ),
    )
    replace_later_control_block(
        document,
        "15. Conformance, Compatibility Testing, and Certification Evidence",
        "IX-CNF-001 — Schema validation",
        "Required practices",
        (
            "This chapter applies controls IX-CNF-001 through IX-CNF-004 as defined "
            "once in Chapter 14. The identifiers are cross-references here and are not "
            "reissued or independently waivable."
        ),
    )

    assert replace_everywhere(document, "Forget Memory", "Restrict Memory Association") == 1
    assert replace_everywhere(document, "/hal/v1/cmp-09/forget-memory", "/hal/v1/cmp-09/restrict-memory-association") == 1
    assert replace_everywhere(document, "Memory Forgotten", "Memory Association Restricted") == 1
    assert replace_everywhere(document, "hal.cmp-09.memory.forgotten.v1", "hal.cmp-09.memory.association-restricted.v1") == 1

    base = (
        "All CMP-09 contracts require `HAL-Internal-Authority-v1`, the HAL envelope, "
        "structured errors, semantic versioning, bounded limits, and conformance "
        "evidence. Their payload semantics remain exactly those stated by Book IV."
    )
    extended = (
        base
        + " Restrict Memory Association and Memory Association Restricted MUST "
        "implement Book II's protected-deletion and tombstone model: they restrict "
        "accessibility or derived associations while preserving historical experience, "
        "immutable transition evidence, and Experience Ledger custody."
    )
    assert replace_everywhere(document, base, extended) == 1
    save_atomic(document, path)


def main() -> None:
    edit_book_ii()
    edit_book_iv()
    edit_book_vi()
    edit_book_ix()
    print("Updated compiled DOCX editions for Books II, IV, VI, and IX.")


if __name__ == "__main__":
    main()
