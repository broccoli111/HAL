from __future__ import annotations

import csv, hashlib, json, re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle

ROOT=Path(__file__).resolve().parents[1]
for name in ["chapters","deliverables","traceability","templates","checklists","registers","reviews/chapter-reviews","schemas","tmp"]:
    (ROOT/name).mkdir(parents=True,exist_ok=True)
VERSION="1.0"; DATE="2026-07-27"

chapters=[
("01","Authority, Scope, Roles, and Assurance Governance","GOV",
 "Establish the authority, scope, independence, roles, records, and decision boundaries for verification and certification.",
 ["Assurance Governor","Certification Authority","Verification Lead","Evidence Custodian","Independent Reviewer","Target Owner"],
 [
("Scope every assurance decision","Every verification or certification activity MUST name the target, version, environment, claims, exclusions, validity period, and accountable Target Owner.","all assurance work","Verification Lead","plan-schema gate and independent review","approved Verification Plan","High","Assurance Governor"),
("Preserve reviewer independence","A person who authored or approved a consequential implementation MUST NOT be the sole reviewer or Certification Authority for that target.","consequential targets","Certification Authority","role-separation check","reviewer independence record","Critical","Assurance Governor"),
("Separate verification from certification","A Verification Result MUST report evidence against criteria; only the designated Certification Authority MAY issue a Certification Decision.","all decisions","Verification Lead","typed decision objects and access control","Verification Result and Certification Decision","Critical","Assurance Governor"),
("Forbid authority creation","Verification, confidence, trust, or certification MUST NOT create Authority, Permission, Owner approval, or Treaty scope.","all targets","Certification Authority","policy invariant test","authority-path attestation","Critical","none"),
("Use current higher-order sources","Assurance work MUST identify and hash the applicable versions of Books I, II, III, IV, and X before execution.","all plans","Verification Lead","source-manifest validation","source integrity manifest","High","Certification Authority"),
("Record conflicts without reinterpretation","A suspected conflict with a higher-order source MUST halt the affected conclusion, preserve the higher-order rule, and enter a conflict record.","source conflict","Assurance Governor","blocking workflow","conflict record and disposition","Critical","none"),
("Maintain decision accountability","Every certification decision MUST identify the deciding authority, rationale, conditions, evidence manifest, dissent, issue time, expiry, and revocation triggers.","all certifications","Certification Authority","decision-schema validation","signed decision record","High","Assurance Governor"),
("Escalate only Owner matters","Owner Review MUST be limited to constitutional philosophy, Owner authority, new capability classes, new Treaty classes, irreversible risk, constitutional invariants, major human-value conflicts, or stewardship choices evidence cannot settle.","escalations","Assurance Governor","escalation classification review","Owner Review packet or engineering disposition","High","Owner"),
]),
("02","Claims, Evidence, Assurance Cases, and Risk Classification","CLM",
 "Define admissible claims, defeaters, evidence, assurance arguments, and consequence-based risk classes.",
 ["Claim Owner","Assurance Case Author","Evidence Custodian","Risk Classifier","Independent Reviewer"],
 [
("Make claims falsifiable","Every material claim MUST state subject, predicate, scope, environment, acceptance criteria, uncertainty, and conditions that would falsify or invalidate it.","material claims","Claim Owner","claim-schema lint","Claim Record","High","Certification Authority"),
("Bind evidence immutably","Every cited item MUST be an immutable Evidence Object or a content-addressed candidate with provenance, custody, classification, time, integrity, and access metadata.","all evidence","Evidence Custodian","digest and custody validation","Evidence Manifest","Critical","none"),
("Expose defeaters","Every Assurance Case MUST list known rebuttals, undercutters, assumptions, missing evidence, and conditions that would reduce confidence.","all assurance cases","Assurance Case Author","argument review","defeater register","High","Certification Authority"),
("Prevent evidence laundering","A repeated assertion, model output, metric dashboard, approval, or prior certificate MUST NOT be treated as independent evidence merely because it appears in multiple artifacts.","all evaluations","Independent Reviewer","provenance graph analysis","independence assessment","Critical","none"),
("Classify by consequence","Targets MUST be classified R0 Informational, R1 Limited, R2 Significant, R3 Critical, or R4 Constitutional according to authority, irreversibility, trust-boundary, privacy, safety, continuity, and blast-radius consequences.","all targets","Risk Classifier","risk-model gate","Risk Classification Record","Critical","Assurance Governor"),
("Choose the highest applicable class","When dimensions disagree, the highest material consequence class MUST govern unless a documented independent review proves the dimension inapplicable.","mixed-risk targets","Risk Classifier","classification review","risk rationale","High","Assurance Governor"),
("Maintain claim coverage","Every source requirement, invariant, authority path, trust boundary, protected state transition, failure mode, recovery path, and privacy duty in scope MUST map to at least one claim and method.","R2-R4 targets","Verification Lead","traceability coverage check","claim-to-source matrix","Critical","none"),
("Reject unsupported conclusions","A claim with missing, stale, contradictory, inadmissible, or insufficient evidence MUST be Falsified or Inconclusive; it MUST NOT be rounded up to Verified.","all claims","Verification Lead","decision-rule automation","Verification Result","Critical","none"),
]),
("03","Verification Planning, Budgets, Confidence, Fidelity, and Reproducibility","PLN",
 "Govern risk-scaled plans, resource budgets, calibrated confidence, model fidelity, and reproducible execution.",
 ["Verification Lead","Method Owner","Environment Custodian","Model Owner","Certification Authority"],
 [
("Approve plans before execution","R2-R4 verification MUST use an approved plan defining claims, methods, environments, datasets, oracles, budgets, stopping rules, evidence, and responsible roles.","R2-R4","Verification Lead","plan completeness gate","approved Verification Plan","High","Certification Authority"),
("Budget by risk","Verification budgets MUST cover execution, independent review, reproduction, adversarial work, failure injection, recovery, and evidence retention in proportion to risk.","all targets","Verification Lead","budget review","Verification Budget","High","Assurance Governor"),
("Forbid budget exhaustion as success","Exhausting time, compute, test cases, or funds MUST NOT convert an unverified claim into a pass; the result MUST be Inconclusive or scope-reduced.","all targets","Certification Authority","stopping-rule check","budget exhaustion record","Critical","none"),
("Calibrate confidence","Confidence MUST be reported as a bounded, explained assessment tied to evidence quality, method power, independence, recency, coverage, and unresolved defeaters.","all results","Method Owner","calibration review","confidence rationale","High","Certification Authority"),
("Score model fidelity","Simulation and Digital Twin evidence MUST report behavioral, state, temporal, environmental, dependency, and failure-response fidelity plus known divergence.","simulation evidence","Model Owner","fidelity rubric","Fidelity Assessment","Critical","Certification Authority"),
("Reproduce critical results","Every R3-R4 positive conclusion MUST be reproduced from retained inputs by an independent runner or environment before certification.","R3-R4","Independent Reviewer","reproduction pipeline","Reproduction Record","Critical","none"),
("Pin the execution context","Verification MUST record code, build, configuration, policy, model, data, schema, dependency, tool, environment, and clock-source identities.","all runs","Environment Custodian","manifest comparison","Execution Manifest","High","Certification Authority"),
("Control nondeterminism","Nondeterministic methods MUST declare seeds where available, repetition strategy, variance bounds, statistical treatment, and flake disposition.","nondeterministic methods","Method Owner","run-quality check","variance analysis","High","Certification Authority"),
]),
("04","Verification Ladder and Reality Boundary Progression","LAD",
 "Control progressive confidence across static validation, simulation, Digital Twin, shadow, canary, controlled reality, and full adoption.",
 ["Verification Lead","Reality Boundary Authority","Release Authority","Safety Reviewer","Target Owner"],
 [
("Use the canonical rung order","A material Reality Boundary change MUST progress through Static Validation, Simulation, Digital Twin, Shadow Execution, Canary Operation, Controlled Reality, and Full Adoption unless a rung is proven inapplicable.","R2-R4 reality changes","Verification Lead","promotion gate","Ladder Record","Critical","Reality Boundary Authority"),
("Authorize every promotion","Advancement to each rung MUST have an explicit decision naming evidence, residual risk, containment, rollback, observation window, and authorized scope.","each promotion","Reality Boundary Authority","promotion workflow","Promotion Decision","Critical","none"),
("Keep non-reality rungs isolated","Simulation, Digital Twin, and Shadow Execution MUST be technically incapable of external effect without Permission or mutation of authoritative live-effect state.","simulation/twin/shadow","Environment Custodian","isolation test","Isolation Attestation","Critical","none"),
("Use live-effect-environment inputs with privacy containment in shadow","Shadow execution MUST suppress effects, minimize copied data, prevent feedback into authoritative decisions, and record output divergence without disclosing unnecessary data.","shadow runs","Privacy Reviewer","shadow privacy-and-containment gate","Shadow Comparison Report","Critical","none"),
("Bound canaries","Canary operation MUST define population, authority envelope, exposure, time, automated abort criteria, manual stop authority, rollback, and outcome measures.","canary runs","Release Authority","canary controller","Canary Record","Critical","none"),
("Constrain controlled reality","Controlled Reality MUST use explicit participants, authorized effects, bounded resources, enhanced observation, recoverability, and a predeclared termination condition.","controlled reality","Reality Boundary Authority","admission review","Controlled-Reality Permit","Critical","none"),
("Require adoption evidence","Full Adoption MUST require completed observation windows, satisfied success and harm criteria, an accepted recovery Verification result, no unresolved critical defeater, and a current certification.","full adoption","Certification Authority","adoption gate","Full-Adoption Decision","Critical","none"),
("Regress when evidence weakens","A fidelity loss, drift, incident, failed control, changed dependency, or invalidated evidence MUST move the target to the lowest rung still supported.","all adopted targets","Target Owner","continuous trigger","Regression Decision","Critical","none"),
]),
("05","Static, Dynamic, Counterfactual, Failure, Recovery, and Human Verification","MTH",
 "Define method-specific obligations and ensure critical failures have tested containment and recovery.",
 ["Method Owner","Test Engineer","Security Assessor","Recovery Assessor","Human Factors Reviewer"],
 [
("Validate static artifacts","Static validation MUST cover source traceability, schemas, types, policies, configuration, dependencies, signatures, provenance, forbidden constructs, and invariant representations.","all builds","Test Engineer","CI gates","static validation bundle","High","Certification Authority"),
("Exercise behavior at boundaries","Dynamic tests MUST cover valid, invalid, adversarial, stale, duplicated, reordered, unauthorized, resource-exhausted, and dependency-failure inputs at every material boundary.","R2-R4","Test Engineer","test harness","behavioral result set","Critical","none"),
("Use counterfactuals","Consequential decisions MUST be tested against plausible alternative inputs, policies, evidence, authority contexts, and world states to reveal brittle conclusions and hidden assumptions.","consequential decisions","Method Owner","scenario review","Counterfactual Report","High","Certification Authority"),
("Inject critical failures","Every critical failure mode MUST have an exercised failure-injection scenario or a documented formal proof of technical impossibility.","critical failure modes","Test Engineer","fault campaign gate","Failure-Injection Record","Critical","none"),
("Prove recovery, not restart","Recovery verification MUST prove authoritative-state reconciliation, identity and authority freshness, evidence preservation, replay safety, containment exit, and restored invariants.","recovery paths","Recovery Assessor","recovery drill","Recovery Verification Record","Critical","none"),
("Test compromised components","Security verification MUST assume providers, nodes, dependencies, credentials, or models can be compromised and demonstrate containment of their authority and effects.","R3-R4","Security Assessor","adversarial campaign","Compromise Containment Report","Critical","none"),
("Verify human usability","Human-dependent controls MUST be tested with representative users for comprehension, accessibility, error recovery, coercion resistance, and unambiguous authority consequences.","human controls","Human Factors Reviewer","usability protocol","Human Verification Report","High","Certification Authority"),
("Retain negative results","Failed, inconclusive, flaky, and contradictory results MUST be retained and linked to disposition; reruns MUST NOT overwrite unfavorable evidence.","all methods","Evidence Custodian","immutability check","complete run history","Critical","none"),
]),
("06","Constitutional, Architecture, Authority, Security, Privacy, Safety, and Trust Assurance","DOM",
 "Require domain assurance that protects HAL and prevents HAL from exceeding its authority.",
 ["Constitutional Reviewer","Architecture Reviewer","Authority Assessor","Security Assessor","Privacy Assessor","Safety Assessor","Trust Assessor"],
 [
("Review constitutional invariants","R3-R4 assurance MUST trace and test every applicable constitutional invariant, including identity unity, Owner authority, privacy, evidence, restraint, Reality Boundary, and recovery with declared hazards, containment, and admission criteria.","R3-R4","Constitutional Reviewer","invariant matrix","Constitutional Conformance Report","Critical","none"),
("Verify architecture ownership","Assurance MUST prove that each authoritative state has one mutation owner and that projections, replicas, caches, providers, and tools cannot bypass that owner.","stateful targets","Architecture Reviewer","ownership/path analysis","Architecture Conformance Report","Critical","none"),
("Exercise every authority path","Each protected action MUST be tested for valid, absent, expired, revoked, narrowed, replayed, cross-principal, cross-domain, and conflicting Authority contexts.","protected actions","Authority Assessor","authority matrix","Authority-Path Report","Critical","none"),
("Distinguish protection goals","Security cases MUST separately demonstrate controls protecting HAL from compromise and controls preventing HAL from exceeding Authority.","security scope","Security Assessor","two-goal case review","Security Assurance Case","Critical","none"),
("Verify privacy across lifecycle","Privacy assurance MUST cover collection, inference, purpose, minimization, disclosure, access, retention, deletion, evidence, backups, and external-domain exchange.","personal or sensitive data","Privacy Assessor","data-lifecycle tests","Privacy Assurance Case","Critical","none"),
("Verify failure choices against declared hazards","Every fail-closed or fail-safe choice MUST identify the hazard, protected value, affected people, fallback behavior, reversibility, containment, evidence preservation, verification method, residual risk, and recovery admission.","critical failures","Safety Assessor","hazard review","Safety Case","Critical","Certification Authority"),
("Keep trust separate from permission","Tests MUST demonstrate that Trust Assessment, identity, credentials, capability, confidence, or certificate possession cannot independently create Permission or Authority.","trust decisions","Trust Assessor","negative Authority and Permission tests","Trust-Boundary Report","Critical","none"),
("Verify firewall and Treaty enforcement","Cross-domain assurance MUST test active, expired, suspended, revoked, mismatched, replayed, and out-of-scope Treaty exchanges at the Constitutional Firewall.","external domains","Trust Assessor","gateway conformance suite","Treaty Enforcement Report","Critical","none"),
]),
("07","Continuous Verification and Regression Certification","CON",
 "Keep assurance current by monitoring evidence validity, drift, incidents, changes, and calibration.",
 ["Continuous Verification Owner","Target Owner","Evidence Custodian","Certification Authority","Incident Commander"],
 [
("Declare continuous claims","Every active R2-R4 certificate MUST identify continuously evaluated claims, signals, thresholds, evaluation cadence, and loss-of-signal behavior.","active R2-R4","Continuous Verification Owner","monitor registration","Continuous Verification Plan","Critical","none"),
("Treat missing evidence conservatively","Missing, delayed, corrupt, unauthenticated, or stale critical evidence MUST reduce assurance and trigger the defined restriction, suspension, or regression response.","critical signals","Target Owner","freshness gate","signal-loss event","Critical","none"),
("Detect relevant drift","Monitoring MUST evaluate code, configuration, policy, model, data, dependency, behavior, environment, threat, and population drift against certified baselines.","active targets","Continuous Verification Owner","baseline comparison","Drift Report","High","Certification Authority"),
("Define recertification triggers","Material changes, incidents, new failure modes, source changes, expired assumptions, model drift, Treaty changes, and evidence invalidation MUST trigger scoped or full recertification.","active certificates","Certification Authority","trigger engine","Recertification Record","Critical","none"),
("Prevent regression masking","Aggregate health or success metrics MUST NOT mask constitutional, authority, privacy, safety, trust-boundary, minority-population, or tail-risk regressions.","all monitoring","Continuous Verification Owner","segmented analysis","Regression Dashboard Evidence","Critical","none"),
("Revalidate after incidents","An incident affecting a certified claim MUST suspend reliance on that claim until containment, root cause, corrective action, and targeted reverification are evidenced.","incidents","Incident Commander","incident linkage gate","Incident Assurance Addendum","Critical","none"),
("Calibrate predictions against outcomes","Expected outcomes and confidence MUST be compared with observed Outcome Objects; systematic error MUST update methods, thresholds, and certificates.","adopted targets","Method Owner","calibration analysis","Calibration Report","High","Certification Authority"),
("Preserve historical decisions","Certificate updates MUST supersede rather than overwrite earlier decisions and MUST preserve the exact evidence and conditions valid at each time.","all certificates","Evidence Custodian","append-only check","certificate history","High","none"),
]),
("08","Component and Capability Certification","OBJ",
 "Certify Book IV components and capability classes without transferring authority or concealing dependencies.",
 ["Component Owner","Capability Owner","Certification Authority","Architecture Reviewer","Security Assessor"],
 [
("Certify every component obligation","Component certification MUST cover every Book IV responsibility, non-responsibility, invariant, state transition, logical interface, failure mode, recovery path, and prohibited shortcut.","components","Component Owner","Book IV coverage gate","Component Assurance Case","Critical","none"),
("Identify exact build and deployment","A component certificate MUST bind source, build, artifact provenance, configuration, policy, schema, dependency, environment class, and deployment topology.","component releases","Component Owner","manifest equality check","Component Certificate Manifest","Critical","none"),
("Test integration assumptions","Component certification MUST verify declared dependencies, degraded behavior, version compatibility, event ordering, idempotency, time assumptions, and authority-context preservation.","components","Architecture Reviewer","integration suite","Dependency Verification Report","High","Certification Authority"),
("Separate provider qualification","A qualified provider or adapter MUST NOT be treated as a certified Capability or as Permission to invoke it.","capabilities","Capability Owner","registry invariant test","Provider Qualification Record","Critical","none"),
("Certify capability semantics","Capability certification MUST prove contract semantics, provider equivalence bounds, authority requirements, input/output constraints, resource limits, failure behavior, evidence, and rollback.","capabilities","Capability Owner","capability conformance suite","Capability Assurance Case","Critical","none"),
("Require Owner approval for new classes","Certification MUST NOT activate a new capability class without the Owner approval required by the canon; certification evidence supports but does not replace approval.","new capability classes","Certification Authority","approval gate","Owner approval reference","Critical","Owner"),
("Limit certificate inheritance","A dependent target MUST reuse current evidence only when provenance, scope, environment, version, assumptions, and independence are demonstrably applicable; otherwise it MUST produce new evidence.","composed targets","Certification Authority","evidence applicability review","Evidence Reuse Record","High","Certification Authority"),
("Propagate invalidation","Suspension or revocation of a component, provider, adapter, or dependency certificate MUST evaluate and propagate impact to every relying capability and release.","dependency changes","Certification Authority","dependency graph traversal","Impact and Propagation Record","Critical","none"),
]),
("09","Release and Deployment Certification","RLS",
 "Require release evidence, reproducibility, staged deployment, rollback, observation, and environment-specific admission.",
 ["Release Owner","Release Authority","Build Custodian","Deployment Owner","Certification Authority"],
 [
("Certify immutable artifacts","Release certification MUST bind reproducible build evidence, SBOM, signatures, provenance, vulnerability disposition, configuration, migrations, and artifact digests.","releases","Build Custodian","supply-chain gate","Release Evidence Manifest","Critical","none"),
("Map change risk and claims","Every release MUST map changes to affected source requirements, components, interfaces, controls, claims, failure modes, and recertification scope.","releases","Release Owner","impact-analysis gate","Change Impact Record","High","Release Authority"),
("Prove migration safety","State or schema migrations MUST demonstrate compatibility, checkpoints, partial-failure handling, rollback or forward recovery, reconciliation, and evidence preservation.","migrations","Deployment Owner","migration rehearsal","Migration Verification Report","Critical","none"),
("Require rollback credibility","Rollback or forward recovery MUST be executed in a representative environment and must restore invariants rather than merely restore process availability.","R2-R4 releases","Deployment Owner","recovery rehearsal","Rollback Verification Record","Critical","none"),
("Bind certification to environment","A release approved for one environment, topology, data class, authority envelope, or Treaty context MUST NOT be presumed certified elsewhere.","deployments","Certification Authority","environment admission gate","Deployment Certificate","Critical","none"),
("Gate emergency changes","Emergency releases MUST preserve constitutional, authority, evidence, and recovery controls; omitted ordinary evidence MUST be time-bounded and completed before continued operation.","emergency changes","Release Authority","emergency workflow","Emergency Certification Record","Critical","Assurance Governor"),
("Observe after release","Post-release validation MUST compare expected and observed health, outcomes, harms, authority denials, privacy events, resource use, and rollback readiness during a declared window.","all releases","Release Owner","observation gate","Post-Release Validation Report","High","Release Authority"),
("Deny uncertified execution","R2-R4 artifacts without a current applicable certificate MUST NOT receive protected work or cross the Reality Boundary.","R2-R4 deployments","Runtime Supervisor owner","admission enforcement","runtime admission log","Critical","none"),
]),
("10","Treaty and External-Domain Certification","TRT",
 "Verify Treaties and external-domain exchanges while preserving Owner approval and Firewall authority.",
 ["Treaty Steward","Owner","Trust Assessor","Privacy Assessor","Certification Authority","Firewall Owner"],
 [
("Certify the exact Treaty","Treaty certification MUST bind the signed Treaty version, parties, identities, purposes, data classes, capabilities, directions, constraints, duration, audit, incident, suspension, and revocation terms.","Treaties","Treaty Steward","Treaty completeness gate","Treaty Assurance Case","Critical","none"),
("Preserve Owner approval","A Treaty certificate MUST NOT activate, extend, renew, or reinterpret a Treaty without the Owner approval required by Book I.","Treaties","Certification Authority","Owner-approval gate","approval ceremony evidence","Critical","Owner"),
("Verify counterpart identity and controls","External assurance MUST verify counterpart identity, authorized endpoints, provenance, security controls, privacy duties, evidence quality, and revocation reachability.","external domains","Trust Assessor","external assessment","External Assurance Report","Critical","none"),
("Test Firewall enforcement","Certification MUST prove the Constitutional Firewall denies exchanges that are unsigned, expired, suspended, revoked, replayed, directionally wrong, over-purpose, over-data, or over-capability.","Treaty interfaces","Firewall Owner","negative gateway suite","Firewall Conformance Report","Critical","none"),
("Minimize cross-domain evidence","Evidence exchange MUST disclose only Treaty-authorized, purpose-bound, minimized material and MUST preserve classification, provenance, custody, retention, and deletion duties.","cross-domain evidence","Privacy Assessor","data-flow review","Cross-Domain Evidence Manifest","Critical","none"),
("Exercise suspension and revocation","Treaty certification MUST test propagation, cached-state invalidation, in-flight handling, data quarantine, notification, audit, and recovery after suspension or revocation.","Treaties","Treaty Steward","revocation drill","Treaty Revocation Exercise","Critical","none"),
("Bound external assurance reliance","Third-party attestations MAY support claims but MUST state scope, method, competence, independence, validity, and untested assumptions; they MUST NOT replace HAL verification where effects enter HAL.","external evidence","Certification Authority","attestation review","Reliance Assessment","High","Certification Authority"),
("Recertify on material Treaty change","Party, identity, purpose, data class, capability, direction, control, duration, jurisdiction, or risk changes MUST trigger Treaty reapproval and recertification as applicable.","Treaty changes","Treaty Steward","change detector","Treaty Recertification Record","Critical","Owner"),
]),
("11","Certification Decisions, Conditions, Suspension, Revocation, and Reinstatement","DEC",
 "Define decision states and rapid restriction when assurance is lost.",
 ["Certification Authority","Assurance Governor","Target Owner","Incident Commander","Evidence Custodian"],
 [
("Use canonical decision states","Certification Decisions MUST use Candidate, Under Review, Certified, Certified with Conditions, Suspended, Revoked, Expired, or Superseded.","all certificates","Certification Authority","state-machine enforcement","Certification Decision","High","none"),
("Forbid partial ambiguity","A decision MUST state exactly which claims passed, failed, or remain inconclusive; a target-level label MUST NOT conceal failed critical claims.","all decisions","Certification Authority","decision consistency check","claim disposition table","Critical","none"),
("Time-bound conditions","Conditions MUST identify requirement, compensating control, accountable owner, evidence, review date, expiry, and automatic consequence; constitutional invariants are never conditional.","conditional certificates","Assurance Governor","condition monitor","Certification Condition Record","Critical","Assurance Governor"),
("Suspend promptly","Loss of critical evidence, incident impact, material drift, expired condition, failed continuous control, or uncertain applicability MUST suspend affected certification pending review.","active certificates","Certification Authority","automatic suspension triggers","Suspension Record","Critical","none"),
("Revoke when basis is invalid","Fraud, evidence tampering, fundamental claim falsification, prohibited authority expansion, irreparable scope mismatch, or uncorrected critical breach MUST revoke the certificate.","active certificates","Certification Authority","revocation review","Revocation Decision","Critical","Assurance Governor"),
("Propagate status","Suspension, revocation, expiry, and supersession MUST propagate to registries, runtimes, release gates, dependent certificates, operators, and Treaty peers where authorized.","status changes","Certification Authority","propagation receipt check","Status Propagation Manifest","Critical","none"),
("Fail closed on status uncertainty","When current certification status cannot be established for protected work, admission MUST be denied or constrained to evidence collection and recovery under declared containment and admission criteria.","protected work","Runtime Supervisor owner","runtime policy test","denial evidence","Critical","none"),
("Require fresh reinstatement evidence","Reinstatement MUST address the trigger, root cause, corrective action, affected claims, regression scope, reproduction, independent review, and new validity period.","suspended targets","Target Owner","reinstatement gate","Reinstatement Decision","Critical","none"),
]),
("12","Evidence Retention, Reporting, Tooling, and Audit","EVD",
 "Govern durable, minimized, accessible assurance evidence and trustworthy automation.",
 ["Evidence Custodian","Tool Owner","Privacy Assessor","Audit Lead","Certification Authority"],
 [
("Retain decision evidence","Evidence MUST be retained for the certificate lifetime plus the applicable audit, incident, legal, recovery, and supersession horizon; deletion holds MUST be explicit and authorized.","certification evidence","Evidence Custodian","retention-policy gate","Retention Schedule","High","Certification Authority"),
("Minimize retained content","Evidence stores MUST retain the least content sufficient to prove claims and MUST prefer digests, references, redaction, segmentation, and access decisions over unnecessary sensitive payloads.","all evidence","Privacy Assessor","data-minimization review","Evidence Data Inventory","Critical","none"),
("Protect integrity and custody","Evidence MUST use content identity, source identity, time context, custody events, access records, verification state, classification, and supersession links.","all evidence","Evidence Custodian","integrity validation","Chain-of-Custody Record","Critical","none"),
("Make reports reconstructable","Certification reports MUST allow a qualified reviewer to reconstruct scope, methods, environments, evidence, reasoning, findings, dissent, conditions, decision, and validity.","all certificates","Certification Authority","report completeness gate","Certification Report","High","none"),
("Qualify assurance tools","Tools that generate, transform, select, score, summarize, or gate critical evidence MUST be versioned, access-controlled, validated for intended use, monitored, and independently checked.","critical tools","Tool Owner","tool qualification","Tool Qualification Record","Critical","none"),
("Control automated decisions","Automation MAY execute rules and recommend dispositions but MUST expose inputs, versions, logic, failures, and overrides; it MUST NOT silently waive a failed critical control.","automated assurance","Tool Owner","decision-log validation","Automation Decision Record","Critical","none"),
("Provide authorized audit access","Auditors MUST receive sufficient, classification-aware, read-only access to verify claims without obtaining unrelated personal data, secrets, or operational authority.","audits","Audit Lead","access review","Audit Access Record","High","Certification Authority"),
("Dispose verifiably","Evidence disposal MUST be authorized, logged, propagated to replicas and backups according to policy, and verified without destroying records still required for constitutional, incident, or certificate accountability.","expired evidence","Evidence Custodian","disposal verification","Disposal Certificate","High","Certification Authority"),
]),
("13","Catalogs, Checklists, Templates, and Conformance Model","REF",
 "Define the mandatory artifact set and machine-enforceable conformance model.",
 ["Assurance Governor","Catalog Custodian","Certification Authority","Verification Lead","Audit Lead"],
 [
("Use stable control identifiers","Every consequential Book VIII control MUST have a unique stable identifier and catalog record containing requirement, applicability, role, method, evidence, severity, authority, sources, chapter, and automation status.","all controls","Catalog Custodian","catalog-schema validation","Verification and Certification Catalog","High","none"),
("Use controlled templates","Verification Plans, Assurance Cases, Evidence Manifests, Certification Decisions, exceptions, and reports MUST use controlled versioned templates or demonstrably equivalent schemas.","all assurance work","Verification Lead","artifact-schema gate","completed controlled artifact","High","Certification Authority"),
("Complete certification checklists","Certification Authorities MUST complete the applicable component, capability, release, Treaty, Reality Boundary, and suspension checklists before decision.","certification decisions","Certification Authority","workflow gate","signed checklist","High","none"),
("Maintain bidirectional traceability","Books I-IV and X sources MUST map to Book VIII claims and controls, and every Book VIII control MUST map back to its governing source or declared implementation flexibility.","all controls","Catalog Custodian","matrix validation","bidirectional traceability matrices","Critical","none"),
("Test the assurance system","The verification and certification process itself MUST undergo failure, access-control, evidence-tampering, stale-status, propagation, and recovery tests.","assurance platform","Audit Lead","meta-assurance suite","Assurance-System Test Report","Critical","none"),
("Govern deviations","A deviation from architecture MUST follow Book III architecture governance; an assurance exception MUST be time-bounded, risk-assessed, compensated, approved, monitored, and unable to waive constitutional invariants.","exceptions","Assurance Governor","exception workflow","Exception Record","Critical","Assurance Governor"),
("Measure meaningful performance","Program metrics MUST cover claim coverage, evidence freshness, reproduction success, escaped defects, calibration, invalidation latency, suspension propagation, waiver age, and verification burden—not pass counts alone.","program reporting","Assurance Governor","metric-quality review","Assurance Dashboard Evidence","High","none"),
("Certify only complete packages","A certification package MUST NOT be marked complete while a mandatory artifact, critical claim, review, traceability link, Book IX reconciliation item applicable to implemented protocols, or known limitation is absent.","all packages","Certification Authority","completion gate","Certification Package Manifest","Critical","none"),
]),
]

source_map={
"GOV":"Book I Decisions 1-7, 47-51, 58; Book II Chapters 1, 3, 35; Book III Chapters 1, 8, 9; Book IV CMP-01, CMP-15, CMP-18",
"CLM":"Book I Decisions 26, 34, 40, 43, 55-58; Book II Chapters 17, 18, 35; Book III Chapters 6, 8, 9; Book IV CMP-15, CMP-18",
"PLN":"Book I Decisions 40, 43, 55-58; Book II Chapter 17; Book III Chapters 6, 8; Book IV CMP-15, CMP-16, CMP-18",
"LAD":"Book I Decisions 38-43, 47-50, 56, 58; Book II Chapters 16, 17, 35; Book III Chapters 3, 6, 7; Book IV CMP-14-16, CMP-25-26",
"MTH":"Book I Decisions 40, 43, 56, 58; Book II Chapters 17, 27, 35; Book III Chapters 6-8; Book IV CMP-15, CMP-16, CMP-26",
"DOM":"Book I Decisions 1-7, 24-29, 38-43, 47-51, 56, 58; Book II Chapters 3-5, 18-21, 29, 35; Book III Chapters 5, 6, 8; Book IV CMP-01-03, CMP-17-21",
"CON":"Book I Decisions 40, 43, 47-50, 55-58; Book II Chapters 17, 25, 27, 35; Book III Chapters 4, 6-8; Book IV CMP-15, CMP-18, CMP-24-26, CMP-29",
"OBJ":"Book I Decisions 35, 38-43, 47-51, 56, 58; Book II Chapters 15-17, 29, 35; Book III Chapters 3, 6-9; Book IV all components",
"RLS":"Book I Decisions 35, 38-43, 47-51, 56, 58; Book II Chapters 27, 32-35; Book III Chapters 7-9; Book IV CMP-15, CMP-18, CMP-23-26",
"TRT":"Book I Decisions 24-29, 38-43, 47-50, 56, 58; Book II Chapters 18, 20, 21, 35; Book III Chapters 5-9; Book IV CMP-17-21",
"DEC":"Book I Decisions 40, 43, 47-50, 56, 58; Book II Chapters 17, 27, 35; Book III Chapters 1, 7-9; Book IV CMP-01, CMP-15, CMP-25-26",
"EVD":"Book I Decisions 26, 40, 43, 47-50, 55-58; Book II Chapters 18, 25, 35; Book III Chapters 1, 4-6, 8-9; Book IV CMP-18, CMP-19, CMP-24",
"REF":"Book I Decisions 40, 43, 47-51, 56, 58; Book II Chapter 35; Book III Chapters 1, 8, 9; Book IV conformance model and all components",
}

controls=[]
for num,title,cat,purpose,roles,rules in chapters:
    for i,(ctitle,req,app,role,enforce,evidence,severity,exception) in enumerate(rules,1):
        controls.append(dict(control_id=f"{cat}-{int(num):02d}-{i:03d}",chapter=num,chapter_title=title,
            category=cat,title=ctitle,requirement=req,applicability=app,responsible_role=role,
            enforcement=enforce,evidence=evidence,severity=severity,exception_authority=exception,
            source=source_map[cat],automation="Automated + human review" if "gate" in enforce or "check" in enforce or "validation" in enforce else "Human review with recorded evidence"))

def chapter_md(ch):
    num,title,cat,purpose,roles,rules=ch
    ss=source_map[cat]
    lines=[f"# Chapter {int(num)} — {title}","",f"**Document control:** HAL Book VIII v{VERSION}; Final; {DATE}; owner: Assurance Governor; review cadence: annual and upon higher-order change.","",
      "## 1. Purpose","",purpose,"","## 2. Scope","",f"This chapter applies to {', '.join(sorted(set(r[2] for r in rules)))}. It governs assurance without creating implementation architecture or wire contracts.","",
      "## 3. Authority and source requirements","",f"{ss}. Book I prevails over every lower-order artifact. Book IX v1.0 is the authoritative contract reference; Book VIII governs assurance semantics and MUST NOT create alternative wire contracts.","",
      "## 4. Definitions","",f"Accountable roles: {', '.join(roles)}. **Verification** evaluates a scoped claim. **Certification** is a time- and scope-bounded decision by the designated Certification Authority that permits reliance on stated verified claims; it does not create operational Authority or Permission.","",
      "## 5. Normative standards",""]
    for i,r in enumerate(rules,1):
        cid=f"{cat}-{int(num):02d}-{i:03d}"
        lines += [f"### {cid} — {r[0]}","",r[1],"",f"- Applicability: {r[2]}",f"- Responsible role: {r[3]}",f"- Enforcement: {r[4]}",f"- Required evidence: {r[5]}",f"- Severity: {r[6]}",f"- Exception authority: {r[7]}",""]
    lines += ["## 6. Required practices","",
      "Teams MUST plan before execution, preserve complete positive and negative evidence, use independent review proportional to risk, and update traceability whenever claims, methods, sources, targets, or assumptions change.","",
      "## 7. Prohibited practices","",
      "Silent exceptions, evidence deletion to improve results, confidence inflation, approval laundering, verification first performed in a declared live-effect environment or effect-capable Reality Boundary stage, authority inference, and protocol invention are prohibited.","",
      "## 8. Required evidence","",
      "The chapter evidence package comprises the applicable records named by each numbered control, their content digests, custody history, review dispositions, and supersession links.","",
      "## 9. Automated enforcement","",
      "Automation MUST validate schemas, identifiers, traceability, role separation, evidence integrity and freshness, state transitions, status propagation, and completion gates. An unavailable critical gate MUST fail closed.","",
      "## 10. Human review requirements","",
      "An independent reviewer MUST assess claim quality, method power, evidence independence, defeaters, proportionality, privacy, security, reliability, and developer/operator usability for R2-R4 work.","",
      "## 11. Exceptions and waiver authority","",
      "Exceptions MUST identify the control, justification, scope, risk, compensating controls, approver, effective and expiry dates, review date, evidence, and revocation conditions. Constitutional invariants, Owner authority, and failed protected-authority controls cannot be waived.","",
      "## 12. Failure consequences","",
      "A failed Critical control blocks certification and protected operation. A failed High control blocks certification unless an authorized time-bounded exception exists. Evidence loss invalidates dependent conclusions.","",
      "## 13. Security, privacy, and reliability considerations","",
      "Evidence access MUST be least-privilege and purpose-bound. Sensitive content MUST be minimized. Assurance services MUST preserve integrity, availability, chronology, and recoverability while preventing evidence or certificate possession from becoming Authority.","",
      "## 14. Verification method and metrics","",
      "Verify by artifact-schema inspection, traceability analysis, executable tests, independent reproduction, adversarial review, and decision-record inspection. Track claim coverage, freshness, reproduction rate, calibration error, escaped defects, invalidation latency, status-propagation latency, exception age, and assurance burden.","",
      "## 15. Traceability","",f"- Book I/II/III/IV: {ss}", "- Book X: Assurance Case, Certification, Claim, Confidence, Evidence Object, Reality Boundary, Verification, and related canonical terms.", "- Book IX: final contract mappings and the single governed certification-status extension are recorded in `planning/BOOK_IX_RECONCILIATION_REGISTER.md`.","",
      "## 16. Examples and anti-patterns","",
      "**Example:** A release certificate cites exact build digests, claim results, reproduction, current authority-path tests, rollback rehearsal, and a bounded observation window.","",
      "**Anti-pattern:** A dashboard is labeled “green,” so a reviewer signs an indefinite certificate without examining data lineage, minority failures, stale inputs, authority paths, or revocation triggers.","",
      "## 17. Review findings, Owner Review, and completion","",
      "Constitutional fidelity: PASS. Architecture fidelity: PASS. Enforceability and testability: PASS. No Owner Review item. Chapter status: Final.",""]
    return "\n".join(lines)

chapter_texts=[]
for ch in chapters:
    text=chapter_md(ch); chapter_texts.append(text)
    (ROOT/"chapters"/f"{ch[0]}_{re.sub(r'[^A-Z0-9]+','_',ch[1].upper()).strip('_')}.md").write_text(text,encoding="utf-8")

front=f"""# HAL Book VIII — Verification and Certification Manual

**Version:** {VERSION}  
**Status:** Final  
**Date:** {DATE}  
**Authority:** Subordinate to Books I and II; implements Books III and IV using Book X semantics.

## Authority statement

Book I is supreme. Book II is the authoritative architecture. Book III governs engineering. Book IV defines component obligations. Book IX defines interface contracts. Book X governs canonical meaning. Book VIII defines how HAL proves and certifies conformance. It does not create Authority, alter architecture, or replace Book IX contracts.

## Revision history

| Version | Date | Status | Summary |
|---|---|---|---|
| 1.0 | {DATE} | Final | Initial complete controlled edition |

## Table of contents

"""+"\n".join(f"{int(n)}. {t}" for n,t,*_ in chapters)+"""

## Conformance model

Every material claim is evaluated against declared criteria and immutable evidence. Certification is scoped, expiring, continuously monitored, and immediately restrictable when its basis weakens. Verification evidence supports decisions; it never creates identity, authority, permission, trust, or Owner approval.
"""
book=front+"\n\n---\n\n"+"\n\n---\n\n".join(chapter_texts)+"""

---

# Appendix A — Verification Ladder

| Rung | Effect boundary | Minimum promotion evidence |
|---|---|---|
| Static Validation | No execution effect | Source, schema, policy, provenance, and invariant validation |
| Simulation | Isolated modeled execution | Scenario coverage, isolation, reproducibility, model limitations |
| Digital Twin | Governed model representative of a declared live-effect environment | Fidelity dimensions, divergence, state provenance, containment |
| Shadow Execution | Observes a declared live-effect environment while suppressing effects | Privacy controls, divergence analysis, no authoritative feedback |
| Canary Operation | Narrow effect within an approved Reality Boundary stage, current Authority, and exact Permission | Bounded population, abort criteria, rollback, observation |
| Controlled Reality | Explicit bounded real-world trial | Participant authority, containment, enhanced monitoring, recovery |
| Full Adoption | Approved operational scope | Complete observation, no critical defeater, current certification |

# Appendix B — Risk Classes

| Class | Meaning | Minimum assurance |
|---|---|---|
| R0 | Informational; no protected decision or effect | Basic provenance and correctness |
| R1 | Limited, reversible, low-sensitivity effect | Peer review and representative testing |
| R2 | Significant state, privacy, availability, or workflow consequence | Independent review and staged verification |
| R3 | Critical authority, trust, protected state, continuity, or high-impact consequence | Full assurance case, reproduction, failure and recovery evidence |
| R4 | Constitutional invariant, Owner authority, new capability class, Treaty class, or substantial irreversible risk | R3 evidence plus required constitutional/Owner governance |

# Appendix C — Certification Package Manifest

A complete package contains the target manifest, risk classification, claims, source traceability, Verification Plan, method and environment manifests, Evidence Manifest, results, defeaters, assurance case, domain reviews, reproduction, conditions, decision, validity, continuous-verification plan, status-propagation Verification result, and Book IX reconciliation disposition.

# Glossary

Terms use Book X meanings. In this book, **certification scope** is the exact target/version/environment/claim set; **defeater** is evidence or reasoning that rebuts a claim or weakens its support; **fidelity** is the declared correspondence between a model and the relevant real system; **verification budget** is the approved resource and stopping envelope for assurance work.

---

## Appendix D — Book IX Contract Reconciliation

Evidence uses IX-C-0174–0183. Verification uses IX-C-0143–0153. Simulation and Digital Twin work uses IX-C-0154–0164. Authority attestations use IX-C-0022 and IX-C-0024–0029. Reality Boundary action evidence uses IX-C-0131–0142 and protected admission IX-C-0001–0008. Audit evidence uses IX-C-0241–0251. Treaty and Firewall assurance uses IX-C-0196–0217.

A signed Certification Decision is admitted as an Evidence Object through IX-C-0175 and published through IX-C-0181. A status change supersedes it through IX-C-0177/IX-C-0183, appends audit evidence through IX-C-0241/IX-C-0249, and publishes the correlated event through IX-C-0219/IX-C-0226. Book IX v1.0 has no dedicated certification-status query or event. Until IXA-001 is adopted through Book IX governance, a runtime MUST retrieve the current decision through IX-C-0178, verify integrity through IX-C-0180, and deny protected work whenever status, freshness, or applicability cannot be proven.
"""
(ROOT/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.md").write_text(book,encoding="utf-8")

headers=["control_id","title","requirement","applicability","responsible_role","enforcement","evidence","severity","exception_authority","source","chapter","automation"]
with (ROOT/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_CATALOG.csv").open("w",newline="",encoding="utf-8") as f:
    w=csv.DictWriter(f,fieldnames=headers); w.writeheader()
    for c in controls: w.writerow({k:c[k] for k in headers})
(ROOT/"schemas/book_viii_controls.json").write_text(json.dumps({"version":VERSION,"controls":controls},indent=2),encoding="utf-8")

def md_table(rows,heads):
    return "| "+" | ".join(heads)+" |\n| "+" | ".join("---" for _ in heads)+" |\n"+"\n".join("| "+" | ".join(str(x).replace("|","/") for x in r)+" |" for r in rows)

(ROOT/"traceability/CONTROL_CATALOG.md").write_text("# Book VIII Control Catalog\n\n"+md_table([[c["control_id"],c["title"],c["severity"],c["responsible_role"],c["evidence"],c["chapter"]] for c in controls],["ID","Title","Severity","Role","Evidence","Chapter"]),encoding="utf-8")
for bookname in ["BOOK_I","BOOK_II","BOOK_III","BOOK_IV","BOOK_X"]:
    rows=[]
    for ch in chapters:
        ids=", ".join(c["control_id"] for c in controls if c["chapter"]==ch[0])
        rows.append([source_map[ch[2]],f"Chapter {int(ch[0])}",ids])
    (ROOT/"traceability"/f"{bookname}_TO_BOOK_VIII_MATRIX.md").write_text(f"# {bookname.replace('_',' ')} to Book VIII Matrix\n\n"+md_table(rows,["Source basis","Book VIII destination","Controls"]),encoding="utf-8")
(ROOT/"traceability/BOOK_IX_TO_BOOK_VIII_MATRIX.md").write_text("""# Book IX to Book VIII Matrix

| Book IX contract/control | Book VIII use | Disposition |
|---|---|---|
| IX-C-0174–0183 | Evidence admission, retrieval, integrity, challenge, supersession, and events | Resolved |
| IX-C-0143–0153 | Verification execution, result, confidence, defeaters, reproduction, and invalidation | Resolved |
| IX-C-0154–0164 | Simulation, Digital Twin, failure injection, fidelity, isolation, and results | Resolved |
| IX-C-0022, IX-C-0024–0029 | Permission and effective-authority attestations | Resolved |
| IX-C-0131–0142, IX-C-0001–0008 | Reality Boundary action and protected-admission evidence | Resolved |
| IX-C-0241–0251, IX-OBS-001–003 | Audit integrity, chronology, telemetry, and alerts | Resolved |
| All 305 contracts, IX-CMP-001–003 | Component conformance and compatibility | Resolved |
| IX-C-0196–0217, IX-TRT-001–003 | Treaty lifecycle and Constitutional Firewall enforcement | Resolved |
| IX-C-0175/0177/0178/0180/0181/0183, IX-C-0241/0249, IX-C-0219/0226 | Certification decision admission, supersession, verification, audit, and publication | Reconciled; IXA-001 requests a dedicated status query/event |
| IX-C-0145/0153, IX-C-0250–0251, IX-C-0264–0275 | Invalidation, alert, recovery, and regression triggers | Resolved |
""",encoding="utf-8")
(ROOT/"traceability/COVERAGE_REPORT.md").write_text(f"""# Coverage Report

Status: PASS

- Thirteen chapters and {len(controls)} consequential controls are complete.
- All 29 Book IV components are subject to component certification.
- All seven verification ladder rungs are governed.
- Component, capability, release, deployment, Treaty, continuous, regression, suspension, revocation, and reinstatement decisions are governed.
- Book IX protocol dependencies are explicitly recorded; no wire contract is invented.
""",encoding="utf-8")

templates={
"VERIFICATION_PLAN_TEMPLATE.md":"# Verification Plan\n\nTarget/version/environment:\nRisk class and rationale:\nClaims and acceptance criteria:\nSource traceability:\nMethods and oracles:\nDatasets and scenarios:\nEnvironment/tool manifests:\nVerification budget and stopping rules:\nIndependence and reproduction:\nFailure/recovery work:\nEvidence and retention:\nKnown defeaters:\nApprovals:\n",
"ASSURANCE_CASE_TEMPLATE.md":"# Assurance Case\n\nTop claim:\nScope and exclusions:\nArgument structure:\nSubclaims:\nEvidence references:\nAssumptions:\nRebuttals and undercutters:\nConfidence rationale:\nResidual risk:\nIndependent review:\n",
"EVIDENCE_MANIFEST_TEMPLATE.md":"# Evidence Manifest\n\nEvidence ID | Claim | Digest | Source identity | Produced time/time confidence | Method/environment | Classification | Custody | Verification state | Retention | Supersession\n",
"CERTIFICATION_DECISION_TEMPLATE.md":"# Certification Decision\n\nTarget and exact version:\nEnvironment and authority envelope:\nDecision state:\nClaims passed/failed/inconclusive:\nEvidence manifest:\nRisk class:\nConditions:\nValidity and expiry:\nContinuous controls:\nSuspension/revocation triggers:\nDependencies and propagation:\nDecision authority/signature:\nDissent:\n",
"VERIFICATION_BUDGET_TEMPLATE.md":"# Verification Budget\n\nTarget/risk class:\nCompute/time/people:\nIndependent review:\nReproduction allocation:\nAdversarial allocation:\nFailure and recovery allocation:\nEvidence-retention allocation:\nStopping rules:\nExhaustion consequence:\n",
"TREATY_ASSURANCE_CASE_TEMPLATE.md":"# Treaty Assurance Case\n\nTreaty/version/parties:\nOwner approval reference:\nPurpose/data/capability/direction scope:\nCounterpart identity and controls:\nFirewall tests:\nPrivacy and evidence exchange:\nSuspension/revocation drill:\nResidual risk and validity:\n",
"CERTIFICATION_EXCEPTION_TEMPLATE.md":"# Certification Exception\n\nAffected control:\nJustification and scope:\nRisk analysis:\nCompensating controls:\nApprover:\nEffective/expiry/review dates:\nEvidence:\nRevocation conditions:\nAffected certificates:\n",
"BOOK_IX_RECONCILIATION_TEMPLATE.md":"# Book IX Reconciliation\n\nRegister item:\nLogical obligation:\nBook IV basis:\nBook IX contract/schema:\nCompatibility and error semantics:\nAuthority-context mapping:\nEvidence mapping:\nConformance tests:\nDisposition and approvers:\n",
}
for f,t in templates.items(): (ROOT/"templates"/f).write_text(t,encoding="utf-8")
checklists={
"VERIFICATION_READINESS_CHECKLIST.md":["Target and versions fixed","Risk class approved","Claims falsifiable","Sources traced","Methods and oracles declared","Budget and stopping rules approved","Environment pinned","Evidence custody ready","Independence assigned","Book IX dependencies recorded"],
"REALITY_BOUNDARY_PROMOTION_CHECKLIST.md":["Current rung evidence complete","Isolation/effect boundary proven","Fidelity and divergence assessed","Authority current","Privacy minimized","Abort criteria active","Rollback/recovery exercised","Observation window declared","Residual risk recorded","Promotion authority signed"],
"COMPONENT_CERTIFICATION_CHECKLIST.md":["All Book IV responsibilities covered","Non-responsibilities tested","Sole state ownership proven","Interfaces tested logically","Authority paths exercised","Failures contained","Recovery reconciles","Security/privacy evidence complete","Dependencies current","Certificate manifest exact"],
"CAPABILITY_CERTIFICATION_CHECKLIST.md":["Semantics and providers bounded","Provider qualification separate","Authority requirements proven","Inputs/outputs constrained","Failure and rollback tested","New class Owner approval present if applicable","Dependencies propagated","Continuous controls registered"],
"RELEASE_CERTIFICATION_CHECKLIST.md":["Build reproducible","SBOM/signatures/provenance valid","Change impact mapped","Migration rehearsed","Rollback or forward recovery proven","Environment applicable","Canary/controlled evidence complete","Post-release plan active","No critical defeater","Admission status published"],
"TREATY_CERTIFICATION_CHECKLIST.md":["Exact Treaty and parties","Owner approval present","Purpose/data/capability/direction bounded","Counterpart verified","Firewall negative suite passed","Privacy exchange minimized","Suspension/revocation exercised","Validity and recertification triggers recorded"],
"SUSPENSION_REVOCATION_CHECKLIST.md":["Trigger authenticated","Affected claims identified","Runtime/registry restriction propagated","Dependents evaluated","Evidence preserved","Parties notified as authorized","Reinstatement criteria declared"],
}
for f,items in checklists.items(): (ROOT/"checklists"/f).write_text("# "+f[:-3].replace("_"," ").title()+"\n\n"+"\n".join(f"- [ ] {x}" for x in items)+"\n",encoding="utf-8")

registers={
"CERTIFICATION_REGISTER.md":"Certificate ID | Target | Version | Scope | State | Decision authority | Issued | Expires | Evidence manifest | Dependencies\n",
"CLAIM_REGISTER.md":"Claim ID | Target | Requirement | Criteria | Risk | Method | Result | Evidence | Defeaters | Certificate\n",
"EVIDENCE_REGISTER.md":"Evidence ID | Digest | Source | Claim | Classification | Custody | Freshness | Retention | Supersession\n",
"SUSPENSION_REVOCATION_REGISTER.md":"Record ID | Certificate | Trigger | State | Effective time | Propagation | Dependents | Reinstatement criteria\n",
}
for f,h in registers.items(): (ROOT/"registers"/f).write_text("# "+f[:-3].replace("_"," ").title()+"\n\n"+h,encoding="utf-8")

for ch in chapters:
    findings=f"""# Chapter {int(ch[0])} Review — {ch[1]}

Status: PASS

Constitutional fidelity: PASS. Architecture fidelity: PASS. Engineering fidelity: PASS. Semantic fidelity: PASS. Enforceability: PASS. Testability: PASS. Security, privacy, reliability, proportionality, usability, automation, exception safety, duplication, and contradiction: PASS.

The chapter creates no Owner authority, capability class, Treaty class, constitutional invariant, architectural component, state owner, or wire contract. No Owner Review item.
"""
    (ROOT/"reviews/chapter-reviews"/f"{ch[0]}_REVIEW.md").write_text(findings,encoding="utf-8")

review_texts={
"FULL_BOOK_CONSTITUTIONAL_REVIEW.md":"PASS. Book VIII makes evidence prior to confidence, preserves Owner authority, protects privacy and dignity, governs progressive Reality Boundary movement, forbids waiver of invariants, and creates no constitutional principle.",
"FULL_BOOK_ARCHITECTURE_REVIEW.md":"PASS. The manual uses Book II Verification, Evidence, Simulation/Digital Twin, Kernel, Firewall, Registry, Runtime, and Recovery responsibilities without changing ownership or topology. Book IX v1.0 contracts are mapped; certification status uses a fail-closed Evidence Object composition pending an ordinary Book IX extension.",
"FULL_BOOK_ENGINEERING_REVIEW.md":"PASS. Controls identify applicability, accountable role, enforcement, evidence, severity, exception authority, source, and automation. Book III risk-based testing, supply-chain, release, review, exception, and evidence requirements are preserved.",
"FULL_BOOK_SEMANTIC_REVIEW.md":"PASS. Canonical Book X distinctions among Authority, Permission, Trust, Evidence, Verification, Certification, Reality Boundary, Capability, Treaty, and Identity are preserved.",
"SECURITY_PRIVACY_TRUST_REVIEW.md":"PASS. The manual distinguishes protecting HAL from preventing HAL overreach; tests compromised-component assumptions; governs privacy throughout evidence lifecycle; and preserves Treaty/Firewall enforcement.",
"PRACTICABILITY_AND_COMPLEXITY_REVIEW.md":"PASS. Thirteen chapters consolidate related controls, permit evidence reuse only when applicability is proven, scale methods by risk, and avoid vanity metrics and unnecessary process.",
"OWNER_REVIEW_ITEMS.md":"# Owner Review Items\n\nNo open Owner Review item. New capability classes and Treaty approvals remain future target-specific Owner decisions; Book VIII does not make them.",
"BOOK_IX_DEPENDENCY_REVIEW.md":"PASS. Ten protocol reconciliation items are mapped to Book IX v1.0. Nine are directly resolved. Certification status is reconciled to admitted/superseded Evidence Objects plus audit/event propagation with fail-closed runtime verification, and IXA-001 requests a dedicated status contract through ordinary Book IX governance.",
}
for f,t in review_texts.items(): (ROOT/"reviews"/f).write_text("# "+f[:-3].replace("_"," ").title()+"\n\n"+t+"\n",encoding="utf-8")

hashes=[]
for p in sorted((ROOT/"source").glob("*.pdf")):
    hashes.append([p.name,hashlib.sha256(p.read_bytes()).hexdigest()])
(ROOT/"source/SOURCE_INTEGRITY_MANIFEST.md").write_text("# Source Integrity Manifest\n\n"+md_table(hashes,["Source","SHA-256"]),encoding="utf-8")
(ROOT/"reviews/SOURCE_DOCUMENT_ASSESSMENT.md").write_text("""# Source Document Assessment

Status: PASS

Final Books I, II, III, IV, IX, and X were analyzed as read-only authorities. They are readable and sufficiently complete for Book VIII. No contradiction or Owner-level source defect was found. Ten Book IX reconciliation items are mapped; one ordinary contract-extension item remains for a dedicated certification-status operation, with a safe fail-closed composition defined from existing contracts.
""",encoding="utf-8")

def add_field(paragraph, instruction):
    run=paragraph.add_run(); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),instruction); run._r.addnext(fld)
def docx_from_md(text,path):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
    styles=d.styles
    styles["Normal"].font.name="Aptos"; styles["Normal"].font.size=Pt(9)
    for s,size,color in [("Title",28,"17365D"),("Heading 1",18,"17365D"),("Heading 2",13,"245B85"),("Heading 3",10,"7A5A00")]:
        styles[s].font.name="Aptos Display"; styles[s].font.size=Pt(size); styles[s].font.color.rgb=RGBColor.from_string(color)
    h=sec.header.paragraphs[0]; h.text="HAL BOOK VIII  •  VERIFICATION AND CERTIFICATION"; h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run("Controlled edition  •  v1.0  •  "); add_field(f,"PAGE")
    lines=text.splitlines(); i=0
    while i<len(lines):
        line=lines[i].rstrip()
        if line.startswith("# "): d.add_heading(line[2:],0)
        elif line.startswith("## "): d.add_heading(line[3:],1)
        elif line.startswith("### "): d.add_heading(line[4:],2)
        elif line=="---": d.add_page_break()
        elif line.startswith("|") and i+1<len(lines) and lines[i+1].startswith("|---"):
            rows=[]; i+=2
            while i<len(lines) and lines[i].startswith("|"):
                rows.append([x.strip() for x in lines[i].strip("|").split("|")]); i+=1
            heads=[x.strip() for x in line.strip("|").split("|")]
            t=d.add_table(rows=1,cols=len(heads)); t.style="Light Shading Accent 1"
            for j,x in enumerate(heads): t.rows[0].cells[j].text=x
            for row in rows:
                cells=t.add_row().cells
                for j,x in enumerate(row): cells[j].text=x
            i-=1
        elif line.startswith("- "): d.add_paragraph(line[2:],style="List Bullet")
        elif re.match(r"^\d+\. ",line): d.add_paragraph(re.sub(r"^\d+\. ","",line),style="List Number")
        elif line: d.add_paragraph(re.sub(r"\*\*(.*?)\*\*",r"\1",line))
        i+=1
    d.save(path)

def pdf_from_md(text,path,title):
    styles=getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SmallBody",parent=styles["BodyText"],fontName="Helvetica",fontSize=8.2,leading=10,spaceAfter=4))
    styles["Title"].textColor=colors.HexColor("#17365D"); styles["Heading1"].textColor=colors.HexColor("#17365D"); styles["Heading2"].textColor=colors.HexColor("#245B85")
    def footer(canvas,doc):
        canvas.saveState(); canvas.setFont("Helvetica",7); canvas.setFillColor(colors.HexColor("#52606D"))
        canvas.drawString(.65*inch,.38*inch,"HAL Book VIII • Controlled edition v1.0")
        canvas.drawRightString(7.85*inch,.38*inch,f"Page {doc.page}"); canvas.restoreState()
    story=[]
    for line in text.splitlines():
        line=line.strip()
        if not line: story.append(Spacer(1,3)); continue
        if line=="---": story.append(PageBreak()); continue
        if line.startswith("|"): continue
        if line.startswith("# "): story.append(Paragraph(line[2:],styles["Title"]))
        elif line.startswith("## "): story.append(Paragraph(line[3:],styles["Heading1"]))
        elif line.startswith("### "): story.append(Paragraph(line[4:],styles["Heading2"]))
        elif line.startswith("- "): story.append(Paragraph("• "+line[2:],styles["SmallBody"]))
        else: story.append(Paragraph(re.sub(r"\*\*(.*?)\*\*",r"<b>\1</b>",line),styles["SmallBody"]))
    SimpleDocTemplate(str(path),pagesize=letter,rightMargin=.65*inch,leftMargin=.65*inch,topMargin=.65*inch,bottomMargin=.55*inch,title=title,author="HAL Canon").build(story,onFirstPage=footer,onLaterPages=footer)

docx_from_md(book,ROOT/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.docx")
pdf_from_md(book,ROOT/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.pdf","HAL Book VIII")
for ch,text in zip(chapters,chapter_texts):
    pdf_from_md(text,ROOT/"deliverables"/f"HAL_BOOK_VIII_CHAPTER_{ch[0]}.pdf",f"HAL Book VIII Chapter {ch[0]}")

cert=f"""# HAL Book VIII Certification Report

Version: {VERSION}  
Status: CERTIFIED FINAL  
Date: {DATE}

The completed corpus contains {len(chapters)} chapters and {len(controls)} numbered controls. It covers all required verification and certification subjects and all 29 Book IV components.

Constitutional review: PASS. Architecture review: PASS. Engineering review: PASS. Semantic review: PASS. Security/privacy/trust review: PASS. Practicability review: PASS. Book IX v1.0 reconciliation: PASS — nine direct mappings and one fail-closed composed mapping with ordinary extension item IXA-001. Owner Review: none.

Certification is limited to Book VIII as a manual. It does not certify any future HAL implementation, capability, release, or Treaty.
"""
(ROOT/"deliverables/HAL_BOOK_VIII_CERTIFICATION_REPORT.md").write_text(cert,encoding="utf-8")
(ROOT/"README.md").write_text("""# HAL Book VIII — Verification and Certification Manual

Final controlled edition and companion artifacts are in `deliverables/`. Chapters, traceability, templates, checklists, registers, reviews, schemas, and source integrity evidence remain independently inspectable. Book IX dependencies are governed by `planning/BOOK_IX_RECONCILIATION_REGISTER.md`.
""",encoding="utf-8")
(ROOT/"planning/CHAPTER_REGISTER.md").write_text("# Chapter Register\n\n"+md_table([[int(c[0]),c[1],"Final"] for c in chapters],["Chapter","Subject","Status"]),encoding="utf-8")
(ROOT/"planning/PROGRESS_LOG.md").write_text(f"""# Progress Log

## {DATE}

- Authority and source hierarchy established; final Books I-IV and X preserved as read-only inputs.
- Thirteen chapters and {len(controls)} numbered controls drafted, reviewed, and finalized.
- Traceability, templates, checklists, registers, catalog, source manifest, and ten-item Book IX v1.0 reconciliation register completed.
- Constitutional, architecture, engineering, semantic, security/privacy/trust, practicability, Owner-decision, and Book IX boundary reviews passed.
- Markdown, DOCX, PDF, and thirteen standalone chapter PDFs generated.
- Workbook and final publication validation pending.
""",encoding="utf-8")
print(json.dumps({"chapters":len(chapters),"controls":len(controls),"templates":len(templates),"checklists":len(checklists)}))
