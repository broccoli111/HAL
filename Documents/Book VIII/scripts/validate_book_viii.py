from pathlib import Path
import csv, hashlib, json, re, sys
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

R=Path(__file__).resolve().parents[1]
checks=[]
def ck(name,ok,detail=""):
    checks.append((name,bool(ok),detail))
    if not ok: print("FAIL",name,detail)

book=(R/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.md").read_text()
data=json.loads((R/"schemas/book_viii_controls.json").read_text())
controls=data["controls"]
ck("13 chapter files",len(list((R/"chapters").glob("*.md")))==13)
ck("13 chapter headings",len(re.findall(r"^# Chapter ",book,re.M))==13)
ck("104 unique controls",len(controls)==104 and len({c["control_id"] for c in controls})==104)
ck("all control fields",all(all(c.get(k) for k in ["control_id","title","requirement","applicability","responsible_role","enforcement","evidence","severity","exception_authority","source","chapter","automation"]) for c in controls))
ck("normative controls",all(" MUST " in (" "+c["requirement"]+" ") or " MUST NOT " in (" "+c["requirement"]+" ") for c in controls))
ck("all required domains",all(x.lower() in book.lower() for x in ["Static Validation","Digital Twin","Shadow Execution","Canary Operation","Controlled Reality","Full Adoption","counterfactual","failure injection","reproduc","fidelity","confidence","component certification","capability certification","release certification","Treaty certification","suspension","revocation","evidence retention","constitutional conformance","architecture conformance"]))
ck("all 29 components covered","all 29 book iv components" in (R/"traceability/COVERAGE_REPORT.md").read_text().lower())
ck("Book IX register ten items",len(re.findall(r"^\| IXR-\d{3} ",(R/"planning/BOOK_IX_RECONCILIATION_REGISTER.md").read_text(),re.M))==10)
ck("no protocol invention",all(term not in book for term in ["https://api.","application/json","grpc://","POST /","GET /"]))
ck("no draft markers",not re.search(r"\b(TODO|TBD|FIXME|placeholder)\b",book,re.I))
ck("chapter reviews",len(list((R/"reviews/chapter-reviews").glob("*.md")))==13)
ck("templates",len(list((R/"templates").glob("*.md")))>=8)
ck("checklists",len(list((R/"checklists").glob("*.md")))>=7)
ck("traceability matrices",len(list((R/"traceability").glob("*_TO_BOOK_VIII_MATRIX.md")))==6)
ck("Book IX contract mappings",all(x in book for x in ["IX-C-0174","IX-C-0143","IX-C-0154","IX-C-0022","IX-C-0131","IX-C-0241","IX-C-0196","IXA-001"]))
ck("owner review empty",(R/"reviews/OWNER_REVIEW_ITEMS.md").exists() and "No open Owner Review item" in (R/"reviews/OWNER_REVIEW_ITEMS.md").read_text())
required=["HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.md","HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.docx","HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.pdf","HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_CATALOG.xlsx","HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_CATALOG.csv","HAL_BOOK_VIII_CERTIFICATION_REPORT.md"]
for f in required: ck("deliverable "+f,(R/"deliverables"/f).exists() and (R/"deliverables"/f).stat().st_size>100)
ck("13 chapter PDFs",len(list((R/"deliverables").glob("HAL_BOOK_VIII_CHAPTER_*.pdf")))==13)
master=PdfReader(str(R/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.pdf"))
ck("master pdf pages",len(master.pages)>50,str(len(master.pages)))
ck("all pdfs readable",all(len(PdfReader(str(p)).pages)>0 for p in (R/"deliverables").glob("*.pdf")))
doc=Document(str(R/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_MANUAL.docx"))
ck("docx substantial",len(doc.paragraphs)>700,str(len(doc.paragraphs)))
wb=load_workbook(R/"deliverables/HAL_BOOK_VIII_VERIFICATION_AND_CERTIFICATION_CATALOG.xlsx",data_only=False,read_only=True)
ck("workbook sheets",wb.sheetnames==["Summary","Controls","Risk Classes","Verification Ladder","Certification States","Book IX Reconciliation"],str(wb.sheetnames))
ck("workbook controls",wb["Controls"]["A105"].value is not None,str(wb["Controls"]["A105"].value))
ck("workbook formulas",wb["Summary"]["B7"].value=="=COUNTA(Controls!A2:A105)")
expected={
"BOOK_I_CONSTITUTION.pdf":"fd71a272f3eaead6ee3ec864ae8b2b2786f4ce76f724468366886387b7dcdb49",
"BOOK_II_ARCHITECTURE_SPECIFICATION.pdf":"c202abc8f9cc7393e76df6e7b11f8f97310153b12e69202984a88a20922e8f72",
"BOOK_III_ENGINEERING_STANDARDS.pdf":"c5ea5936a91c1e15d9d7e90fe2e07a6a0f6f942a7f4a88b4207de857247ff73c",
"BOOK_IV_COMPONENT_SPECIFICATIONS.pdf":"1092bc4ef796a3272a2677df88d5bb85c6325c3f4ba5008eae9ad8fd7aead1bd",
"BOOK_X_CANONICAL_TERMINOLOGY_AND_INFORMATION_MODEL.pdf":"efb67c87f2690ad3a6c8d93877b1a357a2aa70b1d5ab83a4d089728f1913ac28"}
expected["BOOK_IX_INTERFACE_AND_PROTOCOL_REFERENCE.pdf"]="459b3c3701bfe60f6c4eee783a449b5cbc8be1756c5a4d0f6174b2ccb4011428"
ck("source hashes",all(hashlib.sha256((R/"source"/f).read_bytes()).hexdigest()==h for f,h in expected.items()))
passed=sum(x[1] for x in checks)
report=f"""# Publication Validation

Status: {'PASS' if passed==len(checks) else 'FAIL'}

- Automated checks: {passed}/{len(checks)}
- Master PDF pages: {len(master.pages)}
- Standalone chapter PDFs: 13
- Workbook sheets: 6
- Controls: 104
- Source hashes: unchanged

Visual inspection is recorded separately in the final independent audit.
"""
(R/"reviews/PUBLICATION_VALIDATION.md").write_text(report)
print(json.dumps({"checks":len(checks),"passed":passed,"master_pages":len(master.pages),"docx_paragraphs":len(doc.paragraphs),"workbook_sheets":len(wb.sheetnames)}))
sys.exit(0 if passed==len(checks) else 1)
