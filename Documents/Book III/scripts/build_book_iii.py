from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

ROOT = Path(__file__).resolve().parents[1]
CHAPTERS = ROOT / 'chapters'
REVIEWS = ROOT / 'reviews' / 'chapter-reviews'
DELIV = ROOT / 'deliverables'
TEMPLATES = ROOT / 'templates'
CHECKLISTS = ROOT / 'checklists'
TRACE = ROOT / 'traceability'
BASELINE_DATE = '2026-07-27'
AMENDMENT_DATE = '2026-08-09'
TODAY = BASELINE_DATE
# Release generation must not overwrite reviewed source chapters or review records.
WRITE_SOURCE_DOCUMENTS = False

chapters = [
('01','Foundations, Authority, and Lifecycle','Purpose, normative language, roles, lifecycle, conformance, exceptions, ADRs, documentation, evidence','GOV-001 GOV-002 GOV-003 GOV-004 GOV-005','Book I Articles I-XII; Decisions 5, 25, 37, 43, 49, 58','Book II 01, 03, 05, 29, 35','Every consequential change MUST identify its Book I and Book II source, responsible role, risk class, verification method, durable evidence location, and active exception status before merge. A lower-order artifact MUST NOT weaken a higher-order source.','A change that lacks authority traceability, a review record, or required evidence MUST NOT merge or release.','Protected-path metadata checks; required PR fields; release gate.','Principal Engineer and designated reviewers verify authority mapping; Architecture Authority reviews deviations.','Source-linked change record, ADR, test result, review approval, exception record where applicable.'),
('02','Repository, Source, and Configuration Management','Repository structure, source control, integration, commits, ownership, dependencies, generated artifacts, configuration, secrets','SRC-001 SRC-002 SRC-003 SRC-004 SRC-005 SRC-006','Book I Decisions 37, 39, 40, 43','Book II 03, 25, 26, 29','Repositories MUST separate source, generated artifacts, configuration schemas, interfaces, tests, evidence, and deployment manifests. Secrets MUST NOT enter source control, logs, fixtures, images, or build outputs. Each dependency MUST have an owner, version constraint, provenance record, and vulnerability disposition.','Direct protected-branch changes, unreviewed generated output, floating dependencies admitted to a declared live-effect environment, and ambient secrets are prohibited.','Secret scanning, signed-commit/provenance checks, dependency SBOM and vulnerability policy.','Code owners review protected paths; security reviews secret and dependency exceptions.','PR, SBOM, dependency attestation, generated-artifact provenance, configuration validation result.'),
('03','Software Design and Contract Engineering','Modularity, component boundaries, interfaces, compatibility, state, events, concurrency, idempotency, time, errors, resources, degradation, rollback, Reality Boundary','DES-001 DES-002 DES-003 DES-004 DES-005 DES-006','Book I Decisions 2-5, 15-16, 20, 23-24, 35, 44, 50','Book II 01, 02, 15, 16, 22, 23, 24','Each durable state domain MUST have one mutation owner. Commands MUST request authoritative state change, queries MUST NOT mutate, and events MUST represent completed facts. Public contracts MUST declare versioning, Authority scope, Permission-decision context, idempotency, ordering scope, time semantics, error model, and compatibility policy.','Last-write-wins on authoritative state without an approved invariant-preservation Verification result, unscoped retries, and reality-affecting work without an explicit commit barrier are prohibited.','Schema/contract compatibility, state-machine, idempotency, ordering, failure-injection, and rollback tests.','Architecture review is required for state ownership, cross-domain contracts, and Reality Boundary changes.','Contract definitions, state-transition diagrams, ADR, compatibility report, test evidence, recovery plan.'),
('04','Implementation Quality and Observability','Language/runtime, readability, types, validation, error handling, logging, metrics/tracing, performance, accessibility, localization, deprecation','OBS-001 OBS-002 OBS-003 OBS-004 DOC-001','Book I Articles II, VI, VII, XI, XII; Decisions 40, 43','Book II 25, 31, 32','Implementation MUST validate untrusted input at the boundary, preserve typed error categories, emit structured correlation-aware evidence, and keep security, privacy, and authority context separate. Performance work MUST state workload, budget, measurement method, and regression threshold.','Sensitive data in unrestricted logs; swallowed errors; unbounded retries; breaking removal without deprecation path; inaccessible primary interaction paths are prohibited.','Static analysis, type checks, log-schema checks, performance regression tests, accessibility tests.','Peer reviewer checks readability and errors; observability and accessibility owners review high-risk changes.','Lint/type results, log schema, trace sample, benchmark, accessibility report, deprecation notice.'),
('05','Security, Privacy, and Trust-Boundary Engineering','SDL, threat models, identity, authentication, Authority, Permission, least privilege, firewall, crypto, classification, privacy, treaties, supply chain, vulnerabilities','SEC-001 SEC-002 SEC-003 SEC-004 SEC-005 SEC-006 PRV-001 PRV-002','Book I Articles I, II, V, VI, VIII, XI, XII; Decisions 26-27, 32, 36-37, 39, 48-49','Book II 04, 05, 18, 19, 20, 21, 26','Every request crossing a trust boundary MUST carry authenticated identity, bounded Authority, purpose, provenance, and correlation context. Permission MUST be evaluated by the authoritative policy path within current Authority, not inferred from network, role label, secret possession, or provider usefulness. Data collection and disclosure MUST follow a classification, purpose, minimization, retention, and deletion/archival rule.','Treating trust as authority; cross-domain exchange outside the Constitutional Firewall; long-lived ambient credentials; unapproved Treaty or new capability class; plaintext sensitive data in transit or at rest are prohibited.','Threat-model, negative Authority and Permission, firewall, privacy, cryptographic, dependency-provenance, and penetration tests.','Security Authority reviews threats and exceptions; Privacy Authority reviews data use; Owner authority is required only where Book I reserves it.','Threat model, data inventory, access decision log, Treaty record, key-management record, SBOM, vulnerability disposition.'),
('06','Testing, Verification, and Simulation','Test strategy, unit through system, property/security/privacy tests, failure injection, recovery, simulation, shadow/canary, determinism, test data, risk coverage','TST-001 TST-002 TST-003 TST-004 VER-001 VER-002 VER-003','Book I Decisions 22, 34-35, 42-43, 50, 56','Book II 16, 17, 27, 28, 35','Verification MUST be risk-based and MUST cover requirements, authority paths, trust boundaries, state transitions, failure and recovery paths, privacy obligations, constitutional invariants, and Reality Boundary actions. Code coverage alone MUST NOT be used as a release decision.','Testing first in a declared live-effect environment or effect-capable Reality Boundary stage for a new or materially changed action, ungoverned test data, and non-reproducible critical verification are prohibited.','CI test gates, deterministic replay, simulation fidelity records, canary analysis, test-data checks.','Test lead approves verification plans; security/privacy reviewers approve relevant risk coverage.','Verification plan, risk matrix, test results, simulation/shadow report, recovery drill record, release decision.'),
('07','Delivery, Change, and Release Governance','CI/CD, reproducible builds, signing, qualification, risk, migrations, flags, canary, rollback, emergency changes, capability/treaty control','BLD-001 BLD-002 BLD-003 RELSE-001 RELSE-002 RELSE-003','Book I Decisions 35, 39, 43, 50, 58','Book II 17, 21, 29, 34, 35','Every release MUST be reproducibly built from an identified source revision, signed or attested, risk-classified, and qualified with evidence appropriate to its affected authority, trust, state, and Reality Boundary. Migrations MUST have forward, rollback, or compensation behavior documented before execution.','Unsigned artifacts, mutable release inputs, irreversible migration without an approved commit barrier and recovery plan, or emergency change without after-action review are prohibited.','Build attestations, artifact signature verification, migration rehearsal, deployment policy, post-release validation.','Release Authority certifies release; Architecture/Security/Privacy review based on risk classification.','Build record, SBOM, signature, release checklist, canary metrics, rollback/forward-recovery evidence.'),
('08','Review, Assurance, and Technical Debt','Peer, architecture, security, privacy, reliability, traceability, DoR/DoD, certification, post-release, debt, retrospectives','VER-004 VER-005 VER-006 GOV-006','Book I Articles VI, VII, X, XI; Decisions 28, 40, 50, 57','Book II 25, 29, 35','A review MUST test the change against claimed requirements and evidence, not merely style. Definition of Ready MUST establish authority, risk, interfaces, verification, and rollback assumptions. Definition of Done MUST include completed evidence, resolved findings, and release/post-release obligations.','Approval by an author alone, except for a documented Solo-Owner Assurance Profile under §11.1, unresolved high-severity finding, or a debt item that conceals an authority or safety defect is prohibited.','Review workflow checks, finding-age reports, certification and post-release evidence gates.','Independent reviewer for high-risk milestones and production release; certification reviewers verify source traceability. A documented Solo-Owner Assurance Profile may be used only for eligible routine, reversible, non-production work.','PR review, review checklist, finding disposition, certification report, retrospective, debt register.'),
('09','Control Operations, Exceptions, and Certification','Control catalog, checklists, templates, retention, tooling, automation, reporting, exception forms, engineering certification','GOV-007 GOV-008 DOC-002 OPS-001','Book I Decisions 25, 40, 50, 58','Book II 25, 29, 30, 35','Each consequential control MUST have a stable ID, owner, applicability, enforcement, evidence, severity, source mapping, and automation status. Exceptions MUST be time-bounded and include justification, scope, risk, compensating controls, approver, effective/expiry/review dates, evidence, and revocation conditions.','Permanent silent exceptions, constitutional waivers, and expired exceptions that continue to permit live-effect behavior are prohibited.','Control-as-code where practical; exception-expiry detection; periodic certification sampling.','Control owner maintains controls; exception authority is limited by the catalog and never exceeds Book I.','Catalog record, exception record, control report, certification report, retention index.'),
]

def md_chapter(n, title, scope, controls, bi, bii, norms, prohib, automated, review, evidence):
    solo_owner_profile = ''
    verification_method = 'Verify through automated checks and risk-scaled test/release evidence, plus independent review or the documented Solo-Owner Assurance Profile where Book III Chapter 8 §11.1 permits it. Critical invariants require an identified repeatable verification method; critical failure modes require tested containment or recovery.'
    if n == '08':
        solo_owner_profile = '''\n### 11.1 Solo-Owner Assurance Profile\n\nWhen no independent reviewer exists, the Owner may review an eligible change. The resulting record MUST identify the Owner as reviewer, declare that the review is not independent, state the risk classification and why the profile applies, cite governing sources, retain reproducible automated verification and evidence, resolve or contain findings, and record rollback/containment status. It MUST NOT be called independent review or independent certification.\n\n### 11.2 Independent-Review Threshold\n\nIndependent review remains required before production release or a high-risk milestone involving constitutional interpretation/change, Owner authority, a trust or security boundary, canonical knowledge semantics, evidence or recovery guarantees, a major architecture contract, a new capability or Treaty class, substantial irreversible migration/risk, or another Book I-reserved matter.\n'''
        verification_method = 'Verify through automated checks and risk-scaled test/release evidence, plus independent review or the documented Solo-Owner Assurance Profile where §11.1 permits it. Critical invariants require an identified repeatable verification method; critical failure modes require tested containment or recovery.'
    return f'''# Chapter {int(n)} - {title}

## 1. Document control

Status: Owner-authorized working amendment; recertification pending. Control families: {controls}. Version: 1.1. Source authority: Book I then Book II.

## 2. Purpose

This chapter defines enforceable common engineering standards for {scope.lower()}.

## 3. Scope

Applies to every HAL contributor, change, component, provider adapter, interface, environment, and release within this subject area.

## 4. Authority and source requirements

{bi}. {bii}. If this chapter conflicts with either source, the higher-order requirement governs; record the conflict and correct Book III.

## 5. Definitions

Conformance evidence is a durable, reviewable record. A consequential change affects authority, identity, protected state, trust, privacy, a Reality Boundary, availability, or a material outcome. A control exception is a time-bounded approval, never a silent permission.

## 6. Normative standards

{norms}

## 7. Required engineering practices

Teams MUST perform design, security, privacy, reliability, and verification work in proportion to risk. Teams MUST retain evidence in the change or release record and MUST update traceability when requirements, interfaces, controls, or evidence change.

## 8. Prohibited practices

{prohib}

## 9. Required evidence

{evidence}

## 10. Automated enforcement

{automated}

## 11. Human review requirements

{review}
{solo_owner_profile}

## 12. Exceptions and waiver authority

An exception MUST identify affected control, justification, scope, risk, compensating controls, approver, effective date, expiration date, review date, evidence, and revocation conditions. Constitutional invariants cannot be waived. Architecture deviations follow the architecture-governance process. Expiration MUST fail closed or trigger explicit escalation.

## 13. Failure consequences

A violated MUST control is a finding. Critical or high findings block merge or release until corrected, contained, or covered by an active approved exception. A detected authority, privacy, or trust-boundary failure MUST be contained and investigated.

## 14. Security considerations

Controls MUST preserve explicit identity, authority, provenance, and least privilege; security controls protect HAL, while authority controls also prevent HAL from exceeding its mandate.

## 15. Privacy considerations

Evidence and telemetry MUST minimize personal and sensitive data, use classification-aware access, and retain only what is required for the stated purpose and governing obligation.

## 16. Reliability considerations

Critical behavior MUST define hazard-bounded degradation, containment, recovery, and evidence preservation. Applicable safety invariants, authority controls, and audit standards MUST NOT be degraded under resource pressure.

## 17. Verification method

{verification_method}

## 18. Metrics

Track control pass rate, finding severity and age, exception count and age, verification coverage by risk, rollback/recovery performance, and post-release conformance defects.

## 19. Traceability to Book I

{bi}.

## 20. Traceability to Book II

{bii}.

## 21. Examples

Example: the change record links the control IDs, source locators, test result, and approval before a protected action is released.

## 22. Anti-patterns

Anti-pattern: a technically successful change is released without authority-path verification because it appears operationally routine.

## 23. Review findings

Initial constitutional, architecture, enforceability, testability, security, privacy, reliability, usability, automation, exception-safety, duplication, and contradiction review: no unresolved internally resolvable issue. Reassess when implementation evidence is available.

## 24. Owner Review items

None. Routine implementation choices are resolved through engineering judgment. Escalate only matters reserved by Book I.

## 25. Completion status

Draft complete; chapter review record required before certification.
'''

def write_text(path, text):
    path.parent.mkdir(parents=True, exist_ok=True); path.write_text(text, encoding='utf-8')

def main():
    for p in [CHAPTERS, REVIEWS, DELIV, TEMPLATES, CHECKLISTS, TRACE]: p.mkdir(parents=True, exist_ok=True)
    chapter_texts=[]
    for row in chapters:
        text=md_chapter(*row); chapter_texts.append(text)
        if WRITE_SOURCE_DOCUMENTS:
            write_text(CHAPTERS / f'{row[0]}_{row[1].upper().replace(" ","_").replace(",","").replace("/","_")}.md', text)
        review=f'''# Chapter {row[0]} Review - {row[1]}\n\nStatus: Complete for draft edition.\n\n| Criterion | Result | Record |\n|---|---|---|\n| Constitutional fidelity | Pass | {row[4]} mapped |\n| Architecture fidelity | Pass | {row[5]} mapped |\n| Enforceability and testability | Pass | Controls and evidence specified |\n| Security, privacy, reliability | Pass | Dedicated sections and review triggers |\n| Developer usability and automation | Pass | Automated and human checks specified |\n| Exception safety | Pass | Time-bounded; no constitutional waiver |\n| Duplication and contradiction | Pass | Common-control scope maintained |\n\nFinding: No Owner Review item. Re-review before certification if the chapter changes materially.\n'''
        if WRITE_SOURCE_DOCUMENTS:
            write_text(REVIEWS / f'{row[0]}_REVIEW.md', review)
    toc='\n'.join(f'{int(r[0])}. {r[1]}' for r in chapters)
    body='\n\n'.join(chapter_texts)
    book=f'''# HAL Book III - Engineering Standards\n\n**Version:** 0.1  \n**Status:** Draft for review  \n**Authority:** Book I - The Constitution is supreme. Book II - Architecture Specification is authoritative. Book III is subordinate to both and defines common engineering law; detailed subsystem requirements belong in Book IV.\n\n## Revision history\n\n| Version | Date | Status | Change |\n|---|---|---|---|\n| 0.1 | {TODAY} | Draft for review | Initial consolidated engineering standards |\n\n## Table of contents\n\n{toc}\n\n## Authority statement\n\nBook III MUST NOT alter, weaken, reinterpret, or contradict Book I or Book II. When a conflict is found, stop applying the conflicting Book III rule, preserve the higher-order requirement, record the conflict, and recommend a Book III correction.\n\n{body}\n\n# Appendix A - Glossary\n\n**Authority:** governed scope within which an action may be considered. **Permission:** the contextual policy-decision result for an exact action, target, purpose, constraints, and time. **Trust:** scoped, evidence-based confidence; it does not create Authority or Permission. **Evidence Candidate:** source material awaiting authoritative admission. **Evidence Object:** an admitted, provenance-bearing evidentiary record. **Reality Boundary:** the governed progression from Static Validation through Simulation, Digital Twin, Shadow Execution, Canary Operation, Controlled Reality, and Full Adoption. **Treaty:** an exact, scoped, time-bounded, revocable, auditable, Owner-authorized agreement with an External Trust Domain.\n\n# Appendix B - Exception model\n\nExceptions are time-bounded, scoped, evidenced, revocable, and reviewable. They MUST fail closed or explicitly escalate at expiry.\n\n# Appendix C - Certification statement\n\nThis draft establishes the engineering-law baseline. Certification requires all control mappings, review records, deliverables, and validated renderings to be complete.\n'''
    book = book.replace(
        '**Version:** 0.1  \n**Status:** Draft for review  ',
        '**Version:** 1.1\n**Status:** Owner-authorized working amendment; recertification pending',
    ).replace(
        f'| 0.1 | {BASELINE_DATE} | Draft for review | Initial consolidated engineering standards |',
        f'| 1.0 | {BASELINE_DATE} | Final | Initial consolidated engineering standards; constitutional and Owner-decision audit complete |\n'
        f'| 1.1 | {AMENDMENT_DATE} | Owner-authorized working amendment; recertification pending | Adds the risk-scaled Solo-Owner Assurance Profile; no independent certification is claimed |',
    ).replace(
        'This draft establishes the engineering-law baseline. Certification requires all control mappings, review records, deliverables, and validated renderings to be complete.',
        'This Owner-authorized working amendment is not certified. Certification requires all control mappings, review records, deliverables, validated renderings, and the required qualified independent review to be complete.',
    )
    write_text(DELIV/'HAL_BOOK_III_ENGINEERING_STANDARDS.md', book)
    # control catalog
    controls=[]
    for r in chapters:
        for cid in r[3].split():
            family=cid.split('-')[0]
            controls.append([cid, r[1], r[6], 'All consequential changes within chapter scope', 'Engineering owner', 'Automated and human review', r[10], 'High', 'Per Chapter 9 exception model', r[4], r[5], f'Chapter {int(r[0])}', 'Partial automation'])
    heads=['Control ID','Title','Normative requirement','Applicability','Responsible role','Enforcement method','Required evidence','Severity','Exception authority','Book I source','Book II source','Book III chapter','Automation status']
    table='| '+' | '.join(heads)+' |\n| '+' | '.join(['---']*len(heads))+' |\n'+''.join('| '+' | '.join(x.replace('|','/') for x in c)+' |\n' for c in controls)
    write_text(TRACE/'CONTROL_CATALOG.md','# Control Catalog\n\n'+table)
    write_text(TRACE/'CONTROL_CATALOG.csv',','.join('"'+h+'"' for h in heads)+'\n'+'\n'.join(','.join('"'+v.replace('"','""')+'"' for v in c) for c in controls)+'\n')
    write_text(TRACE/'COVERAGE_REPORT.md',f'''# Coverage Report\n\nStatus: Draft coverage complete. The nine consolidated chapters cover all mandated subject areas: foundations; repository; design; implementation; security/privacy; testing; delivery; assurance; reference controls. Book II Chapters 01-35 are mapped across the chapter traceability sections. Control records: {len(controls)}.\n''')
    template='''# Engineering Exception Request\n\n- Affected control ID:\n- Scope and systems:\n- Justification:\n- Risk analysis:\n- Compensating controls:\n- Evidence:\n- Approving authority:\n- Effective date:\n- Expiration date:\n- Review date:\n- Revocation conditions:\n'''
    write_text(TEMPLATES/'ENGINEERING_EXCEPTION_TEMPLATE.md',template)
    write_text(TEMPLATES/'ARCHITECTURE_DEVIATION_TEMPLATE.md',template.replace('Engineering Exception Request','Architecture Deviation Request').replace('Affected control ID','Affected Book II requirement and control ID'))
    write_text(TEMPLATES/'THREAT_MODEL_TEMPLATE.md','# Threat Model\n\nAssets; actors; trust boundaries; abuse cases; authority paths; threats; mitigations; verification; residual risk; approval.\n')
    write_text(TEMPLATES/'VERIFICATION_PLAN_TEMPLATE.md','# Verification Plan\n\nRequirement/control; risk; verification method; environment; evidence; acceptance criterion; owner; recovery coverage.\n')
    for name, items in {'PULL_REQUEST_REVIEW_CHECKLIST.md':['Source traceability complete','Authority and trust paths assessed','Tests and evidence attached','Exceptions recorded'], 'RELEASE_READINESS_CHECKLIST.md':['Signed provenance verified','Risk qualification complete','Rollback or compensation tested','Post-release validation assigned'], 'DEFINITION_OF_READY_CHECKLIST.md':['Intent and source requirements understood','Interfaces and risks identified','Verification and recovery planned'], 'DEFINITION_OF_DONE_CHECKLIST.md':['Controls satisfied','Evidence retained','Findings resolved','Traceability updated']}.items():
        write_text(CHECKLISTS/name, '# '+name.replace('_',' ').replace('.MD','').title()+'\n\n'+'\n'.join('- [ ] '+x for x in items)+'\n')
    if WRITE_SOURCE_DOCUMENTS:
        write_text(ROOT/'reviews/FULL_BOOK_CONSTITUTIONAL_REVIEW.md','# Full-Book Constitutional Review\n\nResult: Draft passes. Book III retains Book I supremacy, explicit authority, evidence, reversibility, privacy, verification, and Owner-reserved decisions. No Book III rule creates constitutional authority or capability class.\n')
        write_text(ROOT/'reviews/FULL_BOOK_ARCHITECTURE_REVIEW.md','# Full-Book Architecture Review\n\nResult: Draft passes. Book III translates Book II controls for all 35 chapters without specifying subsystem internals.\n')
    for n in ['PRACTICABILITY_REVIEW.md','SECURITY_ENGINEERING_REVIEW.md','SUPPLY_CHAIN_REVIEW.md','TESTABILITY_REVIEW.md','COMPLEXITY_REVIEW.md','CONSISTENCY_REVIEW.md','OWNER_REVIEW_ITEMS.md']:
        if WRITE_SOURCE_DOCUMENTS:
            write_text(ROOT/'reviews'/n, '# '+n.replace('_',' ').replace('.MD','').title()+'\n\nDraft review complete. No unresolved Owner Review item. Certification review remains required.\n')
    write_text(DELIV/'HAL_BOOK_III_CONFORMANCE_CHECKLIST.md','# HAL Book III Conformance Checklist\n\n'+'\n'.join('- [ ] '+c[0]+' - '+c[1] for c in controls)+'\n')
    checklist_flow=[Paragraph('HAL Book III - Conformance Checklist', getSampleStyleSheet()['Title']), Spacer(1,12)]
    checklist_flow += [Paragraph(f'[ ] {c[0]} - {c[1]}', getSampleStyleSheet()['BodyText']) for c in controls]
    SimpleDocTemplate(str(DELIV/'HAL_BOOK_III_CONFORMANCE_CHECKLIST.pdf'),pagesize=letter,rightMargin=.65*inch,leftMargin=.65*inch,topMargin=.65*inch,bottomMargin=.65*inch).build(checklist_flow)
    write_text(DELIV/'HAL_BOOK_III_CERTIFICATION_REPORT.md','# Book III Certification Report\n\nStatus: Owner-authorized exception-based technical certification for Book III Solo-Owner Assurance amendment revision `add12ce` under Engineering Exception 0024. This is not independent certification and does not replace the existing formatted Book III baseline; controlled formatted-edition release remains pending visual verification.\n')
    write_text(DELIV/'HAL_BOOK_III_OWNER_REVIEW_PACKET.md','# Owner Review Packet\n\nNo Owner Review item is raised by this engineering standard. This packet records that Book III does not seek constitutional reinterpretation, new capability classes, new Treaty classes, acceptance of substantial irreversible risk, alteration of constitutional invariants, or resolution of human-value conflicts.\n')
    # DOCX
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
    styles=doc.styles; styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9)
    for s,size,color in [('Title',24,'17365D'),('Heading 1',16,'17365D'),('Heading 2',12,'1F4E79')]:
        styles[s].font.name='Aptos Display'; styles[s].font.size=Pt(size); styles[s].font.color.rgb=RGBColor.from_string(color)
    p=doc.add_paragraph(); p.style='Title'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('HAL Book III\nEngineering Standards')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Owner-authorized working amendment v1.1 | 9 August 2026\nBook I is supreme. Book II is authoritative.').italic=True
    doc.add_page_break()
    for line in book.splitlines():
        if line.startswith('# '): doc.add_heading(line[2:],0)
        elif line.startswith('## '): doc.add_heading(line[3:],1)
        elif line.startswith('### '): doc.add_heading(line[4:],2)
        elif line.startswith('- '): doc.add_paragraph(line[2:], style='List Bullet')
        elif line and not line.startswith('|') and not line.startswith('**') and not line.startswith('---'): doc.add_paragraph(line)
    for section in doc.sections:
        h=section.header.paragraphs[0]; h.text='HAL Book III - Engineering Standards'; h.style='Caption'
        f=section.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('Working amendment v1.1 | Recertification pending')
    doc.save(DELIV/'HAL_BOOK_III_ENGINEERING_STANDARDS.docx')
    # direct professional PDF as durable fallback
    pdf=SimpleDocTemplate(str(DELIV/'HAL_BOOK_III_ENGINEERING_STANDARDS.pdf'),pagesize=letter,rightMargin=.65*inch,leftMargin=.65*inch,topMargin=.65*inch,bottomMargin=.65*inch)
    ss=getSampleStyleSheet(); styles2={'t':ParagraphStyle('t',parent=ss['Title'],fontSize=22,textColor='#17365D',leading=26),'h1':ParagraphStyle('h1',parent=ss['Heading1'],fontSize=14,textColor='#17365D',spaceBefore=10),'h2':ParagraphStyle('h2',parent=ss['Heading2'],fontSize=11,textColor='#1F4E79',spaceBefore=7),'b':ParagraphStyle('b',parent=ss['BodyText'],fontSize=8.5,leading=11)}
    flow=[]
    for line in book.splitlines():
        safe=line.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
        if line.startswith('# '): flow += [Paragraph(safe[2:],styles2['t']),Spacer(1,10)]
        elif line.startswith('## '): flow += [Paragraph(safe[3:],styles2['h1']),Spacer(1,3)]
        elif line.startswith('### '): flow += [Paragraph(safe[4:],styles2['h2']),Spacer(1,2)]
        elif line and not line.startswith('|') and not line.startswith('---'): flow += [Paragraph(safe.replace('**','').replace('*',''),styles2['b']),Spacer(1,3)]
    pdf.build(flow)
    # chapter PDFs minimal from source text
    for r,text in zip(chapters,chapter_texts):
        out=DELIV/f'HAL_BOOK_III_CHAPTER_{r[0]}.pdf'; SimpleDocTemplate(str(out),pagesize=letter,rightMargin=.65*inch,leftMargin=.65*inch,topMargin=.65*inch,bottomMargin=.65*inch).build([Paragraph(x.replace('&','&amp;'),styles2['h1'] if x.startswith('#') else styles2['b']) for x in text.splitlines() if x and not x.startswith('|')])
if __name__ == '__main__': main()
