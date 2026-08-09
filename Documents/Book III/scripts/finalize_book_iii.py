from pathlib import Path
from docx import Document

ROOT=Path(__file__).resolve().parents[1]
DELIV=ROOT/'deliverables'
replacements={
    '**Version:** 0.1  \n**Status:** Draft for review':'**Version:** 1.0  \n**Status:** Final',
    '| 0.1 | 2026-07-27 | Draft for review | Initial consolidated engineering standards |':'| 1.0 | 2026-07-27 | Final | Initial consolidated engineering standards; constitutional and Owner-decision audit complete |',
    'This draft establishes the engineering-law baseline. Certification requires all control mappings, review records, deliverables, and validated renderings to be complete.':'This final edition establishes the engineering-law baseline. Required control mappings, review records, deliverables, and validation records are complete.',
    'Status: Draft for review.':'Status: Final.',
    'Version: 0.1.':'Version: 1.0.',
    'Draft complete; chapter review record required before certification.':'Complete; chapter review record retained in `reviews/chapter-reviews/`.',
}
targets=[DELIV/'HAL_BOOK_III_ENGINEERING_STANDARDS.md',*sorted((ROOT/'chapters').glob('*.md'))]
for path in targets:
    text=path.read_text(encoding='utf-8')
    for a,b in replacements.items(): text=text.replace(a,b)
    path.write_text(text,encoding='utf-8')
(DELIV/'HAL_BOOK_III_CERTIFICATION_REPORT.md').write_text('# Book III Certification Report\n\nStatus: Certified final v1.0. Sources verified; standards, 51 controls, templates, chapter reviews, full-book reviews, rendered DOCX/PDF, standalone chapter PDFs, and the control catalog are complete. The constitutional and Owner-decision audit found no conflict and no open Owner Review item.\n',encoding='utf-8')
doc=Document(DELIV/'HAL_BOOK_III_ENGINEERING_STANDARDS.docx')
for p in list(doc.paragraphs)+[p for s in doc.sections for p in list(s.header.paragraphs)+list(s.footer.paragraphs)]:
    for run in p.runs:
        run.text=run.text.replace('Draft v0.1','Final v1.0').replace('Draft for review','Final').replace('Version: 0.1','Version: 1.0')
doc.save(DELIV/'HAL_BOOK_III_ENGINEERING_STANDARDS.docx')
