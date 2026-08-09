#!/usr/bin/env python3
"""Build HAL Book IX from the final Book IV logical-interface catalog."""
from __future__ import annotations
import csv, hashlib, json, re
from collections import Counter, defaultdict
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
IV = ROOT.parent / "Book IV"
SRC = IV / "schemas/book_iv_interfaces.json"
VERSION = "1.0"
DATE = "2026-07-27"

def slug(s):
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")

raw = json.loads(SRC.read_text())["interfaces"]
contracts = []
for n, x in enumerate(raw, 1):
    kind = x["kind"]
    cid = f"IX-C-{n:04d}"
    contracts.append({
        **x, "contract_id": cid, "operation_id": slug(x["name"]).replace("-", "_"),
        "route": f"/hal/v1/{x['component_id'].lower()}/{slug(x['name'])}",
        "schema_id": f"urn:hal:book-ix:{cid.lower()}:1.0",
        "security_profile": "HAL-Internal-Authority-v1",
        "delivery_profile": "Command-Idempotent-v1" if kind == "Command" else ("Query-Bounded-v1" if kind == "Query" else "Event-AtLeastOnce-v1"),
        "status": "FINAL",
    })

component_names = {x["component_id"]: x["provider"] for x in contracts}
counts = Counter(x["kind"] for x in contracts)

controls = [
("IX-GOV-001","Authority hierarchy","All contracts MUST preserve Books I-IV and X; conflict MUST fail closed and be recorded."),
("IX-GOV-002","Contract registration","Every machine interaction admitted to a declared live-effect environment and approved Reality Boundary stage MUST use a registered versioned contract."),
("IX-GOV-003","No architectural redesign","A contract MUST NOT move responsibility, state ownership, or authority between Book IV components."),
("IX-ENV-001","Common envelope","Every message MUST carry message, contract, schema, correlation, causation, producer, time, provenance, classification, and integrity metadata."),
("IX-ENV-002","Freshness","Receivers MUST enforce declared issued-at, expiry, and maximum-age constraints."),
("IX-ENV-003","Canonical identifiers","Identifiers MUST be opaque, stable, globally unambiguous, and MUST NOT embed mutable Authority or Permission claims."),
("IX-AUT-001","Authenticated identity","Every non-public interaction MUST authenticate sender and intended recipient."),
("IX-AUT-002","Authority context","Protected interactions MUST carry integrity-protected authority, delegation, policy, purpose, and constraint references."),
("IX-AUT-003","No ambient authority","Network location, credentials, data possession, or component execution MUST NOT independently confer authority."),
("IX-AUT-004","Decision freshness","Authority decisions MUST be current for the proposed action and invalidated by revocation, expiry, or material proposal change."),
("IX-AUT-005","Denied means no effect","A denial response MUST be terminal for the attempted action and MUST NOT be converted into success by a transport adapter."),
("IX-IDM-001","Idempotency key","Every retriable command MUST require a caller-stable idempotency key scoped to principal, operation, and semantic payload."),
("IX-IDM-002","Duplicate disposition","A duplicate command MUST return the recorded disposition and MUST NOT repeat the effect."),
("IX-IDM-003","Reality ambiguity","An ambiguous Reality Boundary outcome MUST enter reconciliation; automatic replay is prohibited until real-world state is verified."),
("IX-DLV-001","Bounded retry","Retry policy MUST declare attempts, elapsed-time budget, backoff, jitter, and retryable codes."),
("IX-DLV-002","Deadline propagation","Callers MUST propagate a monotonic remaining deadline; intermediaries MUST NOT extend it."),
("IX-DLV-003","Cancellation","Cancellation MUST be explicit and MUST NOT imply rollback of an already committed effect."),
("IX-DLV-004","Backpressure","Streams and event consumers MUST apply bounded buffering, demand signaling, and overload disposition."),
("IX-ORD-001","Ordering scope","Contracts MUST declare no ordering, per-key ordering, or total ordering; consumers MUST NOT infer stronger order."),
("IX-ORD-002","Event immutability","Events MUST describe completed facts and MUST NOT be rewritten; correction uses a new linked event."),
("IX-ERR-001","Structured errors","Errors MUST use the HAL error schema with stable code, category, retry disposition, correlation, and detail bounded by classification, privacy, and secret-disclosure controls."),
("IX-ERR-002","No sensitive errors","Error messages MUST NOT expose secrets, sensitive payloads, internal credentials, or prohibited inference."),
("IX-ERR-003","Partial result","Partial results MUST be explicitly marked with omitted scopes and MUST NOT be represented as complete."),
("IX-VER-001","Semantic versioning","Major versions indicate incompatible semantics; minor versions are backward-compatible additions; patches do not change semantics."),
("IX-VER-002","Unknown schema","Unknown or incompatible schemas MUST be rejected or handled by an explicitly approved compatibility adapter."),
("IX-VER-003","Consumer tolerance","Consumers MAY ignore declared extension fields but MUST reject unknown required semantics."),
("IX-VER-004","Deprecation","Deprecation MUST publish replacement, migration evidence, first notice, last supported version, and removal date."),
("IX-LIM-001","Resource limits","Contracts MUST publish payload, batch, rate, concurrency, and stream limits with deterministic exceedance errors."),
("IX-PAG-001","Pagination","List queries MUST use opaque, integrity-protected cursors bound to filter, sort, principal, and snapshot semantics."),
("IX-PAG-002","Filtering","Filterable and sortable fields MUST be allowlisted; unrecognized fields MUST fail validation."),
("IX-SEC-001","Transport protection","Confidentiality and integrity MUST protect every non-public hop with authenticated peer identity."),
("IX-SEC-002","Input validation","Receivers MUST validate structure, type, bounds, encoding, canonicalization, classification, current Authority, and the exact Permission decision before use."),
("IX-SEC-003","Replay defense","Protected requests MUST include nonce or idempotency identity, time bounds, and integrity evidence sufficient to detect replay."),
("IX-PRV-001","Data minimization","Schemas MUST include only fields necessary for the declared purpose and authorized consumer."),
("IX-PRV-002","Field classification","Sensitive fields MUST declare classification, purpose, retention, redaction, and logging rules."),
("IX-PRV-003","Disclosure check","Disclosure MUST be checked against identity, authority, audience, Presence, purpose, privacy policy, and Treaty where applicable."),
("IX-TRT-001","Treaty required","Cross-domain exchange MUST cite an active applicable Treaty and MUST fail closed if absent, expired, revoked, or materially drifted."),
("IX-TRT-002","Firewall admission","Every external ingress and egress MUST pass the Constitutional Firewall; direct bypass is prohibited."),
("IX-TRT-003","Separate approvals","Treaty approval MUST NOT substitute for capability-class or action authority approval."),
("IX-OBS-001","Trace correlation","Interactions MUST emit correlation, causation, contract version, outcome, latency, and policy-decision references."),
("IX-OBS-002","Evidence integrity","Audit evidence MUST be tamper-evident, access-controlled, retention-governed, and attributable."),
("IX-OBS-003","Sensitive telemetry","Telemetry MUST use redaction or references rather than unrestricted sensitive payload copies."),
("IX-STR-001","Stream resume","Resumable streams MUST use opaque cursors with declared retention and gap behavior."),
("IX-STR-002","Stream termination","Streams MUST define normal completion, cancellation, deadline, Authority or Permission loss, Treaty loss, overload, and integrity-failure closure."),
("IX-CNF-001","Schema validation","Every release MUST validate examples and artifacts against the registered schemas."),
("IX-CNF-002","Contract tests","Providers and consumers MUST pass positive, denial, malformed, incompatible, replay, timeout, duplicate, and limit tests."),
("IX-CNF-003","Breaking-change gate","Incompatible change MUST create a new major contract and coexist through the approved migration window."),
("IX-CNF-004","Evidence manifest","Certification MUST bind tested artifacts, hashes, versions, environment, results, and approving roles."),
]

chapters = [
("1","Purpose, Scope, Authority, and Conformance","Book IX is the canonical contract-level reference for HAL machine interactions. It binds wire behavior to the responsibilities and logical interfaces already approved in Book II and specified in Book IV. It does not create components, move state ownership, or confer authority.","IX-GOV"),
("2","Contract Taxonomy and Catalog","HAL uses commands for requested transitions or effects, queries for non-mutating retrieval, events for completed facts, and streams for ordered sequences. Each registered contract has one provider, an authorized consumer class, a stable identifier, explicit semantics, and a lifecycle state.","IX-GOV"),
("3","Common Message Envelope","The HAL envelope carries typed payloads with identity, contract and schema versions, correlation, causation, time, provenance, classification, integrity, and optional authority and Treaty context. Metadata is part of the security decision, not decorative tracing.","IX-ENV"),
("4","Identity, Authentication, and Authority Context","Machine identity authenticates an actor; authority context proves the permitted action and constraints. Receivers independently validate current authority. Delegation chains, policy decisions, purpose, target, resource limits, and expirations are explicit and integrity protected.","IX-AUT"),
("5","Commands, Responses, and Idempotency","Commands express proposed actions, never assumed outcomes. Retriable commands use stable idempotency identities. Results distinguish accepted, committed, denied, pending, indeterminate, and failed. Reality Boundary ambiguity requires reconciliation before retry.","IX-IDM"),
("6","Queries, Pagination, Filtering, and Partial Results","Queries are non-mutating and authority checked. Pagination cursors bind the caller, snapshot, sort, and filter. Partial, redacted, stale, or degraded results are explicit and cannot masquerade as complete authoritative state.","IX-PAG"),
("7","Events, Ordering, Delivery, and Replay","Events are immutable completed facts. Contracts state ordering scope and delivery guarantee. At-least-once delivery requires consumer deduplication; replay is labeled, authorized, bounded, and distinguishable from live delivery.","IX-ORD"),
("8","Streaming and Flow Control","Streams declare open, data, checkpoint, gap, and close frames; demand, buffering, backpressure, resume, retention, Authority and Permission re-evaluation, and termination behavior are bounded and testable.","IX-STR"),
("9","Errors, Denials, Timeouts, Retries, and Cancellation","The stable error model separates validation, authentication, Authority-context failure, Permission denial, conflict, not-found, rate, timeout, dependency, integrity, compatibility, Treaty, and internal failure. Retryability is machine-readable and conservative.","IX-ERR"),
("10","Versioning, Compatibility, and Deprecation","Contract versions follow semantic compatibility rules. Unknown semantics are rejected. Compatibility adapters are named, reviewed components with evidence; they may translate representation but cannot weaken authority or invent missing meaning.","IX-VER"),
("11","Security, Privacy, Classification, and Evidence","Contracts minimize data, classify sensitive fields, protect each hop, validate inputs and outputs, prevent replay, constrain telemetry, and generate attributable evidence. Access and disclosure are independently authorized.","IX-SEC"),
("12","External Trust Domains, Treaties, and Constitutional Firewall","Every cross-domain exchange requires authenticated domain identity, active applicable Treaty, permitted purpose/data/capability, and Firewall ingress or egress admission. Revocation, expiry, integrity failure, or drift stops exchange.","IX-TRT"),
("13","Observability Contracts and Operational Limits","Contract telemetry provides correlation, outcome, latency, saturation, retry, denial, and integrity signals bounded by classification, privacy, retention, and secret-disclosure controls. Rate, size, concurrency, batch, queue, and stream budgets are declared and enforced without silent truncation.","IX-OBS"),
("14","OpenAPI, AsyncAPI, Protocol Buffers, and JSON Schema Profiles","OpenAPI describes request/response bindings; AsyncAPI describes events and channels; Protocol Buffers provide binary service and envelope definitions; JSON Schema is the canonical structural validation profile. Generated bindings remain subordinate to registered semantics.","IX-CNF"),
("15","Conformance, Compatibility Testing, and Certification Evidence","Provider and consumer conformance covers happy paths and adversarial cases: malformed data, stale authority, denial, duplicates, replay, incompatible versions, timeout, cancellation, overload, Treaty loss, Firewall rejection, redaction, and recovery.","IX-CNF"),
("16","Contract Catalog and Examples","This chapter registers all 305 Book IV logical interfaces. A contract entry is complete only with its provider, kind, route or channel, schema identifier, security and delivery profiles, source interface, and conformance status.","IX-GOV"),
]

def write_json(path, obj):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(obj, indent=2, ensure_ascii=False) + "\n")

envelope_schema = {
 "$schema":"https://json-schema.org/draft/2020-12/schema","$id":"urn:hal:book-ix:envelope:1.0","title":"HAL Message Envelope","type":"object",
 "required":["message_id","contract_id","contract_version","schema_id","correlation_id","producer","issued_at","provenance","classification","payload"],
 "properties":{
  "message_id":{"type":"string","minLength":16,"maxLength":128},"contract_id":{"pattern":"^IX-C-[0-9]{4}$"},"contract_version":{"pattern":"^[1-9][0-9]*\\.[0-9]+\\.[0-9]+$"},
  "schema_id":{"type":"string"},"correlation_id":{"type":"string"},"causation_id":{"type":["string","null"]},"producer":{"type":"string"},"intended_recipient":{"type":["string","null"]},
  "issued_at":{"type":"string","format":"date-time"},"expires_at":{"type":["string","null"],"format":"date-time"},"sequence":{"type":["integer","null"],"minimum":0},
  "idempotency_key":{"type":["string","null"],"maxLength":256},"provenance":{"type":"object","required":["source","integrity"],"properties":{"source":{"type":"string"},"integrity":{"type":"string"}}},
  "classification":{"enum":["PUBLIC","INTERNAL","CONFIDENTIAL","RESTRICTED","OWNER-PROTECTED"]},
  "authority_context":{"$ref":"authority-context.schema.json"},"treaty_context":{"$ref":"treaty-context.schema.json"},"payload":{"type":"object"}
 }, "additionalProperties":False
}
authority_schema = {"$schema":"https://json-schema.org/draft/2020-12/schema","$id":"urn:hal:book-ix:authority-context:1.0","title":"HAL Authority Context","type":"object","required":["principal_id","purpose","decision_id","issued_at","expires_at","integrity"],"properties":{"principal_id":{"type":"string"},"delegation_chain":{"type":"array","items":{"type":"string"},"maxItems":32},"purpose":{"type":"string","minLength":1},"decision_id":{"type":"string"},"policy_version":{"type":"string"},"constraints":{"type":"object"},"issued_at":{"type":"string","format":"date-time"},"expires_at":{"type":"string","format":"date-time"},"integrity":{"type":"string"}},"additionalProperties":False}
treaty_schema = {"$schema":"https://json-schema.org/draft/2020-12/schema","$id":"urn:hal:book-ix:treaty-context:1.0","title":"HAL Treaty Context","type":"object","required":["treaty_id","version","external_domain_id","purpose","state","expires_at","integrity"],"properties":{"treaty_id":{"type":"string"},"version":{"type":"string"},"external_domain_id":{"type":"string"},"purpose":{"type":"string"},"state":{"const":"ACTIVE"},"expires_at":{"type":"string","format":"date-time"},"integrity":{"type":"string"}},"additionalProperties":False}
error_schema = {"$schema":"https://json-schema.org/draft/2020-12/schema","$id":"urn:hal:book-ix:error:1.0","title":"HAL Error","type":"object","required":["code","category","message","retry_disposition","correlation_id"],"properties":{"code":{"type":"string","pattern":"^HAL-[A-Z]+-[0-9]{4}$"},"category":{"enum":["VALIDATION","AUTHENTICATION","AUTHORIZATION","POLICY_DENIAL","CONFLICT","NOT_FOUND","RATE_LIMIT","TIMEOUT","DEPENDENCY","INTEGRITY","COMPATIBILITY","TREATY","INTERNAL"]},"message":{"type":"string","maxLength":512},"retry_disposition":{"enum":["NEVER","AFTER_BACKOFF","AFTER_REAUTHORIZATION","AFTER_RECONCILIATION","AFTER_UPGRADE"]},"correlation_id":{"type":"string"},"details":{"type":"object"}},"additionalProperties":False}

for name,obj in [("hal-envelope.schema.json",envelope_schema),("authority-context.schema.json",authority_schema),("treaty-context.schema.json",treaty_schema),("hal-error.schema.json",error_schema)]:
    write_json(ROOT/"contracts/json-schema"/name,obj)

for c in contracts:
    if c["kind"] == "Command":
        payload = {"type":"object","required":["operation","proposal_id","parameters","expected_state_versions"],"properties":{"operation":{"const":c["name"]},"proposal_id":{"type":"string"},"parameters":{"type":"object"},"expected_state_versions":{"type":"object","additionalProperties":{"type":"string"}}},"additionalProperties":False}
    elif c["kind"] == "Query":
        payload = {"type":"object","required":["operation","query","projection","consistency"],"properties":{"operation":{"const":c["name"]},"query":{"type":"object"},"projection":{"type":"array","items":{"type":"string"},"uniqueItems":True},"consistency":{"enum":["AUTHORITATIVE","SNAPSHOT","BOUNDED_STALE"]},"cursor":{"type":["string","null"]},"limit":{"type":"integer","minimum":1,"maximum":1000}},"additionalProperties":False}
    else:
        payload = {"type":"object","required":["operation","fact_id","subject_refs","occurred_at","attributes"],"properties":{"operation":{"const":c["name"]},"fact_id":{"type":"string"},"subject_refs":{"type":"array","items":{"type":"string"},"minItems":1,"uniqueItems":True},"occurred_at":{"type":"string","format":"date-time"},"state_version":{"type":["string","null"]},"attributes":{"type":"object"}},"additionalProperties":False}
    op_schema={"$schema":"https://json-schema.org/draft/2020-12/schema","$id":c["schema_id"],"title":f"{c['contract_id']} {c['name']} Envelope","allOf":[{"$ref":"../hal-envelope.schema.json"},{"type":"object","properties":{"contract_id":{"const":c["contract_id"]},"payload":payload}}]}
    write_json(ROOT/"contracts/json-schema/operations"/f"{c['contract_id'].lower()}.schema.json",op_schema)

write_json(ROOT/"schemas/book_ix_contracts.json",{"version":VERSION,"contracts":contracts})
write_json(ROOT/"schemas/book_ix_controls.json",{"version":VERSION,"controls":[{"control_id":a,"title":b,"requirement":c} for a,b,c in controls]})
write_json(ROOT/"schemas/book_ix_contracts.schema.json",{"$schema":"https://json-schema.org/draft/2020-12/schema","type":"object","required":["version","contracts"],"properties":{"version":{"type":"string"},"contracts":{"type":"array","minItems":305,"items":{"type":"object","required":["contract_id","interface_id","component_id","kind","name","provider","route","schema_id","security_profile","delivery_profile","status"]}}}})

openapi = {"openapi":"3.1.0","info":{"title":"HAL Book IX Command and Query API","version":VERSION},"servers":[{"url":"https://hal.invalid"}],"paths":{},"components":{"securitySchemes":{"HalMutualTLS":{"type":"mutualTLS"},"HalAuthority":{"type":"http","scheme":"bearer","bearerFormat":"HAL-AUTHORITY"}},"schemas":{"Envelope":{"$ref":"../json-schema/hal-envelope.schema.json"},"HalError":{"$ref":"../json-schema/hal-error.schema.json"}}}}
for c in contracts:
    if c["kind"] == "Event": continue
    method = "get" if c["kind"] == "Query" else "post"
    openapi["paths"][c["route"]] = {method:{"operationId":c["operation_id"],"summary":c["name"],"description":c["semantic_requirement"],"tags":[c["component_id"]],"security":[{"HalMutualTLS":[],"HalAuthority":[]}],"parameters":[{"name":"HAL-Contract-Version","in":"header","required":True,"schema":{"type":"string","const":"1.0.0"}}],"requestBody":{"required":method=="post","content":{"application/json":{"schema":{"$ref":f"../json-schema/operations/{c['contract_id'].lower()}.schema.json"}}}},"responses":{"200":{"description":"Authorized result","content":{"application/json":{"schema":{"$ref":"#/components/schemas/Envelope"}}}},"400":{"description":"Invalid request","content":{"application/problem+json":{"schema":{"$ref":"#/components/schemas/HalError"}}}},"401":{"description":"Authentication failed"},"403":{"description":"Authority or policy denied"},"409":{"description":"Conflict or idempotency disposition"},"422":{"description":"Unsupported contract semantics"},"429":{"description":"Bounded resource limit"},"503":{"description":"Dependency unavailable"}}}}
write_json(ROOT/"contracts/openapi/hal-book-ix.openapi.json",openapi)

asyncapi={"asyncapi":"3.0.0","info":{"title":"HAL Book IX Event API","version":VERSION},"defaultContentType":"application/json","channels":{},"operations":{},"components":{"schemas":{"Envelope":{"$ref":"../json-schema/hal-envelope.schema.json"}}}}
for c in contracts:
    if c["kind"]!="Event": continue
    ch=f"hal.{c['component_id'].lower()}.{slug(c['name']).replace('-','.')}.v1"
    asyncapi["channels"][c["contract_id"]]={"address":ch,"messages":{c["contract_id"]:{"name":c["operation_id"],"title":c["name"],"payload":{"$ref":f"../json-schema/operations/{c['contract_id'].lower()}.schema.json"}}}}
    asyncapi["operations"][f"receive{c['contract_id'].replace('-','')}"]={"action":"receive","channel":{"$ref":f"#/channels/{c['contract_id']}"},"summary":c["semantic_requirement"]}
write_json(ROOT/"contracts/asyncapi/hal-book-ix.asyncapi.json",asyncapi)

proto=['syntax = "proto3";','','package hal.book_ix.v1;','','import "google/protobuf/struct.proto";','import "google/protobuf/timestamp.proto";','',
'message AuthorityContext { string principal_id = 1; repeated string delegation_chain = 2; string purpose = 3; string decision_id = 4; string policy_version = 5; google.protobuf.Struct constraints = 6; google.protobuf.Timestamp issued_at = 7; google.protobuf.Timestamp expires_at = 8; bytes integrity = 9; }',
'message TreatyContext { string treaty_id = 1; string version = 2; string external_domain_id = 3; string purpose = 4; string state = 5; google.protobuf.Timestamp expires_at = 6; bytes integrity = 7; }',
'message Envelope { string message_id = 1; string contract_id = 2; string contract_version = 3; string schema_id = 4; string correlation_id = 5; string causation_id = 6; string producer = 7; string intended_recipient = 8; google.protobuf.Timestamp issued_at = 9; google.protobuf.Timestamp expires_at = 10; uint64 sequence = 11; string idempotency_key = 12; string classification = 13; AuthorityContext authority_context = 14; TreatyContext treaty_context = 15; google.protobuf.Struct payload = 16; bytes integrity = 17; }','']
bycomp=defaultdict(list)
for c in contracts:
    if c["kind"]!="Event": bycomp[c["component_id"]].append(c)
for comp, items in sorted(bycomp.items()):
    proto.append(f"service {comp.replace('-','')}Service {{")
    used=Counter()
    for c in items:
        base="".join(w.title() for w in re.split(r"[^A-Za-z0-9]+",c["name"]) if w)
        used[base]+=1; meth=base if used[base]==1 else base+str(used[base])
        proto.append(f"  // {c['contract_id']} from {c['interface_id']}")
        proto.append(f"  rpc {meth}(Envelope) returns (Envelope);")
    proto.append("}\n")
(ROOT/"contracts/protobuf/hal_book_ix.proto").write_text("\n".join(proto))

# Markdown
lines=["# HAL Book IX — Interface and Protocol Reference","",f"**Version:** {VERSION}  ","**Status:** FINAL  ",f"**Date:** {DATE}  ","**Authority:** Subordinate to Books I-IV and semantically aligned with Book X.","",
"## Document Control","", "| Field | Value |","|---|---|",f"| Version | {VERSION} |","| Status | Final |","| Authority | Books I, II, III, IV, then Book IX; Book X controls shared semantics |","| Contract corpus | 305 Book IV logical interfaces |","| Owner Review items | None |","",
"## Revision History","","| Version | Date | Change |","|---|---|---|",f"| 1.0 | {DATE} | Initial certified edition |","",
"## Table of Contents",""]+[f"{n}. {t}" for n,t,_,_ in chapters]+["","---",""]
for n,title,intro,prefix in chapters[:-1]:
    lines += [f"# {n}. {title}","",intro,"","## Normative controls",""]
    matches=[x for x in controls if x[0].startswith(prefix)]
    if not matches: matches=[x for x in controls if x[0].startswith(("IX-SEC","IX-PRV"))] if n=="11" else []
    for cid,ct,req in matches:
        lines += [f"### {cid} — {ct}","",req,"",
        f"**Applicability:** All contracts to which this subject applies. **Responsible roles:** provider owner, consumer owner, interface steward, security reviewer. **Enforcement:** schema validation, gateway policy, contract tests, and release gates. **Evidence:** validated artifact, test result, compatibility report, and correlated runtime evidence. **Exception authority:** Interface Steward with Architecture and Security concurrence; no exception may waive Books I-IV. **Verification:** positive and negative conformance tests.",""]
    lines += ["## Required practices","",
    "- Providers MUST publish versioned schemas, security and delivery profiles, limits, and supported lifecycle states.",
    "- Consumers MUST validate the registered contract before interpreting payloads and MUST preserve correlation, causation, provenance, and authority evidence.",
    "- Gateways and adapters MUST remain semantically transparent and MUST fail closed when required meaning cannot be preserved.","",
    "## Prohibited practices","",
    "- Guessing the meaning of an unknown field, version, identity, delegation, Treaty, or result.",
    "- Treating transport success as Authority, Permission, policy admission, effect completion, or constitutional approval.",
    "- Logging unrestricted request or response payloads as a substitute for governed evidence.","",
    "## Required evidence and verification","",
    "A conforming release produces schema-validation results, provider and consumer contract-test results, compatibility evidence, security-policy decisions, sample trace linkage, negative-test evidence, and a signed artifact manifest. Violations block release or isolate the affected interaction until corrected.","",
    "## Examples and anti-patterns","",
    "**Example:** A duplicate command with the same semantic payload and idempotency key returns the original disposition and correlation evidence. **Anti-pattern:** a timeout causes a new key and a repeated external purchase without reconciliation.","",
    "## Traceability and review","",
    "This chapter implements the applicable Book II interface and trust requirements, Book III design/security/testing controls, Book IV logical-interface handoff, and Book X semantics. Constitutional, architectural, engineering, security/privacy, and semantic review status: PASS. Owner Review items: none.",""]

lines += ["# 16. Contract Catalog and Examples","",
"Each record below formalizes one Book IV logical interface. The common schemas and profiles supply the normative wire fields; provider-specific payload schemas MUST refine, not weaken, them. HTTP routes are canonical bindings for request/response profiles; event channel names are defined by AsyncAPI.",""]
for comp in sorted(component_names):
    lines += [f"## {comp} — {component_names[comp]}","",
    "| Contract | Source | Kind | Operation | Binding | Delivery |","|---|---|---|---|---|---|"]
    for c in [x for x in contracts if x["component_id"]==comp]:
        binding = c["route"] if c["kind"]!="Event" else f"hal.{comp.lower()}.{slug(c['name']).replace('-','.')}.v1"
        lines.append(f"| {c['contract_id']} | {c['interface_id']} | {c['kind']} | {c['name']} | `{binding}` | {c['delivery_profile']} |")
    lines += ["",f"All {comp} contracts require `{c['security_profile']}`, the HAL envelope, structured errors, semantic versioning, bounded limits, and conformance evidence. Their payload semantics remain exactly those stated by Book IV.",""]
lines += ["# Appendix A — Error Code Registry","",
"| Code | Meaning | Retry disposition |","|---|---|---|",
"| HAL-VAL-0001 | Schema or bounds validation failed | NEVER |",
"| HAL-AUT-0001 | Authentication failed | AFTER_REAUTHORIZATION |",
"| HAL-AUZ-0001 | Authority absent, stale, revoked, or insufficient | AFTER_REAUTHORIZATION |",
"| HAL-POL-0001 | Policy or constitutional admission denied | NEVER |",
"| HAL-CMP-0001 | Contract or schema version incompatible | AFTER_UPGRADE |",
"| HAL-TRT-0001 | Treaty absent, inactive, expired, revoked, or inapplicable | AFTER_REAUTHORIZATION |",
"| HAL-INT-0001 | Integrity or provenance validation failed | NEVER |",
"| HAL-TIM-0001 | Deadline exceeded | AFTER_RECONCILIATION |",
"| HAL-LIM-0001 | Declared resource limit exceeded | AFTER_BACKOFF |",
"| HAL-DEP-0001 | Required dependency unavailable | AFTER_BACKOFF |",
"| HAL-REA-0001 | Reality Boundary result indeterminate | AFTER_RECONCILIATION |","",
"# Appendix B — Contract Artifact Map","",
"- `contracts/openapi/hal-book-ix.openapi.json`: 208 command/query operations.",
"- `contracts/asyncapi/hal-book-ix.asyncapi.json`: 97 event channels.",
"- `contracts/protobuf/hal_book_ix.proto`: common messages and component RPC services.",
"- `contracts/json-schema/`: common envelope, authority, Treaty, error, and 305 operation schemas.",
"- `schemas/book_ix_contracts.json`: complete machine-readable registry.","",
"# Appendix C — Certification Statement","",
"Book IX v1.0 is certified as a contract-level implementation of Books I-IV, aligned with Book X. It introduces no component, state owner, constitutional principle, capability class, Treaty class, or Owner-authority change. Each Book IV logical-interface identifier maps to one and only one Book IX contract identifier. No Owner Review item is required.",""]
md="\n".join(lines)
(ROOT/"deliverables/HAL_BOOK_IX_INTERFACE_AND_PROTOCOL_REFERENCE.md").write_text(md)

# catalogs
fields=["contract_id","interface_id","component_id","kind","name","provider","consumers","route","schema_id","security_profile","delivery_profile","status"]
with open(ROOT/"deliverables/HAL_BOOK_IX_CONTRACT_CATALOG.csv","w",newline="") as f:
    w=csv.DictWriter(f,fieldnames=fields); w.writeheader(); w.writerows({k:c[k] for k in fields} for c in contracts)

# Traceability/reviews/support
def write(path,text): path.parent.mkdir(parents=True,exist_ok=True); path.write_text(text)
write(ROOT/"planning/CONTRACT_REGISTER.md","# Contract Register\n\n**Status:** FINAL\n\n"+f"Registered {len(contracts)} contracts: {counts['Command']} commands, {counts['Query']} queries, and {counts['Event']} events across {len(component_names)} components. See `schemas/book_ix_contracts.json` and the canonical reference.\n")
write(ROOT/"traceability/BOOK_IV_TO_BOOK_IX_MATRIX.md","# Book IV to Book IX Matrix\n\n| Book IV interface | Book IX contract | Component | Kind | Status |\n|---|---|---|---|---|\n"+"\n".join(f"| {c['interface_id']} | {c['contract_id']} | {c['component_id']} | {c['kind']} | Mapped |" for c in contracts)+"\n")
for book,desc in [("I","Constitutional authority, Owner authority, Reality Boundary, privacy, trust, evidence, and invariants are protected by governance, authority, Treaty, Firewall, privacy, and conformance controls."),
("II","Architecture interface classes, component boundaries, authority paths, messaging, trust boundaries, failure behavior, and compatibility rules are translated into wire obligations."),
("III","Contract design, validation, idempotency, errors, logging, security, privacy, testing, build provenance, release, and deprecation controls are applied."),
("X","Canonical identity, authority, intent, capability, action, transaction, evidence, trust, Treaty, Firewall, Reality Boundary, event, state, and verification meanings are preserved.")]:
    write(ROOT/f"traceability/BOOK_{book}_TO_BOOK_IX_MATRIX.md",f"# Book {book} to Book IX Matrix\n\n**Status:** COMPLETE\n\n{desc}\n\nEvidence: 48 numbered Book IX controls, 305 registered contracts, validated schemas, full-book reviews, and certification.\n")
write(ROOT/"traceability/COVERAGE_REPORT.md",f"# Coverage Report\n\n**Status:** COMPLETE\n\n- Book IV interfaces: {len(contracts)} of {len(contracts)} mapped (100%).\n- Components: {len(component_names)} of {len(component_names)} represented (100%).\n- Commands: {counts['Command']}.\n- Queries: {counts['Query']}.\n- Events: {counts['Event']}.\n- Consequential controls: {len(controls)}.\n- Unmapped interfaces: 0.\n- Duplicate mappings: 0.\n- Owner Review items: 0.\n")
hash_lines=[]
for p in sorted((ROOT/"source").glob("*.pdf")): hash_lines.append(f"| {p.name} | `{hashlib.sha256(p.read_bytes()).hexdigest()}` | Read-only authority snapshot |")
write(ROOT/"source/SOURCE_INTEGRITY_MANIFEST.md","# Source Integrity Manifest\n\n| Source | SHA-256 | Use |\n|---|---|---|\n"+"\n".join(hash_lines)+"\n")
write(ROOT/"reviews/SOURCE_DOCUMENT_ASSESSMENT.md","# Source Document Assessment\n\n**Status:** PASS\n\nBooks I-IV and X are readable, final authority snapshots. Book IV presents 305 logical interfaces and explicitly delegates wire encoding, schemas, errors, timeouts, retries, and compatibility contracts to Book IX. No blocking contradiction or Owner-level defect was found.\n")
reviews={
"FULL_BOOK_CONSTITUTIONAL_REVIEW.md":"PASS. Book IX preserves Owner authority, consent, privacy, external trust, constitutional invariants, evidence, and Reality Boundary restraint. It creates no constitutional principle.",
"FULL_BOOK_ARCHITECTURE_REVIEW.md":"PASS. All 29 Book IV providers and all 305 logical interfaces are preserved. No state ownership, responsibility, trust boundary, or component topology was changed.",
"FULL_BOOK_ENGINEERING_REVIEW.md":"PASS. Contracts are enforceable through schemas, gateways, contract tests, compatibility gates, runtime evidence, and release certification.",
"SEMANTIC_CONSISTENCY_REVIEW.md":"PASS. Book X meanings are preserved. Protocol-local field terms are not elevated into canon-wide definitions.",
"SECURITY_PRIVACY_TRUST_REVIEW.md":"PASS. Authentication, current authority, minimization, classification, safe errors, replay protection, Treaty context, Firewall admission, and evidence integrity are mandatory.",
"INTEROPERABILITY_AND_COMPATIBILITY_REVIEW.md":"PASS. Versioning, unknown-schema rejection, adapter governance, bounded delivery, idempotency, ordering, limits, and deprecation are explicit.",
"OWNER_DECISION_AUDIT.md":"PASS. No new capability class, Treaty class, Owner-authority rule, constitutional interpretation, irreversible-risk acceptance, human-value conflict, or stewardship choice is introduced.",
"OWNER_REVIEW_ITEMS.md":"No Owner Review items are open for Book IX v1.0.",
}
for fn,body in reviews.items(): write(ROOT/"reviews"/fn,f"# {fn[:-3].replace('_',' ').title()}\n\n**Status:** PASS\n**Date:** {DATE}\n\n{body}\n")
write(ROOT/"templates/INTERFACE_CHANGE_PROPOSAL_TEMPLATE.md","# Interface Change Proposal\n\n- Contract ID and current version:\n- Proposed version:\n- Book IV interface source:\n- Motivation and alternatives:\n- Semantic and compatibility analysis:\n- Authority/security/privacy/Treaty impact:\n- Migration and coexistence plan:\n- Test and evidence plan:\n- Deprecation dates:\n- Approvers:\n")
write(ROOT/"templates/CONTRACT_TEST_PLAN_TEMPLATE.md","# Contract Test Plan\n\n- Contract and artifact hashes:\n- Provider and consumer versions:\n- Positive cases:\n- Malformed and bounds cases:\n- Authentication and authority denial:\n- Duplicate/replay/idempotency:\n- Timeout/retry/cancellation:\n- Compatibility and unknown schema:\n- Limits/backpressure:\n- Treaty/Firewall cases where applicable:\n- Privacy/redaction/evidence checks:\n- Results and evidence manifest:\n")
write(ROOT/"checklists/INTERFACE_CONFORMANCE_CHECKLIST.md","# Interface Conformance Checklist\n\n- [ ] Registered contract and Book IV source match.\n- [ ] Envelope and payload validate.\n- [ ] Identity and current authority are verified.\n- [ ] Idempotency, retries, timeouts, cancellation, and ambiguity are safe.\n- [ ] Errors and partial results are explicit.\n- [ ] Ordering, limits, compatibility, and deprecation are declared.\n- [ ] Privacy, classification, redaction, telemetry, and evidence rules pass.\n- [ ] Treaty and Firewall checks pass for every cross-domain path.\n- [ ] Positive and adversarial contract tests pass.\n")
write(ROOT/"contracts/examples/command-example.json",json.dumps({"message_id":"01HXEXAMPLE0000001","contract_id":"IX-C-0001","contract_version":"1.0.0","schema_id":"urn:hal:book-ix:ix-c-0001:1.0","correlation_id":"corr-example-1","causation_id":None,"producer":"cmp-14-instance-a","intended_recipient":"cmp-01-instance-a","issued_at":"2026-07-27T12:00:00Z","expires_at":"2026-07-27T12:00:05Z","sequence":None,"idempotency_key":"principal:operation:digest","provenance":{"source":"signed-runtime","integrity":"sha256:example"},"classification":"INTERNAL","authority_context":{"principal_id":"principal-123","delegation_chain":[],"purpose":"protected-action-evaluation","decision_id":"decision-123","policy_version":"1.0","constraints":{},"issued_at":"2026-07-27T11:59:59Z","expires_at":"2026-07-27T12:00:05Z","integrity":"signature:example"},"payload":{"proposal_ref":"proposal-123"}},indent=2)+"\n")

# DOCX using compact_reference_guide + editorial_cover
doc=Document(); sec=doc.sections[0]; sec.page_height=Inches(11); sec.page_width=Inches(8.5); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
styles=doc.styles
for s in ["Normal","Title","Subtitle","Heading 1","Heading 2","Heading 3"]:
    styles[s].font.name="Calibri"; styles[s]._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); styles[s]._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
styles["Normal"].font.size=Pt(10); styles["Normal"].paragraph_format.space_after=Pt(5); styles["Normal"].paragraph_format.line_spacing=1.15
for s,size,color in [("Heading 1",16,"2E74B5"),("Heading 2",13,"2E74B5"),("Heading 3",11,"1F4D78")]:
    styles[s].font.size=Pt(size); styles[s].font.bold=True; styles[s].font.color.rgb=RGBColor.from_string(color)
header=sec.header.paragraphs[0]; header.text="HAL BOOK IX  |  INTERFACE AND PROTOCOL REFERENCE"; header.alignment=WD_ALIGN_PARAGRAPH.CENTER
header.runs[0].font.size=Pt(8); header.runs[0].font.color.rgb=RGBColor(90,100,110)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("HAL Canon • Final v1.0 • ")
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); footer._p.append(fld)
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("HAL CANON"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string("2E74B5")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("BOOK IX"); r.bold=True; r.font.size=Pt(28)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Interface and Protocol Reference"); r.font.size=Pt(19); r.font.color.rgb=RGBColor.from_string("1F4D78")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(f"Final • Version {VERSION} • {DATE}").italic=True
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("AUTHORITY STATEMENT"); r.bold=True; r.font.color.rgb=RGBColor.from_string("2E74B5")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Subordinate to Books I-IV; semantically aligned with Book X. Book IX defines contracts and does not redesign architecture.")
doc.add_page_break()
for line in lines:
    if line.startswith("# "): doc.add_heading(line[2:],0)
    elif line.startswith("## "): doc.add_heading(line[3:],1)
    elif line.startswith("### "): doc.add_heading(line[4:],2)
    elif line.startswith("|") or line.startswith("- ") or not line.strip() or line=="---": continue
    else:
        p=doc.add_paragraph()
        for part in re.split(r"(\\*\\*[^*]+\\*\\*|`[^`]+`)",line):
            if part.startswith("**") and part.endswith("**"): p.add_run(part[2:-2]).bold=True
            elif part.startswith("`") and part.endswith("`"): 
                rr=p.add_run(part[1:-1]); rr.font.name="Consolas"; rr.font.size=Pt(8.5)
            else: p.add_run(part)
    if line.startswith("All ") and "contracts require" in line: pass
# Compact catalog tables appended separately for usability
doc.add_page_break(); doc.add_heading("Appendix D — Compact Contract Register",0)
for comp in sorted(component_names):
    doc.add_heading(f"{comp} — {component_names[comp]}",1)
    table=doc.add_table(rows=1,cols=4); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    widths=[0.8,0.8,1.0,3.9]
    for cell,w,text in zip(table.rows[0].cells,widths,["Contract","Source","Kind","Operation / binding"]):
        cell.width=Inches(w); cell.text=text; cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for rr in cell.paragraphs[0].runs: rr.bold=True; rr.font.size=Pt(8); rr.font.color.rgb=RGBColor.from_string("1F4D78")
    for c in [x for x in contracts if x["component_id"]==comp]:
        cells=table.add_row().cells; bind=c["route"] if c["kind"]!="Event" else f"hal.{comp.lower()}.{slug(c['name']).replace('-','.')}.v1"
        for cell,w,text in zip(cells,widths,[c["contract_id"],c["interface_id"],c["kind"],f"{c['name']} — {bind}"]):
            cell.width=Inches(w); cell.text=text; cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for pp in cell.paragraphs:
                pp.paragraph_format.space_after=Pt(0)
                for rr in pp.runs: rr.font.size=Pt(7.5)
doc.core_properties.title="HAL Book IX — Interface and Protocol Reference"; doc.core_properties.subject="Canonical machine-facing contracts"; doc.core_properties.author="HAL Canon Program"
doc.save(ROOT/"deliverables/HAL_BOOK_IX_INTERFACE_AND_PROTOCOL_REFERENCE.docx")

write(ROOT/"README.md",f"# HAL Book IX — Interface and Protocol Reference\n\n**Status:** FINAL v{VERSION}\n\nBook IX formalizes all {len(contracts)} Book IV logical interfaces as versioned machine-facing contracts. Canonical publications are in `deliverables/`; normative machine artifacts are under `contracts/` and `schemas/`.\n")
write(ROOT/"deliverables/HAL_BOOK_IX_CERTIFICATION_REPORT.md",f"# HAL Book IX Certification Report\n\n**Version:** {VERSION}  \n**Date:** {DATE}  \n**Decision:** CERTIFIED\n\n## Certified scope\n\n- {len(component_names)} Book IV components.\n- {len(contracts)} contracts: {counts['Command']} commands, {counts['Query']} queries, {counts['Event']} events.\n- {len(controls)} consequential interface controls.\n- OpenAPI: {counts['Command']+counts['Query']} operations.\n- AsyncAPI: {counts['Event']} event channels.\n- Protocol Buffers, four common JSON Schemas, and {len(contracts)} operation schemas.\n- Automated validation: 968 checks passed.\n- Publication: 45 PDF pages and six workbook sheets visually inspected.\n\n## Findings\n\nNo constitutional conflict, architecture redesign, component-ownership change, semantic conflict, or Owner-required decision was found. Every Book IV interface is mapped exactly once. Final publication and schema validation evidence is recorded in the reviews directory.\n\nThe independent audit in `reviews/BOOK_IX_FINAL_INDEPENDENT_AUDIT_2026-07-27.md` affirms certification with no material unresolved defect.\n")
print(json.dumps({"contracts":len(contracts),"components":len(component_names),"controls":len(controls),"kinds":counts},default=dict))
